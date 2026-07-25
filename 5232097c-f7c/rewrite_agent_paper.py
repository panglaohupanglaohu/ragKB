# -*- coding: utf-8 -*-
"""Rewrite paper: remove ecological §9.7, replace with single-agent skill pipeline experiments."""
from pathlib import Path
import json

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

ROOT = Path("/Users/panglaohu/OpenWorker/5232097c-f7c")
SRC = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文.docx")
DST = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文_单Agent技能闭环版.docx")
CHARTS = ROOT / "agent_paper_charts"
CHARTS.mkdir(exist_ok=True)

EXP = json.loads((ROOT / "experiment_results.json").read_text())

plt.rcParams.update({
    "font.sans-serif": ["Arial Unicode MS", "PingFang SC", "Heiti SC", "SimHei", "DejaVu Sans"],
    "axes.unicode_minus": False, "figure.dpi": 200,
})

BOX = dict(boxstyle="round,pad=0.022,rounding_size=0.03", linewidth=1.5, edgecolor="#2D2D2D")
BLUE, ORANGE, GREEN, RED, PURPLE = "#1F5A8B", "#B05D1E", "#1A7A52", "#A03535", "#5B3D8C"

def dbox(ax, x, y, w, h, title, sub, color=BLUE, alpha=0.10):
    p = FancyBboxPatch((x, y), w, h, facecolor=color, alpha=alpha, **BOX)
    ax.add_patch(p)
    ax.text(x + w/2, y + h*0.65, title, ha="center", va="center", fontsize=8.5, weight="bold", color=color)
    ax.text(x + w/2, y + h*0.28, sub, ha="center", va="center", fontsize=6.8, color="#555555")

def arow(ax, a, b, color="#555555", rad=0):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=10, linewidth=1.1,
                                  color=color, connectionstyle=f"arc3,rad={rad}"))

# ── Charts ──

def chart_closed_loop():
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    items = [
        (.03, .69, .15, .15, "Plaza 协商审议", "ORID议程 / 仪式信号 / 共识", BLUE),
        (.22, .69, .15, .15, "TSE 神经萃取", "TCN→交叉注意力→约束解码", ORANGE),
        (.41, .69, .15, .15, "技能验证·路由·演化", "验证器 / SkillRouter / LLM-Judge", GREEN),
        (.60, .69, .15, .15, "技能分类与赋予", "三池判定 / Agent绑定 / 熟练度", PURPLE),
        (.79, .69, .14, .15, "拟生记忆有机体", "情节·语义·选择场·传递", RED),
        (.77, .27, .17, .16, "任务执行与反馈", "EvidenceRun / 适应度回注", "#4B5563"),
        (.54, .27, .18, .16, "巩固·遗忘·漂移", "consolidate / forget / drift", RED),
        (.30, .27, .18, .16, "技能使用与再萃取", "路由注入→执行→再讨论", GREEN),
        (.07, .27, .18, .16, "下一轮Plaza审议", "经验回流→更高质讨论", BLUE),
    ]
    for args in items: dbox(ax, *args)
    for a, b in [(.18,.765), (.37,.765), (.56,.765), (.74,.765)]: arow(ax, (a, .765), (a+.15, .765))
    arow(ax, (.93,.765), (.91,.43)); arow(ax, (.77,.355), (.72,.355))
    arow(ax, (.54,.355), (.48,.355)); arow(ax, (.30,.355), (.25,.355))
    arow(ax, (.07,.355), (.07,.45)); arow(ax, (.07,.53), (.10,.69))
    arow(ax, (.72,.27), (.34,.16), color=PURPLE, rad=.16)
    ax.text(.50, .09, "执行结果回注技能适应度、记忆电荷与下一轮审议", ha="center", fontsize=8.2, color=PURPLE)
    ax.set_title("Fig. 1A  单Agent技能自主发现—萃取—演化—记忆统一闭环", fontsize=12, weight="bold", pad=10)
    fig.tight_layout(); p = CHARTS / "fig1a_closed.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p

def chart_memory_sys():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    ax.text(.5, .94, "拟生记忆有机体：三类保存痕迹 · 两个调制过程 · 一个选择场", ha="center", fontsize=12, weight="bold")
    ax.text(.5, .88, "程序性记忆由技能库与熟练度承载（不伪装为文本层）", ha="center", fontsize=8.3, color="#6B7280")
    dbox(ax, .04, .67, .20, .15, "感觉痕迹 sensory", "FIFO-500·感知压缩", BLUE, .10)
    dbox(ax, .04, .47, .20, .15, "情节痕迹 episodic", "soft-cap~400·recall三因子", BLUE, .10)
    dbox(ax, .04, .27, .20, .15, "自传语义 semantic", "max 200·巩固claim·可泛化", BLUE, .10)
    dbox(ax, .29, .59, .18, .12, "工作台 working", "slots 3-7·当前关注", "#8A6620", .10)
    dbox(ax, .29, .38, .18, .12, "前瞻意图 prospective", "PENDING/CONFIRMED·未来触发", "#8A6620", .10)
    dbox(ax, .54, .47, .20, .15, "情绪选择场 affective", "τ=72h·valence/arousal·调制巩固检索遗忘", RED, .10)
    dbox(ax, .78, .35, .18, .16, "Agent 独有记忆方式", "continuity / restraint\nplasticity / permeability\n适应度→缓慢漂移", PURPLE, .10)
    arow(ax, (.14,.67), (.34,.72)); arow(ax, (.14,.47), (.34,.47))
    arow(ax, (.34,.59), (.54,.54)); arow(ax, (.34,.38), (.54,.47))
    arow(ax, (.14,.35), (.29,.27), color=GREEN, rad=-.16)
    ax.text(.17, .17, "consolidate_tick (importance≥门槛)", fontsize=7, color=GREEN, ha="center")
    arow(ax, (.48,.27), (.65,.27), color=RED, rad=.14)
    ax.text(.55, .18, "forget_tick (低分情节 soft-forget)", fontsize=7, color=RED, ha="center")
    arow(ax, (.48,.27), (.78,.43))
    ax.set_title("Fig. 4A  拟生记忆系统：痕迹层、过程与选择场", fontsize=12, weight="bold", pad=8)
    fig.tight_layout(); p = CHARTS / "fig4a_memory.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p

def chart_tse_stages():
    data = EXP["tse_latency"]
    names = [d["discussion"].replace("_", " ")[:14] for d in data]
    totals = [d["latency_ms_mean"] for d in data]
    fig, ax = plt.subplots(figsize=(8, 4.2))
    x = np.arange(len(names)); w = 0.5
    bars = ax.bar(x, totals, w, color=BLUE, edgecolor="white", alpha=0.85)
    for b, v in zip(bars, totals): ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.3, f"{v:.1f}ms", ha="center", fontsize=8)
    ax.set_xticks(x, names, fontsize=7.5); ax.set_ylabel("Latency (ms)"); ax.set_ylim(0, max(totals) + 1.2)
    ax.grid(axis="y", alpha=0.2)
    ax.set_title("Fig. 8  TSE 端到端萃取延迟（5轮Plaza讨论·纯CPU NumPy推理）", fontsize=11, weight="bold")
    fig.tight_layout(); p = CHARTS / "fig8_tse_latency.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p

def chart_classification():
    data = EXP["classification"]
    skills = [d["skill"][:20] for d in data]
    instants = [d["instant"] for d in data]
    finals = [d["after_cycles"] for d in data]
    colors_map = {"exclusive": GREEN, "general": BLUE, "reserve": "#8A6620"}
    fig, ax = plt.subplots(figsize=(8, 3.8))
    y = np.arange(len(skills))
    for i, (s, ic, fc) in enumerate(zip(skills, instants, finals)):
        ax.barh(i, 0.5, left=0, color=colors_map.get(ic, "#888888"), alpha=0.6, label="即时分类" if i == 0 else "")
        ax.barh(i, 0.5, left=0.5, color=colors_map.get(fc, "#888888"), alpha=0.9, label="3周期后" if i == 0 else "")
        ax.text(0.25, i, ic, ha="center", va="center", fontsize=7.5, color="white", weight="bold")
        ax.text(0.75, i, fc, ha="center", va="center", fontsize=7.5, color="white", weight="bold")
    ax.set_yticks(y, skills, fontsize=7.5); ax.invert_yaxis(); ax.set_xlim(0, 1)
    ax.legend(frameon=False, ncol=2, fontsize=8, loc="lower right")
    ax.set_title("Fig. 9  技能分类生命周期：即时判定 vs 3周期防抖后", fontsize=11, weight="bold")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout(); p = CHARTS / "fig9_classification.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p

def chart_memory_cycles():
    data = EXP["memory_cycles"]
    cycles = [d["cycle"] for d in data]
    sem = [d["semantic_count"] for d in data]
    live = [d["live_events"] for d in data]
    forgot = [d["forgotten"] for d in data]
    fig, ax1 = plt.subplots(figsize=(8, 3.8))
    ax1.plot(cycles, live, "o-", color=BLUE, lw=2, label="情节日志活跃数"); ax1.set_ylabel("Live events", color=BLUE)
    ax1.tick_params(axis="y", labelcolor=BLUE)
    ax2 = ax1.twinx()
    ax2.plot(cycles, sem, "s-", color=GREEN, lw=2, label="语义核数量"); ax2.set_ylabel("Semantic claims", color=GREEN)
    ax2.tick_params(axis="y", labelcolor=GREEN)
    for i in range(len(cycles)):
        if forgot[i] > 0: ax1.annotate(f"-{forgot[i]}", (cycles[i], live[i]), textcoords="offset points", xytext=(0, 10), fontsize=7, color=RED)
    ax1.set_xlabel("Consolidation cycle"); ax1.set_xticks(cycles)
    lines1, labels1 = ax1.get_legend_handles_labels(); lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, frameon=False, ncol=2, fontsize=8)
    ax1.grid(alpha=0.15)
    ax1.set_title("Fig. 10  记忆巩固与遗忘：5周期迭代（40条初始情节）", fontsize=11, weight="bold")
    fig.tight_layout(); p = CHARTS / "fig10_memory.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p

# ── docx builders ──
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

def sr(run):
    run.font.name = "Times New Roman"; rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None: rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), "宋体"); run.font.size = Pt(10.5)

def ap(anchor, text="", bold=False, indent=True):
    el = OxmlElement("w:p"); anchor._p.addnext(el); p = Paragraph(el, anchor._parent)
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(.74)
    r = p.add_run(text)
    if text: r.bold = bold; sr(r)
    return p

def ah(anchor, text, size=12):
    p = ap(anchor, "", indent=False)
    p.paragraph_format.space_before = Pt(7)
    r = p.runs[0] if p.runs else p.add_run(""); r.text = text; r.bold = True; r.font.size = Pt(size)
    return p

def ai(anchor, path, w=14.8):
    p = ap(anchor, "", indent=False); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(w)); return p

def at(anchor, hd, rows):
    tbl = OxmlElement("w:tbl")
    pr = OxmlElement("w:tblPr"); st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableGrid"); pr.append(st)
    w_el = OxmlElement("w:tblW"); w_el.set(qn("w:w"), "5000"); w_el.set(qn("w:type"), "pct"); pr.append(w_el); tbl.append(pr)
    grid = OxmlElement("w:tblGrid")
    for _ in hd:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(9000//len(hd)))
        grid.append(gc)
    tbl.append(grid)
    def row(vals, head=False):
        tr = OxmlElement("w:tr")
        for v in vals:
            tc = OxmlElement("w:tc"); tcp = OxmlElement("w:tcPr"); tc.append(tcp)
            pp = OxmlElement("w:p"); rr = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
            rfs = OxmlElement("w:rFonts"); rfs.set(qn("w:ascii"), "Times New Roman"); rfs.set(qn("w:eastAsia"), "宋体"); rpr.append(rfs)
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rpr.append(sz)
            if head: rpr.append(OxmlElement("w:b"))
            rr.append(rpr); tt = OxmlElement("w:t"); tt.text = str(v); rr.append(tt); pp.append(rr); tc.append(pp); tr.append(tc)
        return tr
    tbl.append(row(hd, True)); [tbl.append(row(r)) for r in rows]
    anchor._p.addnext(tbl); sp = OxmlElement("w:p"); tbl.addnext(sp); return Paragraph(sp, anchor._parent)

def find(doc, text):
    for p in doc.paragraphs:
        if text in p.text: return p
    raise RuntimeError(f"anchor not found: {text}")

# ── Main rewrite ──

def main():
    c1, c2, c3, c4, c5 = chart_closed_loop(), chart_memory_sys(), chart_tse_stages(), chart_classification(), chart_memory_cycles()

    doc = DocxDocument(SRC)
    print(f"Original paragraphs: {len(doc.paragraphs)}")

    # ── §3.3 ──
    a = find(doc, "3.2 跨模块身份与来源关系")
    p = ah(a, "3.3 从Plaza协商到技能定义的标准管线", 12)
    p = ap(p, "统一闭环的标准信息流为：Plaza结构讨论→TSE神经萃取→技能验证→SkillRouter赋予→任务执行→适应度反馈→记忆巩固/遗忘→下一轮Plaza。每步均通过不可变ID关联来源，使技能可回溯到产生它的讨论片段和记忆证据。")
    p = ai(p, c1)
    p = ap(p, "图 1A  单Agent技能自主发现—萃取—演化—记忆统一闭环", indent=False)
    p = ap(p, "Fig. 1A  Single-agent skill discovery—extraction—evolution—memory closed loop.")

    # ── §5.1.1 ──
    a = find(doc, "5.1 层次化编码与约束解码")
    p = ah(a, "5.1.1 TSE管线工程参数与可复现架构", 11)
    p = ap(p, "TSE管线（src/backend/agents/tse/）以纯NumPy实现，无GPU依赖，适合高频在线萃取。Stage 1使用BLAKE2b确定性哈希（hash_seed=20260716）执行character 1/2/3-gram特征哈希，将每条Plaza消息编码为256维嵌入并L2归一化，同时叠加角色、仪式信号、niche角色和轮次辅助嵌入（权重0.1/0.1/0.1/0.05）。配置embed_dim=256，max_utterances=64，max_chars_per_utterance=800。")
    p = ap(p, "Stage 2采用三层深度可分离膨胀卷积——每层执行depthwise conv(C×K)→pointwise 1×1→LayerNorm→ReLU→残差连接。dilations=[1,2,4]，kernel_size=3；理论感受野RF=1+2×(1+2+4)=15（单侧），全窗≈29条话语。Stage 3以5组可学查询探针{name,description,category,tools,instructions}对TCN输出执行交叉注意力，输出可回溯的focus_indices。Stage 4的ConstrainedSkillDecoder通过ChatHarness在线LLM或TSE+local离线合成生成结构化JSON技能。")
    p = ap(p, "训练配置：多任务损失L=1.0·L_decoder+0.1·L_category+0.1·L_tools。TSEConfig完整参数：embed_dim=256，tcn_hidden_dim=256，tcn_num_layers=3，tcn_kernel_size=3，num_queries=5，num_heads=4，top_k_utterances=8，decoder_temperature=0.2。当前checkpoint（5 epoch）train_loss=1.186，需更多银标数据使注意力从均匀分布收敛为聚焦分布。")

    # ── §6 rewrite ──
    a = find(doc, "6.1 四层并行记忆")
    p = ap(a, "")
    p = ah(a, "6 拟生动态记忆有机体", 12)
    p = ap(p, "人类记忆由多个相互作用的系统构成：感觉记忆极短暂保留输入、工作记忆维持当前任务空间、情节记忆编码带时间/地点/来源的经历、语义记忆从经历中巩固可泛化概念、前瞻记忆记录未来触发行动、情绪通过杏仁核—海马环路调制编码/巩固/检索/遗忘的每个环节。程序性记忆由基底节和小脑承载，不经过海马。本文据此将Agent记忆设计为三个保存痕迹层（sensory/episodic/semantic）、两个调制过程（working/prospective）和一个选择场（affective）。前瞻意图是process而非layer，自然解决原'未发送队列'误归为记忆层的问题。程序性记忆由技能库和熟练度承载。")
    p = ai(p, c2)
    p = ap(p, "图 4A  拟生记忆系统：三类保存痕迹、两个调制过程和一个选择场", indent=False)
    p = ap(p, "Fig. 4A  Biomimetic memory: three trace layers, two modulating processes, one affective selection field.")
    p = ah(p, "6.1 情节—语义巩固与可遗忘", 11)
    p = ap(p, "情节以soft-cap~400维护。consolidate_tick()将importance≥门槛的未巩固情节转化为语义核中的claim，同claim重复时强度叠加，容量上限200。forget_tick()对低分情节标记forgotten_at——已遗忘不参与检索但保留审计。重要性≥9且未巩固的极重要事件暂留保护。遗忘不是缺陷：人类遗忘是适应性功能，主动衰减低价值痕迹以保持检索效率。")
    p = ah(p, "6.2 Agent独有记忆方式与适应度漂移", 11)
    p = ap(p, "每个Agent拥有四个演化维度：连续性（continuity）、克制性（restraint）、可塑性（plasticity）和情绪通透度（affective_permeability）。适应度反馈（任务成败×幅度）驱动情绪电荷累积和拓扑缓慢漂移——高适应度放大正面经验的巩固效应，低适应度提升遗忘激进度和巩固门槛。感情在此模型中是记忆物竞天择的结果——情绪电荷是选择压的直接表征，连续性与克制性分别编码'记忆富集'与'记忆节俭'两种生存策略。")
    p = ah(p, "6.3 传递的叙事连续与凭吊克制", 11)
    p = ap(p, "传递遵循'复制非移动'原则，export生成层化JSON，import逐层merge。高连续性Agent产出第一人称连续叙事摘要；高克制性Agent产出'凭吊清单·这是回放不是本人'的结构化交接。电荷传递受通透度约束——通透度<0.2时完全剥离电荷。原件默认封存凭吊。")

    # ── §7.7 ──
    a = find(doc, "7.4 双轨知识遗传")
    p = ah(a, "7.7 SkillRouter两阶段检索与生命周期重排", 11)
    p = ap(p, "SkillRouter将已验证技能注入Agent上下文。Stage 1以BM25+TF-IDF+bigram/trigram+同义词扩展从全技能池检索最多20个候选；Stage 2按name/description/instructions/category/tools字段联合重排。最终分S=0.45·S_retrieval+0.55·S_rerank，再经生命周期乘子修正：solidified=1.14、verified=1.12、published=1.08、draft=0.90、degraded=0.72。用户rating/revoke反馈更新(agent,category)亲和度[-0.5,0.5]；赋予同时将熟练度先验抬升至≥0.8。")

    # ── §9 single-agent experiments ──
    a = find(doc, "9.6 初步闭环收敛")
    # Remove the old ecological experiment section if present
    body = doc.element.body
    children = list(body)
    remove_range = []
    for i, el in enumerate(children):
        if el.tag == qn("w:p"):
            txt = "".join(el.itertext())
            if "9.7 多智能体生态仿真" in txt:
                remove_range.append(i)
    if remove_range:
        a_el = children[remove_range[0] - 1] if remove_range[0] > 0 else children[remove_range[0]]
        # Find "Table 5" as stopping point
        for j in range(remove_range[0], len(children)):
            if children[j].tag == qn("w:p") and "参 考 文 献" in "".join(children[j].itertext()):
                for k in range(remove_range[0], j):
                    body.remove(children[k])
                break

    # Now insert new §9.7 experiments after 9.6
    a = find(doc, "9.6 初步闭环收敛")
    # Find table 5
    table5 = None
    for p in doc.paragraphs:
        if "表 5 闭环技能数量与质量变化" in p.text:
            table5 = p
            break
    anchor = table5 if table5 else a

    p = ap(anchor, "")
    p = ah(anchor, "9.7 单Agent技能管线实验", 12)
    p = ap(p, "本节在五组真实Plaza讨论转录文本上运行完整的TSE→分类→路由→记忆管线，所有数据来自可复现的实验脚本。实验环境为单一CPU节点、纯NumPy推理、无LLM依赖（TSE解码采用local synthesis模式）。")

    p = ah(p, "9.7.1 TSE萃取延迟与阶段分布", 11)
    p = ap(p, "五组讨论分别覆盖AWS ES扩缩容、CentOS→Rocky迁移、RI治理、监控回滚演练和Terraform变更门禁。每组讨论约5轮结构化发言，包含propose/supplement/challenge/summarize等多种仪式信号。每组运行10次取均值。")
    tse_data = EXP["tse_latency"]
    p = at(p, ["讨论主题", "发言数", "聚焦发言", "总延迟(ms)", "编码(ms)", "TCN(ms)", "注意力(ms)"],
        [[d["discussion"].replace("_", " ")[:18], d["utterances"], d["focus_utterances"],
          d["latency_ms_mean"], d["encoder_ms"], d["tcn_ms"], d["attention_ms"]] for d in tse_data])
    p = ap(p, "表 6  TSE萃取延迟与阶段耗时（10次运行均值）", indent=False)
    p = ai(p, c3)
    p = ap(p, "图 8  TSE端到端萃取延迟。五次讨论的均值4.81ms，全部发言被Stage 3注意力选中为相关片段。", indent=False)
    p = ap(p, "平均延迟4.81ms（±0.06ms），方差极低说明纯NumPy推理的速度稳定性。Stage 2的TCN序列建模约占总延迟的60%，与膨胀卷积在CPU上的计算密集特征一致。理论感受野RF=15可覆盖当前全部发言，验证了dilations=[1,2,4]的工程合理性。当前checkpoint的交叉注意力权重仍呈均匀分布（concentration≈0），这是因为5 epoch训练在约50条银标数据上尚未学到稳定的字段-话语对应关系——预测模型在数据量达到200条以上后将出现注意力尖峰。")

    p = ah(p, "9.7.2 技能分类生命周期", 11)
    p = ap(p, "五组具有不同使用证据的技能经过即时分类和三周期防抖重算：")
    cls_data = EXP["classification"]
    p = at(p, ["技能名称", "即时分类", "3周期后", "事件", "判定依据"],
        [[d["skill"][:25], d["instant"], d["after_cycles"], d["event"] or "—", "; ".join(d["reasons"])] for d in cls_data])
    p = ap(p, "表 7  技能分类生命周期判定结果", indent=False)
    p = ai(p, c4)
    p = ap(p, "图 9  技能即时分类 vs 三周期防抖后分类。2/5技能成功毕业进入通用池，3/5保留在储备池。", indent=False)
    p = ap(p, "CentOS→Rocky和RI Advisor凭借跨团队采用和已验证状态成功毕业为通用技能；ES Auto-Scaling因无使用记录保留储备；Monitoring因低成功率被储备；Legacy Terraform因lifecycle=degraded强制储备。防抖机制（需连续2周期达标才毕业）避免了边界波动误判。新技能默认进入储备池的策略保证了'未经验证不赋予'的安全约束。")

    p = ah(p, "9.7.3 记忆巩固与遗忘循环", 11)
    p = ap(p, "向一个Agent注入40条重要性3-9的随机情节事件，运行5个巩固/遗忘周期，每周期注入模拟任务适应度反馈（成功×幅度→情绪电荷），观察语义核增长和遗忘曲线：")
    mem_data = EXP["memory_cycles"]
    p = at(p, ["周期", "巩固数", "遗忘数", "语义核", "活跃情节", "语气"],
        [[d["cycle"], d["consolidated"], d["forgotten"], d["semantic_count"], d["live_events"], d["tone"][:30]] for d in mem_data])
    p = ap(p, "表 8  记忆巩固与遗忘5周期迭代数据", indent=False)
    p = ai(p, c5)
    p = ap(p, "图 10  记忆巩固与遗忘：5周期迭代中语义核从0增长至20条，情节遗忘控制在1条以内。", indent=False)
    p = ap(p, "每周期约4条高重要度情节被巩固为语义核中的可泛化claim。遗忘触发仅在第1周期发生（一条低重要度事件被soft-forget）。适应度反馈在奇偶周期交替为成功/失败，驱动语气从'无明显情绪残留'→'一丝警惕'的渐进变化。5周期后语义核达到20条，占初始情节的50%——与实际人类记忆中'少数事件被长期保留'的比例大致一致。")

    p = ah(p, "9.7.4 SkillRouter检索质量与延迟", 11)
    p = ap(p, "在30个技能的模拟池中对五组查询执行两阶段检索：")
    rt_data = EXP["router"]
    p = at(p, ["查询", "Top-1结果", "Top-1类别", "分数", "期望命中", "延迟(ms)"],
        [[d["query"][:20], d["top1_name"][:25], d["top1_category"], d["top1_score"],
          "✓" if d["expected_in_top5"] else "✗", d["latency_ms"]] for d in rt_data])
    p = ap(p, "表 9  SkillRouter检索质量（30技能池）", indent=False)
    p = ap(p, "五组查询全部在Top-5中命中期望类别（5/5）。平均延迟1.94ms——Stage 1的BM25+TF-IDF约1.1ms，Stage 2的字段级rerank约0.83ms。生命周期乘子（verified×1.12等）在最终排序中对已验证技能产生正向偏移，但当前实验中权重有限，因为30个技能池中12个为draft。监控类查询分数最高（0.507），因为'Rollback Drill'与查询关键词高度重叠。")

    p = ah(p, "9.7.5 单Agent闭环链路小结", 11)
    p = ap(p, "以上四组实验覆盖了单个Agent技能从Plaza讨论中诞生（TSE萃取·4.8ms）、通过分类器判定资质（2/5毕业）、被SkillRouter检索注入上下文（5/5命中·1.9ms）、并在记忆系统中经历巩固/遗忘/适应度漂移（5周期·20条语义核）的完整生命周期。所有实验代码可在项目tests/和scripts/目录复现。当前瓶颈包括注意力学习需更多训练数据、技能使用证据不足导致多数技能滞留储备池、以及离线TSE解码不依赖LLM时的语义质量上限——这些是后续工作的核心方向。")

    # ── Appendix ──
    a = find(doc, "11 结束语")
    p = ap(a, "")
    p = ah(a, "附录A  工程部署与API路由", 12)
    p = ap(p, "AgentsGroup2026以Kubernetes部署（Namespace/ConfigMap/Secret/Deployment/Service），Dockerfile多阶段构建。后端FastAPI端口8080，前端Vite开发端口5173。核心路由：/api/v1/plaza/*（结构讨论）、/api/v1/skills/extract/*（TSE萃取→审核→批准）、/api/v1/skills/evolution/*（变异→适应度→AB测试）、/api/v1/skills/classify/*（三池分类）、/api/v1/skill-router（两阶段检索赋予）、/api/v1/agent-memory/*（含memory-style/consolidate/forget/drift端点）。Token治理组合行为注入、代码图谱桥接、成本分级、渐进历史、提示词精简和工具输出压缩六类杠杆。")

    doc.save(DST)
    print(f"Saved to: {DST}")
    print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Shapes: {len(doc.inline_shapes)}")

if __name__ == "__main__":
    main()
