# -*- coding: utf-8 -*-
"""Rebuild the paper from the original formatting, incorporating ten technical modules
and a new empirical §9.7 with charts rendered from experiments_raw.tsv."""
from __future__ import annotations

import csv
import os
import shutil
from pathlib import Path
from statistics import mean

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path('/Users/panglaohu/OpenWorker/5232097c-f7c')
ORIGINAL = Path('/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文.docx')
DATA = Path('/Users/panglaohu/Downloads/AgentsGroup2026/experiments_raw.tsv')
OUT = ROOT / '协商审议_DARTNet_记忆遗传_统一闭环论文_重写含实验绘图版.docx'
CHART_DIR = ROOT / 'paper_charts'
CHART_DIR.mkdir(exist_ok=True)

# Font fallback suitable for macOS CJK rendering.
plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'PingFang SC', 'Heiti SC', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['figure.dpi'] = 180


def parse_data():
    with DATA.open(encoding='utf-8') as f:
        rows = list(csv.DictReader(f, delimiter='\t'))
    by_run = {}
    for row in rows:
        by_run.setdefault(row['run'], []).append(row)
    summary = {}
    for run, rr in by_run.items():
        r0 = rr[0]
        summary[run] = {
            'run': run,
            'n': len(rr),
            'skill': float(r0['run_skill_mean']),
            'collab': float(r0['run_collab_mean']),
            'bestT': int(r0['bestT']),
            'gens': int(r0['gens']),
            'regime': r0['regime'] or 'base',
            'tournament': r0['tournament'],
            'abundance': float(r0['abundance']) if r0['abundance'] else None,
            'predator_pressure': float(r0['predator_pressure']) if r0['predator_pressure'] else None,
            'drift_prob': float(r0['drift_prob']) if r0['drift_prob'] else None,
            'niche_capacity': float(r0['niche_capacity']) if r0['niche_capacity'] else None,
            'dominant': r0['dominant'],
            'n_dominant': int(r0['n_dominant']) if r0['n_dominant'] else 0,
        }
    return rows, summary


def chart_competition(summary):
    names = ['division', 'confrontation', 'mixed']
    runs = ['aws+build-division', 'aws+build-confrontation', 'aws+build-mixed']
    skills = [summary[r]['skill'] * 100 for r in runs]
    collabs = [summary[r]['collab'] * 100 for r in runs]
    x = np.arange(len(names)); w = 0.36
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    b1 = ax.bar(x-w/2, skills, w, label='技能行为占比', color='#2F6B9A')
    b2 = ax.bar(x+w/2, collabs, w, label='协作行为占比', color='#D98943')
    ax.set_xticks(x, ['分工\nDivision', '对抗\nConfrontation', '混合\nMixed'])
    ax.set_ylabel('运行均值 (%)')
    ax.set_title('不同群落组装模式下的技能与协作行为')
    ax.set_ylim(0, max(collabs + skills) + 4)
    ax.grid(axis='y', alpha=.25)
    ax.legend(frameon=False, ncol=2, loc='upper right')
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.35, f'{b.get_height():.1f}', ha='center', va='bottom', fontsize=8)
    fig.tight_layout()
    p = CHART_DIR / 'fig7_competition_modes.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig)
    return p


def chart_hump(summary):
    runs = ['habitat-harsh-aws+build','habitat-scarce-aws+build','habitat-pressure-aws+build','habitat-abundant-aws+build']
    d = sorted([summary[r] for r in runs], key=lambda x:x['abundance'])
    x = np.array([z['abundance'] for z in d]); y = np.array([z['skill']*100 for z in d]); c = np.array([z['collab']*100 for z in d])
    fig, ax = plt.subplots(figsize=(7.0,4.0))
    ax.plot(x,y, marker='o', lw=2.1, color='#2F6B9A', label='技能行为占比')
    ax.plot(x,c, marker='s', lw=1.8, ls='--', color='#D98943', label='协作行为占比')
    labels = ['严苛\n0.35','稀缺\n0.45','压力适中\n0.55','丰裕\n1.40']
    for xi, yi, lab in zip(x,y,labels):
        ax.annotate(lab, (xi,yi), xytext=(0,9), textcoords='offset points', ha='center', fontsize=8)
    ax.set_xlabel('资源丰度（abundance）')
    ax.set_ylabel('运行均值 (%)')
    ax.set_title('资源丰度与群体行为：中等压力下的技能涌现峰值')
    ax.grid(alpha=.25); ax.legend(frameon=False)
    ax.set_ylim(0, max(max(y),max(c))+4)
    fig.tight_layout()
    p = CHART_DIR / 'fig8_abundance_hump.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig)
    return p


def chart_agent_profiles(rows):
    selected = [r for r in rows if r['run'] in ('aws+build-confrontation','aws+build-mixed')]
    labels = [f"{r['agent']}\n{r['run'].split('-')[-1]}" for r in selected]
    skill = [float(r['skill_pct']) * 100 for r in selected]
    collab = [float(r['collab_pct']) * 100 for r in selected]
    residual = [float(r['residual_pct']) * 100 for r in selected]
    y = np.arange(len(selected))
    fig, ax = plt.subplots(figsize=(7.2,5.0))
    ax.barh(y, skill, label='技能', color='#2F6B9A')
    ax.barh(y, collab, left=skill, label='协作', color='#D98943')
    ax.barh(y, residual, left=np.array(skill)+np.array(collab), label='残余', color='#C7CDD4')
    ax.set_yticks(y, labels, fontsize=8); ax.invert_yaxis(); ax.set_xlim(0,100)
    ax.set_xlabel('Agent行为构成 (%)'); ax.set_title('对抗与混合模式下的Agent行为剖面')
    ax.legend(frameon=False, ncol=3, loc='lower right'); ax.grid(axis='x', alpha=.2)
    fig.tight_layout()
    p = CHART_DIR / 'fig9_agent_profiles.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig)
    return p


def style_run_like(run, reference=None):
    if reference:
        run._element.rPr = reference._element.rPr
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)


def add_paragraph_after(paragraph, text='', style_name='Normal', bold=False, reference_run=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    try: p.style = style_name
    except KeyError: pass
    r = p.add_run(text); r.bold = bold; style_run_like(r, reference_run)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74) if style_name == 'Normal' else None
    pf.space_after = Pt(4)
    return p


def add_heading_after(paragraph, text, level=2):
    p = add_paragraph_after(paragraph, '', 'Normal')
    r = p.runs[0]; r.text = text; r.bold = True; r.font.size = Pt(12 if level==2 else 11)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    return p


def add_table_after(paragraph, headers, data_rows):
    # Create at document end, then relocate its XML after paragraph.
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for j,h in enumerate(headers):
        cell=table.cell(0,j); cell.text=str(h)
        for run in cell.paragraphs[0].runs: run.bold=True; style_run_like(run)
    for row in data_rows:
        cells=table.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=str(val)
            for run in cells[j].paragraphs[0].runs: style_run_like(run)
    paragraph._p.addnext(table._tbl)
    spacer=add_paragraph_after(Paragraph(table._tbl, paragraph._parent), '', 'Normal')
    return spacer


def add_image_after(paragraph, image_path, width_cm=14.5):
    p = add_paragraph_after(paragraph, '', 'Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return p


def insert_block(doc, anchor_keyword, build_fn, occurrence=0):
    hits=[p for p in doc.paragraphs if anchor_keyword in p.text]
    if len(hits)<=occurrence:
        print('missing anchor:',anchor_keyword); return
    build_fn(hits[occurrence])


def technical_blocks(doc):
    def dart(anchor):
        p=add_heading_after(anchor,'5.1.1 工程实现参数与可复现TSE管线',3)
        p=add_paragraph_after(p,'DART-Net在工程实现中对应TSE（TCN-Skill-Extractor）管线。Stage 1采用BLAKE2b确定性哈希（hash_seed=20260716）生成character 1/2/3-gram特征，并将内容向量与角色、仪式信号、niche角色和轮次嵌入融合为256维话语表示。配置为embed_dim=256、max_utterances=64、max_chars_per_utterance=800；该设计确保离线与在线运行具有确定性。')
        p=add_paragraph_after(p,'Stage 2使用三层深度可分离膨胀一维卷积：kernel_size=3，dilations=[1,2,4]，每层执行depthwise卷积、pointwise投影、LayerNorm、ReLU和残差连接。理论感受野为RF=1+2×Σd=15（单侧），对应双侧全窗约29条话语。Stage 3使用5个字段查询探针{name, description, category, tools, instructions}执行交叉注意力，输出可回溯的focus_indices；top_k_utterances=8。')
        add_paragraph_after(p,'Stage 4约束解码支持ChatHarness在线生成、异步chat_fn调用和TSE+local离线合成三种后端。JSON输出经schema校验并允许grammar_retry=1次修复。训练目标为L=1.0L_decoder+0.1L_category+0.1L_tools，推理时decoder_temperature=0.2，最少/最多输出技能数为1/8。')
    insert_block(doc,'5.1 层次化编码与约束解码',dart)

    def extraction(anchor):
        p=add_heading_after(anchor,'5.2.1 三算法技能萃取与知识簇预处理',3)
        p=add_paragraph_after(p,'生产萃取器将同一讨论同时投射到三个互补的逆向工程算子：(1) 去语境化（De-contextualization）剥离服务名、错误码和业务专名，提取可迁移的动词—关系结构；(2) 反面模式（Anti-pattern）将事故、异议和失败日志反转为约束与避坑规则；(3) 关键路径与最小动作集（Critical Path & Minimum Action Set）从长叙事中保留3—5个决定成败的动作和验证点。')
        add_paragraph_after(p,'对于长文档，系统先按标题或空行分块，再以TF-IDF余弦相似度执行凝聚聚类，最多形成5个知识簇；每簇独立萃取、再经slug去重和审核队列合并。审核队列以pending→llm_prefilling→ready_for_review→approved/rejected/error状态机持久化，保留source_meta、原始响应和审核人信息。')
    insert_block(doc,'5.2 技能Schema与生命周期扩展',extraction)

    def memory(anchor):
        p=add_heading_after(anchor,'6.1.1 四层记忆的工程算法',3)
        p=add_paragraph_after(p,'EpisodicLog记录{id,t,subject,action,detail,place,importance,tags,lastAccessAt}。召回函数采用score=recency+importance+relevance，其中recency=0.995^hours；relevance由query与事件文本的character bigram重叠度计算。该机制同时编码时新性、任务重要性和语义相关性。')
        p=add_paragraph_after(p,'PerceptionStream为容量500的FIFO缓冲区；compress()把多模态刺激汇总为一条情景日志，并输出时间窗、模态频数和fear均值。IntentionQueue维护pending/confirmed/dropped三态意图，支持drop、escalate、keep三类超时策略，并提供dueAt驱动的到期排序。')
        add_paragraph_after(p,'AffectResidue维护valence∈[-1,1]、arousal∈[0,1]与标签强度。其指数衰减时间常数为72小时，重复感受按max(旧强度×1.2,新强度)更新，避免高强度风险信号被快速稀释。')
    insert_block(doc,'6.1 四层并行记忆',memory)

    def transfer(anchor):
        p=add_heading_after(anchor,'6.2.1 Seal—Will—Export—Import的可审计传递协议',3)
        p=add_paragraph_after(p,'传递实现遵守“复制而非移动、保留而非静默销毁”原则。seal生成schema=ag.legacy/v1的遗体快照，保存完整日志、感知摘要、意图和情绪快照；will声明受益方、迁移层、keep_memorial和handover_intentions策略。export生成schema=ag.memory/v1的层化JSON对象。')
        p=add_paragraph_after(p,'import采用逐层merge而非整盘覆盖：日志追加“传递继承”tag和来源标记；感知流追加并截断至500条；意图根据ask_new_owner/auto/drop策略交接；情绪标签按max强度合并、valence/arousal按50%加权融合。系统写入transfer_id、source/beneficiary审计记录和importance=9的“记忆继承”事件。')
        add_paragraph_after(p,'该协议将“身份延续”限制为显式来源分区和可追溯回放，避免把继承数据误表述为后继Agent的原生经历。')
    insert_block(doc,'6.2 Seal—Will—Export—Import遗传链',transfer)

    def evolution(anchor):
        p=add_heading_after(anchor,'7.5 LLM-as-Judge演化评估管线',3)
        p=add_paragraph_after(p,'每个候选变体经过simulate→judge双阶段：模拟器按技能instructions执行评测任务；Judge独立评价instruction_following、output_quality与conciseness三个0—1维度。复合适应度为F_judge=0.4×following+0.4×quality+0.2×conciseness。SkillFitnessReport聚合均分、逐例输出和composite<0.5的失败集合。')
        add_paragraph_after(p,'为抑制“以冗长换得表面完整”的变异，系统引入长度惩罚：当变体/祖先文本长度ratio≤1无惩罚；1<ratio≤1.5在线性区间最多扣减0.2；ratio>1.5则分数归零。')
    insert_block(doc,'7.3 变异、选择、组合与退役',evolution)

    def classifier(anchor):
        p=add_heading_after(anchor,'7.6 三池分类器与防抖毕业机制',3)
        p=add_paragraph_after(p,'技能分类不是人工静态标签。强制储备条件包括：lifecycle=degraded、已有使用但effectiveness<0.4、90天未使用或total_uses=0。通用技能要求至少2个团队采用或至少2类场景验证通过，并要求gate_ok；特有技能要求单团队使用占比≥0.8、effectiveness≥0.6且满足rubric。')
        p=add_table_after(p,['参数','值','用途'],[['EXCLUSIVE_TEAM_SHARE','0.8','特有技能单团队使用占比'],['EXCLUSIVE_MIN_EFFECTIVENESS','0.6','特有技能最低效果'],['GENERAL_MIN_TEAMS','2','通用技能跨团队采用'],['GENERAL_MIN_CATEGORIES','2','通用技能跨场景验证'],['RESERVE_MAX_EFFECTIVENESS','0.4','低效果强制进入储备'],['STALE_DAYS','90','长期未使用回收'],['GRADUATE_STREAK','2','连续达标才毕业'],['DEMOTE_GRACE','1','降级宽限周期']])
        add_paragraph_after(p,'为避免边界波动，毕业需连续2个周期达标，降级允许1个宽限周期。萃取得到的新技能通过seed_reserve_from_extraction()首先写入储备池，随后以使用、验证和跨团队证据驱动毕业。')
    insert_block(doc,'7.4 双轨知识遗传',classifier)

    def deploy(anchor):
        p=add_heading_after(anchor,'附录A 工程部署、接口与Token治理',2)
        p=add_paragraph_after(p,'项目以Kubernetes资源部署：Namespace隔离、ConfigMap管理模型池和运行配置、Secret承载密钥、主服务Deployment承载Plaza/TSE/记忆/技能注册表，Teams Deployment按团队横向扩展，Service提供HTTP与实时流入口。核心FastAPI路由覆盖Plaza、agent-memory、skills/extract、skills/evolution、skills/classify与cost/token治理。')
        p=add_paragraph_after(p,'Token治理层在每次模型调用前组合行为注入、代码图谱桥接、成本分级、渐进历史、提示词精简和工具输出压缩六类杠杆。SavingsStore记录每次优化节约的token和按模型单价折算的成本，为闭环适应度中的成本项提供可审计数据。')
    insert_block(doc,'11 结束语',deploy)


def experiment_block(doc, charts, summary):
    def build(anchor):
        p=add_heading_after(anchor,'9.7 多智能体生态仿真实验：群落组装、资源压力与行为涌现',2)
        p=add_paragraph_after(p,'为补充§9.6的短周期闭环收敛实验，本文进一步分析AgentsGroup2026实验记录experiments_raw.tsv。数据包含9个运行组、45条Agent级观测，每组5个Agent。每条记录给出Agent在运行中用于技能、协作与残余行为的比例，并记录竞合编排、代际数、最佳轮次及生境参数。以下分析仅基于该文件中可复核的原始字段，结论应理解为原型系统的探索性证据。')
        p=add_heading_after(p,'9.7.1 实验设计与指标',3)
        p=add_paragraph_after(p,'实验分为两部分。群落组装部分比较solo、division、confrontation和mixed四种编排：solo为同域团队独立运行，division强调跨团队职责分工，confrontation引入竞争性选择，mixed允许竞合并存。生境部分固定为division编排，设置harsh、scarce、pressure和abundant四种资源环境，资源丰度分别为0.35、0.45、0.55和1.40。主要指标为运行均值skill_pct与collab_pct，分别表示技能使用和协作交互的行为份额。')
        p=add_table_after(p,['运行组','编排/环境','Agent数','技能均值','协作均值','代际','最佳T'],[
            ['aws-solo-division','solo / AWS','5','0.0%','0.9%','1','45'],
            ['build-solo-division','solo / Build','5','10.7%','15.2%','2','75'],
            ['aws+build-division','division / base','5','6.8%','15.0%','1','69'],
            ['aws+build-confrontation','confrontation / base','5','8.8%','10.6%','1','77'],
            ['aws+build-mixed','mixed / base','5','5.0%','8.2%','9','57'],
            ['habitat-harsh','division / harsh','5','4.8%','2.5%','1','60'],
            ['habitat-scarce','division / scarce','5','3.7%','13.6%','1','61'],
            ['habitat-pressure','division / pressure','5','7.9%','13.2%','1','67'],
            ['habitat-abundant','division / abundant','5','2.8%','4.9%','2','86'],
        ])
        p=add_heading_after(p,'9.7.2 群落组装模式的比较',3)
        p=add_image_after(p, charts[0], 14.7)
        p=add_paragraph_after(p,'图7 不同群落组装模式下的技能与协作行为（数据源：experiments_raw.tsv）。')
        p=add_paragraph_after(p,'图7显示，在跨团队基础配置中，confrontation的技能行为均值为8.8%，高于division的6.8%和mixed的5.0%；而division的协作均值为15.0%，高于confrontation的10.6%与mixed的8.2%。这说明竞争性筛选在当前参数下更倾向于放大可区分的技能行为，职责分工则更有利于维持协作密度。mixed组经历9代后出现5项dominant技能，提示更长迭代可能促进跨团队技能占优，但本数据未提供重复运行，不能将其解释为显著性结论。')
        p=add_heading_after(p,'9.7.3 资源丰度与技能涌现',3)
        p=add_image_after(p, charts[1], 14.7)
        p=add_paragraph_after(p,'图8 资源丰度、生境压力与群体行为的关系。')
        p=add_paragraph_after(p,'图8呈现非单调关系：资源丰度从0.35（harsh）增至0.55（pressure）时，技能行为由4.8%升至7.9%；丰度进一步增至1.40（abundant）后，技能行为下降至2.8%。在现有四个设定点上，该结果与“中等环境压力更有利于技能探索与选择”的倒U型假说一致。协作行为在稀缺与适中压力下均保持较高水平（13.6%与13.2%），而在严苛与丰裕两端下降，表明过强约束与过度宽松都可能弱化群体的有效互动。')
        p=add_heading_after(p,'9.7.4 Agent级行为构成与解释边界',3)
        p=add_image_after(p, charts[2], 14.7)
        p=add_paragraph_after(p,'图9 对抗与混合模式下的Agent级行为构成；skill、collab与residual三者按记录定义相加为100%。')
        p=add_paragraph_after(p,'图9揭示运行均值背后的异质性：在confrontation中，aws_lead、aws_mon等Agent呈现较高的技能/协作份额，而部分Build Agent仍主要落在residual类别；mixed组中技能优势分布更分散。该现象支持将技能种群视为具有角色依赖的生态结构，而非均匀扩散的全局资源。由于样本规模为每组5个Agent，且实验文件未提供独立随机种子重复、方差或置信区间，本文不报告显著性检验；后续应采用多种子重复、预注册阈值和跨域任务复验。')
        p=add_heading_after(p,'9.7.5 与统一闭环的关联',3)
        add_paragraph_after(p,'生态实验将§8中的耦合动力学具体化：竞合编排影响Plaza中角色交互和选择压力，进而改变技能种群的使用/协作构成；生境参数相当于对执行成本、失败风险与资源预算的外生扰动；dominant技能及其谱系可由技能分类器、适应度报告和记忆传递记录回溯。因此，生态数据并非独立展示层，而是对“审议—萃取—执行—演化—遗传”闭环在群体尺度上的补充观测。')
    insert_block(doc,'9.6 初步闭环收敛',build)


def main():
    rows, summary = parse_data()
    charts=[chart_competition(summary),chart_hump(summary),chart_agent_profiles(rows)]
    doc=Document(ORIGINAL)
    technical_blocks(doc)
    experiment_block(doc, charts, summary)
    doc.save(OUT)
    print('saved',OUT)
    print('charts',*[str(x) for x in charts],sep='\n')

if __name__=='__main__': main()
