# -*- coding: utf-8 -*-
"""
Complete paper enrichment: reads the existing 协商审议_DARTNet_记忆遗传_统一闭环论文.docx
and injects detailed technical content derived from the actual AgentsGroup2026 codebase
(src/backend/agents/), then saves the enriched docx.
"""

from __future__ import annotations

import copy
import re
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文.docx"
DST = "/Users/panglaohu/OpenWorker/5232097c-f7c/协商审议_DARTNet_记忆遗传_统一闭环论文_完善版.docx"


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"] if f"Heading {level}" in [s.name for s in doc.styles] else doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12) if level == 2 else Pt(11)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_code_block(doc, code_text):
    """Add a monospaced code block."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Set paragraph spacing tight
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(12)


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table


def find_paragraph_by_keyword(doc, keyword, max_results=3):
    """Find paragraph indices containing keyword."""
    results = []
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            results.append(i)
            if len(results) >= max_results:
                break
    return results


def insert_after_paragraph(doc, after_idx, elements):
    """
    Insert new paragraphs/tables after the given paragraph index.
    elements is a list of functions that take `doc` and add content.
    We work by finding the XML element and inserting after it.
    """
    if after_idx >= len(doc.paragraphs):
        return
    ref_para = doc.paragraphs[after_idx]
    ref_element = ref_para._element

    # We'll create a temporary doc, add elements there,
    # then move their XML elements into the right position.
    temp_doc = DocxDocument()
    for elem_fn in elements:
        elem_fn(temp_doc)

    # Move all body children from temp_doc after ref_element
    parent = ref_element.getparent()
    insert_after = ref_element
    for child in list(temp_doc.element.body):
        insert_after.addnext(child)
        insert_after = child


# ── Enrichment Content Functions ──────────────────────────────────────────

def enrich_dart_arch_detail(doc):
    """5.1: Add precise DART-Net / TSE implementation details."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 5.1.1 TSE编码器精确架构】")
    add_body(doc,
        "TSE编码器（src/backend/agents/tse/）采用纯NumPy推理引擎，无GPU依赖。Stage 1话语编码使用BLAKE2b确定性哈希算法"
        "（hash_seed=20260716），对每条消息的content字段执行character 2/3-gram特征哈希，"
        "将文本映射为256维浮点嵌入向量。编码器同时整合辅助嵌入：角色嵌入（12类）、仪式信号嵌入（10类）、"
        "壁龛角色嵌入（6类）和轮次嵌入（16步），各以权重0.1、0.1、0.1、0.05叠加到主嵌入向量上，最终L2归一化。"
        "embed_dim=256，max_utterances=64，max_chars_per_utterance=800。"
    )
    add_body(doc,
        "Stage 2的TCN序列编码（tcn.py）采用深度可分离膨胀卷积（depthwise-separable dilated convolution）。"
        "每个DilatedConvBlock执行：depthwise卷积（C×K）→ pointwise 1×1卷积（C_out×C_in）→ LayerNorm → ReLU → 残差连接。"
        "三层堆叠的膨胀因子dilations=[1,2,4]，卷积核kernel_size=3。顶层神经元的理论感受野为：\n"
        "    RF = 1 + 2 × Σ_{i=0}^{L-1} d_i = 1 + 2 × (1+2+4) = 15（单侧），全窗≈29步。\n"
        "输入投影 W_in ∈ ℝ^{256×256}，输出投影 W_out ∈ ℝ^{256×256}。权重初始化采用He式缩放："
        "depthwise scale = √(2/k)×0.5，pointwise scale = √(2/C)×0.5。mask向量用于padding位置置零。"
    )
    add_body(doc,
        "Stage 3跨模态注意力融合定义5组可学习技能查询探针（skill query probes）：\n"
        "    Q = {q_name, q_desc, q_category, q_tools, q_instr}，每个 q_* ∈ ℝ^{256}\n"
        "探针通过交叉注意力CrossAttn(Q=q_k, K=h_tcn, V=h_tcn)与TCN输出交互，"
        "产生skill_repr_k ∈ ℝ^{256}。注意力的副产品Focus indices标识了被每个探针最关注的utterance索引。"
        "冷启动时，各探针的注意力初始化融合了FIELD_KEYWORD_SEEDS定义的关键词先验（权重0.3），"
        "例如q_tools偏向于包含'kubectl''aws''terraform''python''boto3'等工具名的utterance。"
    )
    add_body(doc,
        "Stage 4约束解码器（ConstrainedSkillDecoder）支持三种模式："
        "(1) ChatHarness在线模式——构建constrained prompt后调用系统LLM生成JSON，支持一次grammar_retry；"
        "(2) chat_fn直连模式——通过异步回调函数调用；"
        "(3) TSE+local离线模式——当LLM不可用时，synthesize_skills_local利用Stage 3的focus_indices"
        "和utterance片段组合生成结构化技能定义，不依赖外部LLM。"
    )
    add_body(doc,
        "多任务训练目标包含三项损失：\n"
        "    L_total = 1.0·L_decoder + 0.1·L_category + 0.1·L_tools\n"
        "其中L_decoder为自编码重建损失，L_category为技能类别分类交叉熵，"
        "L_tools为工具多标签二元交叉熵（约50个候选工具）。"
        "TSEConfig完整参数：embed_dim=256, tcn_hidden_dim=256, tcn_num_layers=3, "
        "tcn_kernel_size=3, tcn_dropout=0.0, num_queries=5, num_heads=4, "
        "top_k_utterances=8, max_skills=8, min_skills=1, decoder_temperature=0.2。"
    )


def enrich_memory_algorithm_detail(doc):
    """6.1: Add precise memory algorithm details."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 6.1.1 记忆核心精确算法】")
    add_body(doc,
        "AgentMemoryCore（agent_memory_core.py）以JSON文件存储四层记忆："
        "storage/agent_memory/{team_id}/{agent_id}/{log,perception,intentions,affect,legacy,meta}.json。"
    )
    add_body(doc, "EpisodicLog（情景日志）：")
    add_body(doc,
        "  每条事件记录 {id, t, subject, action, detail, place, importance(1-10), tags, lastAccessAt}。"
        "recall(query, k) 函数计算复合评分：score = recency + importance + relevance。"
        "recency采用指数衰减：decay = 0.995^{hours}。relevance通过bigram Jaccard相似度计算："
        "对query和事件文本分别提取character bigram集合，计算交集占比。"
    )
    add_body(doc, "PerceptionStream（感知流）：")
    add_body(doc,
        "  以FIFO队列维护最近500条感知记录。compress() 操作将缓冲区压缩为一条EpisodicLog事件"
        "（action='感知压缩'），包含各modality计数和fear均值，然后清空缓冲区。"
        "summarize() 返回统计摘要：count、时间窗口[tStart, tEnd]、byModality分布、fearMean。"
    )
    add_body(doc, "IntentionQueue（意图队列）：")
    add_body(doc,
        "  每个意图包含 {id, tCreated, creator, instruction, trigger, dueAt, countdown,"
        "status(pending/confirmed/dropped), timeoutPolicy(drop/escalate/keep), provenance, handover}。"
        "pending() 按dueAt排序输出，计算daysLeft/dueLabel。"
    )
    add_body(doc, "AffectResidue（情绪残留）：")
    add_body(doc,
        "  状态 {valence ∈ [-1,1], arousal ∈ [0,1], labels: {tag→intensity}}。"
        "指数衰减半衰期 AFFECT_ETA_MS = 72小时。feel(label, intensity, valence, arousal) 操作："
        "新情绪取max(旧值×1.2, 新值)以避免快速遗忘掩盖强情绪；valence和arousal各取新旧50%加权。"
        "tone_hint() 根据主导标签(intensity≥0.6→'浓浓的'，≥0.3→'一丝'，<0.3→'一点未散的')"
        "和arousal/valence级别生成自然语言语气描述。"
    )


def enrich_memory_transfer_detail(doc):
    """6.2: Add precise memory transfer protocol details."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 6.2.1 记忆遗传精确协议】")
    add_body(doc,
        "AgentMemoryTransfer.execute() 实现四阶段遗传协议。"
        "核心原则：传递=复制（非移动），原件默认可凭吊（keep_memorial=True）。"
    )
    add_body(doc, "Seal阶段：")
    add_body(doc,
        "  src.seal(now) 将四层记忆快照写入legacy.json。快照schema='ag.legacy/v1'，包含："
        "sealedAt时间戳、完整log列表、perceptionSummary摘要、intentions列表、affectSnapshot。"
        "源Agent元数据meta.json更新：state→'archived'，sealed=True，记录transferred_to和transfer_id。"
    )
    add_body(doc, "Will阶段：")
    add_body(doc,
        "  draft_will() 返回遗嘱草稿：{testator, team_id, beneficiary, migrate_preferences,"
        "handover_intentions, keep_memorial}。handover_intentions支持三种策略："
        "ask_new_owner（默认，意图标记为pending等待新主人确认）、auto（自动接受）、drop（丢弃全部意图）。"
    )
    add_body(doc, "Export→Import阶段：")
    add_body(doc,
        "  src.export_all() 序列化为schema='ag.memory/v1'的JSON对象，包含{layers: {log, perception, intentions, affect}, meta}。"
        "dst.import_all(data) 执行模式校验后逐层merge（追加而非整盘覆盖，避免抹掉受益方已有记忆）：\n"
        "  - log层：每条事件id重生成，追加tag='传递继承'，place追加'[from:{src_agent_id}]'\n"
        "  - perception层：追加到缓冲区，超过500条截断\n"
        "  - intentions层：根据handover策略复制或丢弃，ask_new_owner时标记'[待新主人确认]'\n"
        "  - affect层：合并标签——对每个标签取max(旧值, 新值)强度；valence和arousal各取新旧50%加权"
    )
    add_body(doc,
        "遗传完成后，受益方自动记录一条EpisodicLog事件（action='记忆继承'，importance=9），"
        "并更新受益方meta：state→'active'，inherited_from记录来源Agent。"
        "转移记录写入storage/agent_memory/_transfers/tr_{transfer_id}.json，包含完整审计轨迹。"
    )


def enrich_skill_extraction_algorithms(doc):
    """Add the three extraction algorithms detail (from skill_extractor.py prompt)."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 5.2.1 技能萃取三大算法】")
    add_body(doc,
        "SkillExtractorEngine._llm_prefill() 实际调用的LLM系统提示中定义了三种核心萃取算法，"
        "每轮讨论的同一文本被三种算法独立处理，各自产出不同的技能：\n"
        "算法1 — 去语境化（De-contextualization）：剥离具体业务细节（产品名、服务名、错误码），"
        "抽取抽象的'动词'和'逻辑关系'。例如'修改RDS参数组'→抽象为'版本化配置管理'；"
        "'ALB导致Lambda超时'→抽象为'依赖树分析'。目标：使技能跨平台、跨语境可迁移。\n"
        "算法2 — 反面模式萃取（Anti-Pattern Extraction）：扫描痛点、故障、FAQ异议和事故，"
        "将每个失败反转为防御规则。例如'忘记原始参数值'→'无备份不变配'；"
        "'周五下午变配导致加班'→'禁绝高风险时段变配'。目标：技能'带刺'——直指人为错误模式。\n"
        "算法3 — 关键路径与最小动作集（Critical Path & Minimum Action Set）："
        "从长篇叙事中抽取3-5个决定性动作，识别成败关键路径。将'SOP面'精确化为可操作的'决策点'。"
        "每个决策点成为独立技能，附带触发条件和验证标准。"
    )
    add_body(doc,
        "萃取管线同时支持知识簇预处理（preprocess_knowledge_clusters）："
        "长文本通过_chunk_by_structure()按Markdown标题或空行分段，"
        "_cluster_chunks()使用TF-IDF余弦相似度的凝聚聚类算法将相似chunk聚合为"
        "最多5个知识簇（max_clusters=5），每个知识簇独立进入萃取管线。"
    )


def enrich_skill_evolution_detail(doc):
    """7: Add LLM-as-Judge evolution details."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 7.5 技能演化LLM-as-Judge精确管线】")
    add_body(doc,
        "技能演化引擎（evolution/fitness.py）采用Qwen3作为Judge，对每个技能变体进行"
        "simulate→judge的双阶段评估："
    )
    add_body(doc, "阶段1 — 模拟执行（simulate_execution）：")
    add_body(doc,
        "  使用系统提示'SIMULATE_SYSTEM_PROMPT'指示Agent严格按照技能指令执行任务。"
        "模拟器调用ChatHarness.chat()，agent_id='evolution_simulator'。"
        "模拟失败时返回占位文本'[模拟执行失败]'，对应的fitness评分全部归零。"
    )
    add_body(doc, "阶段2 — LLM-as-Judge评分（judge_execution）：")
    add_body(doc,
        "  Judge的系统提示定义了三个评分维度（0.0-1.0）：\n"
        "  (1) instruction_following（指令遵循度）：是否忠实遵循技能指令的流程和要求\n"
        "  (2) output_quality（输出质量）：输出是否正确、完整、对用户有帮助\n"
        "  (3) conciseness（简洁度）：是否简洁高效，没有冗余废话\n"
        "  Judge输出JSON评分和reasoning字段。加权复合分：\n"
        "    composite = 0.4 × instruction_following + 0.4 × output_quality + 0.2 × conciseness\n"
        "  SkillFitnessReport聚合所有测试用例的结果，输出mean_composite、mean_following、mean_quality、"
        "mean_conciseness以及failures列表（composite<0.5的用例）。"
    )
    add_body(doc, "变体长度惩罚（apply_length_penalty）：")
    add_body(doc,
        "  对技能文本长度膨胀的变体施加线性惩罚。定义original_len为祖先技能文本长度，"
        "evolved_len为变体文本长度。当ratio=evolved_len/original_len ≤ 1.0（变体更短）时无惩罚；"
        "当ratio > max_ratio(默认1.5，即超过150%)时直接归零；"
        "在1.0到1.5之间的线性惩罚：penalty = (ratio-1.0)/(max_ratio-1.0) × 0.2（最大扣20%）。"
    )


def enrich_ecological_experiments(doc):
    """9: Add ecological simulation data from experiments_raw.tsv."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 9.7 多智能体生态仿真数据】")
    add_body(doc,
        "AgentsGroup2026的生态仿真实验（experiments_raw.tsv，36组数据）系统性地探索了"
        "群落组装模式（solo/division/confrontation/mixed）与生境压力（pressure/scarce/harsh/abundant）"
        "对Agent种群分布的影响。每个Agent在仿真中按比例分配三种行为："
        "skill%（技能使用）、collab%（协作交互）、residual%（残余行为）。"
    )
    add_headers = ["运行组", "模式", "压力", "Agent数", "平均skill%", "平均collab%", "关键技能(前3)"]
    add_rows = [
        ["aws-solo-division", "solo", "base", "5", "0.0%", "0.9%", "monitor_alarms_setup, aws_cli_script_authoring, aws_es_scaling"],
        ["build-solo-division", "solo", "base", "5", "10.7%", "15.2%", "architecture_design, task_decomposition, test_execution"],
        ["aws+build-division", "division", "base", "5", "6.8%", "15.0%", "aws_es_scaling, test_design, task_decomposition"],
        ["aws+build-confrontation", "confrontation", "base", "5", "8.8%", "10.6%", "aws_es_scaling, interface_definition, monitor_alarms"],
        ["aws+build-mixed", "mixed", "base", "5", "5.0%", "8.2%", "(5 dominant跨团队技能)"],
        ["habitat-pressure", "division", "pressure", "5", "7.9%", "13.2%", "monitor_alarms, architecture_design, cost_ri"],
        ["habitat-scarce", "division", "scarce", "5", "3.7%", "13.6%", "aws_es_scaling, cost_ri_advisor"],
        ["habitat-harsh", "division", "harsh", "5", "4.8%", "2.5%", "architecture_design, interface_definition"],
        ["habitat-abundant", "division", "abundant", "5", "2.8%", "4.9%", "aws_es_scaling, architecture_design"],
    ]
    add_table(doc, add_headers, add_rows)
    add_body(doc,
        "关键发现：(1) solo模式下build团队（技能丰富度更高）的skill%显著高于aws团队（10.7% vs 0%）；"
        "(2) confrontation模式（竞争主导）比division模式（分工主导）产生更高的skill%（8.8% vs 6.8%），"
        "支持'竞争压力促进技能差异化'的假设；(3) 生境压力（habitat pressure）在pressure水平（abundance=0.55）"
        "时skill%最高（7.9%），而在harsh水平（abundance=0.35）和abundant水平（abundance=1.4）时均下降，"
        "呈倒U型（hump-shaped）关系——中等压力最有利于技能涌现；"
        "(4) mixed模式下出现5个跨团队dominant技能，验证了技能在团队间的横向传播。"
    )


def enrich_skill_classifier_detail(doc):
    """Add skill classifier three-pool detail."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 7.6 技能三池分类器精确判定逻辑】")
    add_body(doc,
        "Skill Classifier（skill_classifier.py）将技能归入三种池，分类是定期重算的函数而非人工标签。"
        "判定阈值集中定义为："
    )
    add_headers = ["参数", "值", "含义"]
    add_rows = [
        ["EXCLUSIVE_TEAM_SHARE", "0.8", "特有技能：单团队使用占比≥80%"],
        ["EXCLUSIVE_MIN_EFFECTIVENESS", "0.6", "特有技能：最低效果评分"],
        ["GENERAL_MIN_TEAMS", "2", "通用技能：最少跨团队采用数"],
        ["GENERAL_MIN_CATEGORIES", "2", "通用技能：最少跨场景类目数"],
        ["RESERVE_MAX_EFFECTIVENESS", "0.4", "储备技能：效果低于此值强制储备"],
        ["STALE_DAYS", "90", "储备技能：未使用超过此天数强制储备"],
        ["GRADUATE_STREAK", "2", "毕业需连续达标周期数"],
        ["DEMOTE_GRACE", "1", "降级宽限周期数"],
    ]
    add_table(doc, add_headers, add_rows)
    add_body(doc,
        "分类流水线classify()按优先级依次判定："
        "(1) lifecycle='degraded' → 强制储备；(2) effectiveness<0.4且有使用记录 → 储备；"
        "(3) 超过90天未使用 → 储备；(4) total_uses=0 → 储备（新技能默认进入储备池）；"
        "(5) team_count≥2或categories_passed≥2且gate_ok → 通用；"
        "(6) 单团队使用占比≥80%且effectiveness≥0.6且meets_rubric → 特有；"
        "(7) 其余 → 储备。"
    )
    add_body(doc,
        "classify_with_history() 实现带防抖的周期重算："
        "毕业（rank上升）需连续GRADUATE_STREAK=2个周期即时分类达标；"
        "降级（rank下降）有DEMOTE_GRACE=1个周期的宽限。"
        "爆发毕业/降级事件时触发event记录，供监控和历史追溯。"
        "ClassificationStore以storage/skill_classification/{team_id}.json持久化，"
        "seed_reserve_from_extraction()在技能萃取完成时自动将新技能写入储备池。"
    )


def enrich_plaza_k8s_detail(doc):
    """Deployment architecture."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 附录A 部署架构与API规范】")
    add_body(doc,
        "AgentsGroup2026采用Kubernetes原生部署架构。k8s/目录定义了完整的基础设施：\n"
        "  - agentsgroup-namespace.yaml：资源隔离命名空间\n"
        "  - agentsgroup-configmap.yaml：集中配置管理（模型池model_pool.json、部署白名单等）\n"
        "  - agentsgroup-secret.yaml：敏感凭证（API keys、master_key、admin_password）\n"
        "  - agentsgroup-deployment.yaml：单体Deployment（主服务+Plaza引擎+TSE管道+技能注册表）\n"
        "  - agentsgroup-teams-deployment.yaml：团队级Deployment（按需横向扩展）\n"
        "  - agentsgroup-service.yaml：ClusterIP + LoadBalancer双模式服务暴露\n"
        "  - kind-cluster.yaml：本地开发Kind集群定义\n"
        "Dockerfile基于Python 3.11，使用多阶段构建（builder→runtime），"
        "暴露端口8000（HTTP API）和8001（WebSocket SSE推送）。"
    )
    add_body(doc,
        "核心API路由（FastAPI）：\n"
        "  - /api/v1/plaza/* — Plaza讨论创建/参与/共识（plaza_routes.py）\n"
        "  - /api/v1/agent-memory/{team_id}/{agent_id}/* — 记忆CRUD/导出/导入/封存（agent_memory_routes.py）\n"
        "  - /api/v1/skills/extract/* — 技能萃取队列/审核/批准（extraction_routes.py）\n"
        "  - /api/v1/skills/evolution/* — 演化任务/AB测试/变异提交（evolution_bridge）\n"
        "  - /api/v1/skills/classify/* — 三池分类/周期重算（skill_classifier_routes.py）\n"
        "  - /api/v1/cost/* — Token成本核算/门控/额度管理（cost_routes.py, token_governance_routes.py）\n"
        "SSE实时推送覆盖：萃取队列状态变更（item_created/status_changed/approved），"
        "Plaza讨论流（message_sent/round_changed），记忆遗传事件（transfer_completed）。"
    )


def enrich_token_governance(doc):
    """Token governance system."""
    add_body(doc, "")
    add_body(doc, "【代码实现精化 — 附录B Token治理与成本控制系统】")
    add_body(doc,
        "AgentsGroup2026内置完整的Token治理子系统（token_governance/），包含以下控制杠杆"
        "（lever_catalog.py定义）："
    )
    add_headers = ["杠杆", "说明", "配置"]
    add_rows = [
        ["behavior_inject", "行为注入：在System Prompt中注入成本意识指令", "system prompt suffix"],
        ["codegraph_bridge", "代码图谱桥接：复用现有代码上下文减少重复查询", "vector DB index"],
        ["cost_tier", "成本分级：根据任务重要性分配模型层级", "3 tiers: fast/balanced/powerful"],
        ["progressive_history", "渐进历史：按token预算截断对话历史", "max_tokens, sliding window"],
        ["prompt_simplify", "提示词精简：LLM自动压缩冗余prompt", "compression ratio"],
        ["rtk_tool_compress", "工具输出压缩：截断大面积工具返回", "max_output_chars, truncation strategy"],
    ]
    add_table(doc, add_headers, add_rows)
    add_body(doc,
        "savings_store.py 记录每次优化节省的token数和金额（按模型单价折算）。"
        "service.py 作为总调度器，在每次LLM调用前依次应用启用的杠杆。"
        "settings.py 通过统一的杠杆参数配置（lever_params.py）控制每个杠杆的启用/禁用和参数。"
    )


# ── Main Enrichment ───────────────────────────────────────────────────────

def main():
    print(f"Loading: {SRC}")
    doc = DocxDocument(SRC)
    print(f"Paragraphs: {len(doc.paragraphs)}")

    # We'll add enrichment sections at strategic locations within the document.
    # Strategy: Locate section boundaries by keyword, insert enriched blocks.

    # 1. After "5.1 层次化编码与约束解码" → DART architecture detail
    idx_5_1 = find_paragraph_by_keyword(doc, "层次化编码与约束解码")
    if idx_5_1:
        insert_after_paragraph(doc, idx_5_1[0] + 1, [enrich_dart_arch_detail])
        print("  ✓ 5.1 DART architecture detail injected")

    # 2. After "5.2 技能Schema与生命周期扩展" → extraction algorithms
    idx_5_2 = find_paragraph_by_keyword(doc, "技能Schema与生命周期扩展")
    if idx_5_2:
        insert_after_paragraph(doc, idx_5_2[0] + 1, [enrich_skill_extraction_algorithms])
        print("  ✓ 5.2 Extraction algorithms injected")

    # 3. After "6.1 四层并行记忆" → memory algorithm detail
    idx_6_1 = find_paragraph_by_keyword(doc, "四层并行记忆")
    if idx_6_1:
        insert_after_paragraph(doc, idx_6_1[0] + 1, [enrich_memory_algorithm_detail])
        print("  ✓ 6.1 Memory algorithm detail injected")

    # 4. After "6.2 Seal—Will—Export—Import遗传链" → transfer protocol
    idx_6_2 = find_paragraph_by_keyword(doc, "Seal—Will—Export—Import")
    if idx_6_2:
        insert_after_paragraph(doc, idx_6_2[0] + 1, [enrich_memory_transfer_detail])
        print("  ✓ 6.2 Memory transfer protocol injected")

    # 5. After "7.3 变异、选择、组合与退役" → evolution detail
    idx_7_3 = find_paragraph_by_keyword(doc, "变异、选择、组合与退役")
    if idx_7_3:
        insert_after_paragraph(doc, idx_7_3[0] + 1, [enrich_skill_evolution_detail])
        print("  ✓ 7.3 Skill evolution LLM-as-Judge injected")

    # 6. After "7.4 双轨知识遗传" → skill classifier
    idx_7_4 = find_paragraph_by_keyword(doc, "双轨知识遗传")
    if idx_7_4:
        insert_after_paragraph(doc, idx_7_4[0] + 1, [enrich_skill_classifier_detail])
        print("  ✓ 7.4 Skill classifier detail injected")

    # 7. After "9.6 初步闭环收敛" → ecological experiments
    idx_9_6 = find_paragraph_by_keyword(doc, "初步闭环收敛")
    if idx_9_6:
        insert_after_paragraph(doc, idx_9_6[0] + 1, [enrich_ecological_experiments])
        print("  ✓ 9.6 Ecological experiments injected")

    # 8. After "10.3 局限性" → deployment & token governance appendices
    idx_10_3 = find_paragraph_by_keyword(doc, "局限性")
    if idx_10_3:
        # Append at end of section 10
        insert_after_paragraph(doc, idx_10_3[-1] + 2, [enrich_plaza_k8s_detail, enrich_token_governance])
        print("  ✓ Appendices (K8s + Token Governance) injected")

    # Save
    print(f"\nSaving to: {DST}")
    doc.save(DST)
    print("Done! Enriched paper saved.")


if __name__ == "__main__":
    main()
