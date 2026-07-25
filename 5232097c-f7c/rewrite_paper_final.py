# -*- coding: utf-8 -*-
"""Rewrite paper from scratch based on original format + new biomimetic memory architecture + clean diagrams."""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

# ── config ──
ROOT = Path("/Users/panglaohu/OpenWorker/5232097c-f7c")
SRC = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文.docx")
DST = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文_记忆拟生版.docx")
CHARTS = ROOT / "mem_paper_charts"
CHARTS.mkdir(exist_ok=True)

plt.rcParams.update({
    "font.sans-serif": ["Arial Unicode MS", "PingFang SC", "Heiti SC", "SimHei", "DejaVu Sans"],
    "axes.unicode_minus": False, "figure.dpi": 200,
})

# ── clean style helpers ──
BOX = dict(boxstyle="round,pad=0.025,rounding_size=0.03", linewidth=1.5, edgecolor="#333333")

def draw_box(ax, x, y, w, h, title, subtitle, color="#2F6B9A", alpha=0.12):
    patch = FancyBboxPatch((x, y), w, h, facecolor=color, alpha=alpha, **BOX)
    ax.add_patch(patch)
    ax.text(x + w/2, y + h*0.65, title, ha="center", va="center", fontsize=8.8, weight="bold", color=color)
    ax.text(x + w/2, y + h*0.28, subtitle, ha="center", va="center", fontsize=6.8, color="#555555")

def arrow(ax, a, b, color="#4B5563"):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=10, linewidth=1.1, color=color, connectionstyle="arc3,rad=0"))

def curved(ax, a, b, rad=0.18, color="#4B5563"):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=10, linewidth=1.1, color=color, connectionstyle=f"arc3,rad={rad}"))


def chart1_closed_loop():
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    boxes = [
        (.025, .70, .15, .16, "Plaza 协商审议", "ORID 议程 / 仪式信号 / 共识", "#1F5A8B"),
        (.21, .70, .15, .16, "ExecutionPlan", "步骤 / 依赖 / 验收条件", "#1F5A8B"),
        (.40, .70, .15, .16, "TaskHabitatContract", "niche / demanded_skills / budget", "#1A7A52"),
        (.59, .70, .15, .16, "TSE 技能萃取", "TCN 编码 → 交叉注意力 → 解码", "#B05D1E"),
        (.78, .70, .14, .16, "孪生竞标 · Ratchet", "C0+R1–R5 / 质量门 / 棘轮锁定", "#A03535"),
        (.80, .27, .16, .16, "生产派发 / EvidenceRun", "胜出构型 → 证据链 → 回流", "#5B3D8C"),
        (.58, .27, .18, .16, "技能演化 · 路由 · 分类", "LLM-as-Judge / SkillRouter / 三池", "#1A7A52"),
        (.35, .27, .18, .16, "拟生记忆有机体", "感觉·情节·自传语义 + 情绪选择场", "#5B3D8C"),
        (.12, .27, .18, .16, "反馈 · 再审议", "EvidenceRun → 记忆加固 → Plaza", "#1F5A8B"),
    ]
    for args in boxes: draw_box(ax, *args)
    for a, b in [(.175,.78),(.36,.78),(.55,.78),(.74,.78)]:
        arrow(ax, (a,.78), (a+.19,.78))
    arrow(ax, (.92,.78), (.88,.43)); arrow(ax, (.80,.355), (.76,.355))
    arrow(ax, (.58,.355), (.53,.355)); arrow(ax, (.35,.355), (.30,.355))
    arrow(ax, (.12,.355), (.105,.43)); arrow(ax, (.105,.55), (.105,.70))
    curved(ax, (.71,.27), (.44,.16), rad=.15, color="#7B5FB0")
    ax.text(.54, .10, "执行结果回注技能适应度、记忆与下一轮审议", ha="center", fontsize=8.2, color="#5B3D8C")
    ax.set_title("Fig. 1  计划驱动的协商—萃取—孪生—记忆统一闭环", fontsize=12.5, weight="bold", pad=10)
    fig.tight_layout(); p = CHARTS / "fig1_closed_loop.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


def chart2_memory_systems():
    """Biomimetic memory: traces (sensory/episodic/semantic), processes (working/prospective), selection field (affective)."""
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")

    # Title bands
    ax.text(.5, .94, "拟生记忆有机体：三类保存痕迹 · 两个动态过程 · 一个选择场", ha="center", fontsize=12, weight="bold", color="#1A2030")
    ax.text(.5, .88, "程序性记忆由技能库承载（不伪装为文本层）", ha="center", fontsize=8.5, color="#6B7280")

    # Left column: Trace layers (保存的痕迹)
    draw_box(ax, .04, .68, .22, .16, "感觉痕迹", "sensory · FIFO-500\n环境输入 · 工具观察 · 感知压缩", "#1F5A8B", .10)
    draw_box(ax, .04, .47, .22, .16, "情节痕迹", "episodic · soft-cap~400\n时间/地点/来源 · recall三因子", "#1F5A8B", .10)
    draw_box(ax, .04, .26, .22, .16, "自传语义", "semantic · max 200\n从情节巩固的 claim · 可泛化", "#1F5A8B", .10)

    # Center column: Processes (not memory layers)
    draw_box(ax, .30, .60, .20, .13, "工作台", "working · slots 3-7\n当前关注槽位 · 过程非层", "#8A6620", .10)
    draw_box(ax, .30, .38, .20, .13, "前瞻意图", "prospective · PENDING/CONFIRMED\n未来触发行动 · 过程非层", "#8A6620", .10)

    # Right column: Selection field (not content)
    draw_box(ax, .56, .47, .22, .16, "情绪选择场", "affective · τ=72h\nvalence / arousal / labels\n调制巩固·检索·遗忘", "#A03535", .10)

    # Process arrows
    arrow(ax, (.15, .68), (.35, .73)); arrow(ax, (.15, .47), (.35, .48))
    arrow(ax, (.35, .60), (.56, .55)); arrow(ax, (.35, .38), (.56, .47))

    # Consolidation
    curved(ax, (.15, .40), (.30, .25), rad=-.18, color="#1A7A52")
    ax.text(.18, .18, "consolidate_tick\n(importance ≥ 门槛)", fontsize=7, color="#1A7A52", ha="center")

    # Forget
    curved(ax, (.50, .26), (.65, .26), rad=.15, color="#A03535")
    ax.text(.56, .18, "forget_tick\n(低分情节 soft-forget)", fontsize=7, color="#A03535", ha="center")

    # Unique style panel
    draw_box(ax, .80, .35, .18, .18, "Agent 独有记忆方式", "continuity · restraint\nplasticity · permeability\n任务成败 → 缓慢漂移", "#5B3D8C", .10)

    arrow(ax, (.50, .26), (.80, .44))

    ax.set_title("Fig. 4  拟生记忆系统：保存的痕迹、调制过程和 Agent 独有演化方式", fontsize=12, weight="bold", pad=8)
    fig.tight_layout(); p = CHARTS / "fig4_memory_systems.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


def chart3_bidding():
    fig, ax = plt.subplots(figsize=(9.5, 5.2))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")

    draw_box(ax, .04, .66, .17, .17, "基线 C0", "原始 团队/技能/模型", "#1F5A8B", .10)
    for i, (r, t) in enumerate([("R1", "换角色"), ("R2", "换技能"), ("R3", "并行化"), ("R4", "加Review"), ("R5", "模型降档")]):
        draw_box(ax, .26 + i*.14, .67, .12, .14, r, t, "#B05D1E", .10)
    for i in range(5): arrow(ax, (.21, .75), (.26 + i*.14, .75))
    draw_box(ax, .23, .36, .22, .15, "孪生试炼", "success_rate · quality · token", "#1A7A52", .10)
    draw_box(ax, .52, .36, .18, .15, "双门槛筛选", "success ≥ 0.9 且\nquality ≥ 0.9", "#A03535", .10)
    draw_box(ax, .77, .36, .18, .15, "达标排序", "token 升序\n同 token 质量降序", "#B05D1E", .10)
    arrow(ax, (.56, .66), (.38, .51)); arrow(ax, (.45, .43), (.52, .43)); arrow(ax, (.70, .43), (.77, .43))
    draw_box(ax, .30, .07, .24, .15, "Ratchet Ledger", "效率 = quality/max(token,1)\n只进不退 · generation 递增", "#5B3D8C", .10)
    draw_box(ax, .65, .07, .24, .15, "回流与入账", "胜者回写 Plaza\nsimulation 成本分账", "#4B5563", .10)
    arrow(ax, (.86, .36), (.54, .22)); arrow(ax, (.54, .15), (.65, .15))
    ax.set_title("Fig. 5A  候选组合竞标、双门槛筛选与棘轮锁定", fontsize=12, weight="bold", pad=8)
    fig.tight_layout(); p = CHARTS / "fig5a_bidding.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


def chart4_competition(summary):
    names = ["Division", "Confrontation", "Mixed"]
    r = ["aws+build-division", "aws+build-confrontation", "aws+build-mixed"]
    s = [summary[r_]["skill"] * 100 for r_ in r]; c = [summary[r_]["collab"] * 100 for r_ in r]
    x = np.arange(3); w = 0.34
    fig, ax = plt.subplots(figsize=(7, 4.2))
    b1 = ax.bar(x - w/2, s, w, label="Skill %", color="#1F5A8B", edgecolor="white")
    b2 = ax.bar(x + w/2, c, w, label="Collab %", color="#B05D1E", edgecolor="white")
    ax.set_xticks(x, names); ax.set_ylabel("Run mean (%)"); ax.set_ylim(0, max(s + c) + 5)
    ax.grid(axis="y", alpha=.2); ax.legend(frameon=False, ncol=2)
    for bars in (b1, b2):
        for b in bars: ax.text(b.get_x() + b.get_width()/2, b.get_height() + .4, f"{b.get_height():.1f}", ha="center", va="bottom", fontsize=7.5)
    ax.set_title("Fig. 8  群落组装模式对技能与协作行为的影响", fontsize=11, weight="bold")
    fig.tight_layout(); p = CHARTS / "fig8_competition.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


def chart5_hump(summary):
    runs = ["habitat-harsh-aws+build", "habitat-scarce-aws+build", "habitat-pressure-aws+build", "habitat-abundant-aws+build"]
    d = sorted([summary[r_] for r_ in runs], key=lambda z: z["abundance"])
    x = np.array([z["abundance"] for z in d]); y = np.array([z["skill"] * 100 for z in d]); c = np.array([z["collab"] * 100 for z in d])
    fig, ax = plt.subplots(figsize=(7, 4.2))
    ax.plot(x, y, marker="o", lw=2, color="#1F5A8B", label="Skill %")
    ax.plot(x, c, marker="s", lw=1.6, ls="--", color="#B05D1E", label="Collab %")
    ax.set_xlabel("Resource abundance"); ax.set_ylabel("Run mean (%)")
    ax.grid(alpha=.2); ax.legend(frameon=False)
    ax.set_title("Fig. 9  资源丰度与群体行为的非单调关系", fontsize=11, weight="bold")
    for xi, yi, lab in zip(x, y, ["Harsh\n0.35", "Scarce\n0.45", "Pressure\n0.55", "Abundant\n1.40"]):
        ax.annotate(lab, (xi, yi), xytext=(0, 8), textcoords="offset points", ha="center", fontsize=7)
    fig.tight_layout(); p = CHARTS / "fig9_hump.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


def chart6_profiles(rows):
    sel = [r for r in rows if r["run"] in ("aws+build-confrontation", "aws+build-mixed")]
    labels = [r["agent"] for r in sel]; s = [float(r["skill_pct"]) * 100 for r in sel]
    cl = [float(r["collab_pct"]) * 100 for r in sel]; res = [float(r["residual_pct"]) * 100 for r in sel]
    y = np.arange(len(sel)); fig, ax = plt.subplots(figsize=(7, 5))
    ax.barh(y, s, label="Skill", color="#1F5A8B"); ax.barh(y, cl, left=s, label="Collab", color="#B05D1E")
    ax.barh(y, res, left=np.array(s) + np.array(cl), label="Residual", color="#C5CDD4")
    ax.set_yticks(y, labels, fontsize=7); ax.invert_yaxis(); ax.set_xlim(0, 100)
    ax.set_xlabel("Agent behaviour (%)"); ax.legend(frameon=False, ncol=3, loc="lower right")
    ax.grid(axis="x", alpha=.2)
    ax.set_title("Fig. 10  Agent 级行为构成剖面", fontsize=11, weight="bold")
    fig.tight_layout(); p = CHARTS / "fig10_profiles.png"; fig.savefig(p, bbox_inches="tight"); plt.close(fig); return p


# ── data ──
import csv
DATA = Path("/Users/panglaohu/Downloads/AgentsGroup2026/experiments_raw.tsv")
rows_all = list(csv.DictReader(DATA.open(), delimiter="\t"))
by = {}
for r in rows_all: by.setdefault(r["run"], []).append(r)
summary = {}
for run, rr in by.items():
    r0 = rr[0]; summary[run] = {"run": run, "skill": float(r0["run_skill_mean"]), "collab": float(r0["run_collab_mean"]), "abundance": float(r0["abundance"]) if r0["abundance"] else None}

charts = [
    chart1_closed_loop(),
    chart2_memory_systems(),
    chart3_bidding(),
    chart4_competition(summary),
    chart5_hump(summary),
    chart6_profiles(rows_all),
]

# ── docx builders ──
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

def sr(run):
    run.font.name = "Times New Roman"; rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None: rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), "宋体"); run.font.size = Pt(10.5)

def ap(anchor, text="", bold=False, indent=True):
    el = OxmlElement("w:p"); anchor._p.addnext(el); p = Paragraph(el, anchor._parent)
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(.74)
    r = p.add_run(text)
    if text: r.bold = bold; sr(r)
    return p

def ah(anchor, text, size=12):
    p = ap(anchor, "", indent=False)
    p.paragraph_format.space_before = Pt(7)
    r = p.runs[0] if p.runs else p.add_run(""); r.text = text; r.bold = True; r.font.size = Pt(size)
    return p

def ai(anchor, path, w=14.8):
    p = ap(anchor, "", indent=False); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(w)); return p

def at(anchor, hd, rows):
    tbl = OxmlElement("w:tbl")
    pr = OxmlElement("w:tblPr"); st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableGrid"); pr.append(st)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), "5000"); w.set(qn("w:type"), "pct"); pr.append(w); tbl.append(pr)
    grid = OxmlElement("w:tblGrid"); [grid.append((lambda: (gc := OxmlElement("w:gridCol"), gc.set(qn("w:w"), str(9000//len(hd))), gc)[2])()) for _ in hd]; tbl.append(grid)
    def row(vals, head=False):
        tr = OxmlElement("w:tr")
        for v in vals:
            tc = OxmlElement("w:tc"); tcp = OxmlElement("w:tcPr"); tc.append(tcp)
            pp = OxmlElement("w:p"); rr = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
            rfs = OxmlElement("w:rFonts"); rfs.set(qn("w:ascii"), "Times New Roman"); rfs.set(qn("w:eastAsia"), "宋体"); rpr.append(rfs)
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rpr.append(sz)
            if head: rpr.append(OxmlElement("w:b"))
            rr.append(rpr); tt = OxmlElement("w:t"); tt.text = str(v); rr.append(tt); pp.append(rr); tc.append(pp); tr.append(tc)
        return tr
    tbl.append(row(hd, True)); [tbl.append(row(r)) for r in rows]
    anchor._p.addnext(tbl); sp = OxmlElement("w:p"); tbl.addnext(sp); return Paragraph(sp, anchor._parent)

def find(doc, text):
    for p in doc.paragraphs:
        if text in p.text: return p
    raise RuntimeError(f"anchor not found: {text}")

# ── paper rewrite ──
def rewrite():
    doc = DocxDocument(SRC)
    print(f"Original paragraphs: {len(doc.paragraphs)}")

    # ── §3.3 TaskHabitatContract ──
    a = find(doc, "3.2 跨模块身份与来源关系")
    p = ah(a, "3.3 从执行计划到任务生境契约", 12)
    p = ap(p, "统一闭环并不直接把Plaza结论交给生态仿真，而是先编译为TaskHabitatContract。输入ExecutionPlan经依赖拓扑排序形成有序NicheWindow序列；每个窗口包含step_id、demanded_skills、responsible_role、acceptance、base_ticks= max(8,12+4·|K|)、depends_on与inferred_skills标记。总预算B=clamp(Σbⱼ,40,500)，世代预算G=clamp(2+⌈|N|/3⌉,1,10)。契约以SHA-256截断指纹绑定plan_id与revision，使演练结果可回溯到具体计划版本。")
    p = ai(p, charts[0])
    p = ap(p, "图 1A  计划驱动的协商—TSE萃取—数字孪生竞标—拟生记忆统一闭环", indent=False)
    p = ap(p, "Fig. 1A  Plan-driven deliberation—TSE extraction—digital twin bidding—biomimetic memory closed loop.")

    # ── §6 rewrite: biomimetic memory ──
    # Remove old §6 anchors and insert fresh content after §5.3
    a = find(doc, "6.1 四层并行记忆")
    # Go back a bit to insert a rewritten §6
    p = ap(a, "")  # spacer
    p = ah(a, "6 拟生动态记忆有机体", 12)
    p = ap(p, "人类记忆由多个相互作用的系统构成：感觉记忆（sensory memory）极短暂保留输入痕迹供注意选择；工作记忆（working memory）维持当前任务的容量受限空间；情节记忆（episodic memory）编码带时间、地点和自我来源的经历；语义记忆（semantic memory）从经历中巩固出可泛化的概念和规律；前瞻记忆（prospective memory）记录未来触发时要执行的行动；情绪系统通过杏仁核—海马环路调制记忆的编码、巩固、检索（优先记住高唤醒事件）和遗忘（适应性衰减低价值痕迹）。程序性记忆（procedural memory）由基底节和小脑承载，不经过海马——在本文中，它由技能库、熟练度和执行轨迹承载，不复制为文本记忆层。")
    p = ap(p, "本文据此将Agent记忆设计为三个保存痕迹层（sensory感知流、episodic情节日志、semantic自传语义核）、两个调制过程（working工作台、prospective前瞻意图缓冲区）和一个选择场（affective情绪电荷场）。前瞻意图是process而非layer——它记录'以后要做'，不充当记忆存储，自然解决原设计中将'未发送队列'误归为记忆层的问题。情绪不存储事实，而是通过适应度选择压调制情节→语义的巩固概率、低分情节的遗忘强度以及对话检索时的语气和相关性偏置。")
    p = ai(p, charts[1])
    p = ap(p, "图 4A  拟生记忆系统：三类保存痕迹（感觉·情节·自传语义）、两个调制过程（工作台·前瞻意图）和一个选择场（情绪电荷）。程序性记忆由技能库承载。", indent=False)
    p = ap(p, "Fig. 4A  Biomimetic memory systems: three trace layers, two processes, and one affective selection field.")
    p = ah(p, "6.1 情节—语义巩固与可遗忘", 11)
    p = ap(p, "情节日志以soft-cap~400维护。consolidate_tick()将importance≥门槛（默认5，随适应度漂移）的未巩固情节转化为语义核中的claim，记录source_event_ids可回溯来源。同一claim重复出现时strength叠加。语义核容量上限200，超出时淘汰最弱最旧项。forget_tick()对低分情节执行soft-forget（标记forgotten_at而非物理删除），超出soft-cap的部分强制遗忘，再按遗忘激进度的5%额外淘汰；importance≥9且未巩固的极重要事件暂留保护。遗忘不销毁数据——已遗忘情节不参与recall检索但保留在审计中。")
    p = ah(p, "6.2 Agent独有记忆方式与动态漂移", 11)
    p = ap(p, "每个Agent在首次绑定时从不可见的初始化原型（小满式连续或沈弥安式克制）获得初始记忆方式，包含四个维度：连续性（continuity）、克制性（restraint）、可塑性（plasticity）和情绪通透度（affective_permeability）。这些维度不充当静态标签，而是以Agent实际经历——任务成败、物竞存活时长和适应度变化——为输入进行渐变漂移。高适应度Agent的可塑性放大了正面经验的巩固效应，低适应度Agent则提升遗忘激进度和巩固门槛；连续性越高，情节容量上限越宽。每个维度有独立clamp区间（0.15–0.95），保持架构稳定。")
    p = ap(p, "用户可手动微调倾向但不直接改写记忆内容；后续经历仍会持续漂移。演化史以history列表记录每次漂移与调整的fitness_delta和时间戳。该设计将感情解释为记忆物竞天择的结果——情绪电荷是选择压的直接表征，连续性与克制性是分别对'记忆富集'与'记忆节俭'两种生存策略的量化编码。")
    p = ah(p, "6.3 工作台与前瞻意图（过程·非层）", 11)
    p = ap(p, "工作台由Agent独有slots槽位（3–7，随可塑性漂移）构成的容量受限空间。push_working()将当前关注点推入槽位并去重顶到队首；槽满择旧淘汰。任务完成/失败时自动推入当前任务关注点；物竞存活时写入存活统计。前瞻意图维护pending/confirmed/dropped三态和dueAt驱动排序；传递时按ask_new_owner/auto/drop策略交接。二者在API和页面中均标注为过程而非记忆层。")
    p = ah(p, "6.4 传递：叙事的连续与凭吊的克制", 11)
    p = ap(p, "传递遵循'复制非移动'原则。export生成包含所有5个文件+semantic的层化JSON；import逐层merge而非覆盖。传递完成后生成transfer_narrative：高连续性Agent产出'致后来者·一段未断的连续'叙事摘要，将关键情节和语义核组织为第一人称连续句；高克制性Agent产出'凭吊清单·这是回放不是本人'的结构化交接。电荷传递受通透度约束：通透度<0.2时完全剥离电荷；<0.5时仅传递35%强度。原件默认keep_memorial封存。")

    # ── §7.7 SkillRouter ──
    a = find(doc, "7.4 双轨知识遗传")
    p = ah(a, "7.7 SkillRouter两阶段检索、生命周期重排与反馈学习", 11)
    p = ap(p, "SkillRouter将已验证技能注入Agent上下文。Stage 1在全技能池执行BM25、TF-IDF余弦、中文bigram/trigram、描述短语、指令深匹配与同义词扩展，取max(20,3K)个候选；Stage 2按name/description/instructions/category/tools字段联合重排。基础分S=0.45S_retrieval+0.55S_rerank。生命周期乘子修正排序：solidified=1.14、verified=1.12、published=1.08、team_local=1.00、draft=0.90、degraded=0.72。用户rating/revoke反馈更新(agent,category)亲和度[-0.5,0.5]；赋予同时将熟练度先验抬升至≥0.8。")

    # ── §8.4/8.5 Bidding + EvidenceRun ──
    a = find(doc, "8.3 三维控制界面")
    p = ah(a, "8.4 计划驱动竞标与正向棘轮", 11)
    p = ap(p, "对同一任务契约，竞标编排器由基线C0生成最多4个单算子候选：R1换角色、R2换技能绑定、R3并行化无依赖步骤、R4加Review回边、R5模型降档。每个候选在数字孪生产生(success_rate,quality_score,token_consumed,collab_heat)。仅success≥0.9且quality≥0.9的候选进入达标组，达标组按token升序排名；不达标组仅诊断。胜者效率η=quality/max(token,1)，由RatchetLedger以scenario_best:{task_type}共享键实施正向棘轮——仅在η不低于当前值且满足最小增量时推进generation，退步被拒绝。竞标token记为simulation成本不混入生产效能。")
    p = ai(p, charts[2])
    p = ap(p, "图 5A  候选组合竞标、双门槛筛选与棘轮锁定", indent=False)
    p = ah(p, "8.5 EvidenceRun统一证据链与唯一适应度约束", 11)
    p = ap(p, "EvidenceRun是技能验证、任务执行、演化比较和成本门禁共享的追加式证据对象——关联team/agent/skill/task/evolution/cost/plaza/session/request九类ID，保存runtime/command/exit_code/artifact_dir/stdout/stderr/metrics_before/after与detail，采用SHA-256截断evidence_hash按月落盘。生态演练中T_i=survival_ticks唯一原生适应度；skill_ticks/collab_ticks/residual_ticks是对每个存活tick主因的解释性归因，满足三者之和等于T_i，不引入第二适应度。")

    # ── Ecological experiments (after 表5) ──
    a = find(doc, "Table 5 Closed-loop changes")
    p = ap(a, "")
    p = ah(a, "9.7 多智能体生态仿真：群落组装、资源压力与行为涌现", 12)
    p = ap(p, "§9.6短周期闭环仅3次迭代。本节分析experiments_raw.tsv的9组45条Agent观测（每组5个Agent），探索竞合编排和生境压力对群体行为的影响。所有值来自可复核原始字段，结论为探索性证据。")
    p = ah(p, "9.7.1 实验设计与指标", 11)
    p = ap(p, "群落组装对照solo/division/confrontation/mixed；生境对照固定division在harsh(0.35)/scarce(0.45)/pressure(0.55)/abundant(1.40)四种丰度下运行。")
    p = at(p, ["运行组", "编排", "压力", "N", "技能%", "协作%", "代", "最优T"], [
        ["aws-solo-division", "solo", "base", "5", "0.0", "0.9", "1", "45"],
        ["build-solo-division", "solo", "base", "5", "10.7", "15.2", "2", "75"],
        ["aws+build-division", "division", "base", "5", "6.8", "15.0", "1", "69"],
        ["aws+build-confrontation", "confront.", "base", "5", "8.8", "10.6", "1", "77"],
        ["aws+build-mixed", "mixed", "base", "5", "5.0", "8.2", "9", "57"],
        ["habitat-harsh", "division", "0.35", "5", "4.8", "2.5", "1", "60"],
        ["habitat-scarce", "division", "0.45", "5", "3.7", "13.6", "1", "61"],
        ["habitat-pressure", "division", "0.55", "5", "7.9", "13.2", "1", "67"],
        ["habitat-abundant", "division", "1.40", "5", "2.8", "4.9", "2", "86"],
    ])
    p = ap(p, "表 6  九组生态仿真运行汇总（数据源：experiments_raw.tsv）", indent=False)
    p = ah(p, "9.7.2 群落组装模式比较", 11)
    p = ai(p, charts[3])
    p = ap(p, "图 8  群落组装模式对技能与协作行为的影响。", indent=False)
    p = ap(p, "confrontation技能行为均值(8.8%)高于division(6.8%)和mixed(5.0%)；division协作均值(15.0%)最高。竞争性筛选更有利于放大可区分技能行为。mixed组9代后出现5项dominant技能。")
    p = ah(p, "9.7.3 资源丰度与技能涌现", 11)
    p = ai(p, charts[4])
    p = ap(p, "图 9  资源丰度与群体行为的非单调（倒U型）关系。", indent=False)
    p = ap(p, "技能行为从4.8%升至7.9%(pressure)再降至2.8%(abundant)，呈倒U型——中等压力最有利于技能探索与选择。协作在适中压力下保持高位，极端环境下降。")
    p = ah(p, "9.7.4 Agent级行为剖面", 11)
    p = ai(p, charts[5])
    p = ap(p, "图 10  Agent级行为构成（skill+collab+residual=100%）。", indent=False)
    p = ap(p, "Agent间存在显著异质性，支持将技能种群视为角色依赖的生态结构。每组仅5个Agent且无种子重复，不报告显著性检验。")
    p = ah(p, "9.7.5 与统一闭环的关联", 11)
    p = ap(p, "竞合编排影响Plaza中角色交互与选择压；生境参数为外生成本/风险/预算扰动；dominant技能可由分类器、适应度报告和记忆传递记录回溯。生态数据是闭环在群体尺度的补充观测。")

    # ── Appendix ──
    a = find(doc, "11 结束语")
    p = ap(a, "")
    p = ah(a, "附录A  工程部署与API路由", 12)
    p = ap(p, "AgentsGroup2026以Kubernetes部署（Namespace/ConfigMap/Secret/Deployment/Service），Dockerfile多阶段构建。后端FastAPI运行端口8080，前端Vite开发端口5173。核心路由：/api/v1/plaza/*、/api/v1/agent-memory/*（含memory-style、consolidate、forget、drift端点）、/api/v1/skills/extract/*、/api/v1/skills/evolution/*、/api/v1/skills/classify/*、/api/v1/cost/*。Token治理组合行为注入、代码图谱桥接、成本分级、渐进历史、提示词精简和工具输出压缩六类杠杆。")

    doc.save(DST)
    print(f"Saved to: {DST}")
    print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Shapes: {len(doc.inline_shapes)}")

if __name__ == "__main__":
    rewrite()
