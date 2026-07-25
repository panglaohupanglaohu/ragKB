# -*- coding: utf-8 -*-
"""Rebuild paper: add 10 tech modules + new §9.7 with charts & tables."""

import csv
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path("/Users/panglaohu/OpenWorker/5232097c-f7c")
ORIGINAL = Path(
    "/Users/panglaohu/Downloads/"
    "协商审议_DARTNet_记忆遗传_统一闭环论文.docx"
)
DATA = Path("/Users/panglaohu/Downloads/AgentsGroup2026/experiments_raw.tsv")
OUT = ROOT / "协商审议_DARTNet_记忆遗传_统一闭环论文_重写含实验绘图版.docx"
CHART_DIR = ROOT / "paper_charts"
CHART_DIR.mkdir(exist_ok=True)

plt.rcParams["font.sans-serif"] = [
    "Arial Unicode MS", "PingFang SC", "Heiti SC", "SimHei", "DejaVu Sans",
]
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["figure.dpi"] = 180


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def parse_data():
    with DATA.open(encoding="utf-8") as f:
        rows = list(csv.DictReader(f, delimiter="\t"))
    by_run = {}
    for r in rows:
        by_run.setdefault(r["run"], []).append(r)
    summary = {}
    for run, rr in by_run.items():
        r0 = rr[0]
        summary[run] = {
            "run": run,
            "n": len(rr),
            "skill": float(r0["run_skill_mean"]),
            "collab": float(r0["run_collab_mean"]),
            "bestT": int(r0["bestT"]),
            "gens": int(r0["gens"]),
            "regime": r0["regime"] or "base",
            "abundance": (
                float(r0["abundance"]) if r0["abundance"] else None
            ),
        }
    return rows, summary


# ---------------------------------------------------------------------------
# Chart generators
# ---------------------------------------------------------------------------

def chart_competition(su):
    division_runs = ["aws+build-division", "aws+build-confrontation", "aws+build-mixed"]
    skills = [su[r]["skill"] * 100 for r in division_runs]
    collabs = [su[r]["collab"] * 100 for r in division_runs]

    x = np.arange(3)
    w = 0.36
    fig, ax = plt.subplots(figsize=(7, 4))
    b1 = ax.bar(x - w / 2, skills, w, label="Skill %", color="#2F6B9A")
    b2 = ax.bar(x + w / 2, collabs, w, label="Collab %", color="#D98943")
    ax.set_xticks(x, ["Division", "Confrontation", "Mixed"])
    ax.set_ylabel("Run mean (%)")
    ax.set_title("Fig.7 Community assembly modes vs. skill & collaboration behavior")
    ax.set_ylim(0, max(collabs + skills) + 5)
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=2, loc="upper right")
    for bars in (b1, b2):
        for b in bars:
            ax.text(
                b.get_x() + b.get_width() / 2,
                b.get_height() + 0.4,
                f"{b.get_height():.1f}",
                ha="center", va="bottom", fontsize=8,
            )
    fig.tight_layout()
    p = CHART_DIR / "fig7_competition_modes.png"
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def chart_hump(su):
    habitat_runs = [
        "habitat-harsh-aws+build",
        "habitat-scarce-aws+build",
        "habitat-pressure-aws+build",
        "habitat-abundant-aws+build",
    ]
    d = sorted([su[r] for r in habitat_runs], key=lambda x: x["abundance"])
    x = np.array([z["abundance"] for z in d])
    y = np.array([z["skill"] * 100 for z in d])
    c = np.array([z["collab"] * 100 for z in d])

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.plot(x, y, marker="o", lw=2.1, color="#2F6B9A", label="Skill %")
    ax.plot(x, c, marker="s", lw=1.8, ls="--", color="#D98943", label="Collab %")
    labels = ["Harsh (0.35)", "Scarce (0.45)", "Pressure (0.55)", "Abundant (1.40)"]
    for xi, yi, lab in zip(x, y, labels):
        ax.annotate(lab, (xi, yi), xytext=(0, 9), textcoords="offset points",
                    ha="center", fontsize=8)
    ax.set_xlabel("Resource abundance")
    ax.set_ylabel("Run mean (%)")
    ax.set_title("Fig.8 Resource abundance vs. population behavior")
    ax.grid(alpha=0.25)
    ax.legend(frameon=False)
    ax.set_ylim(0, max(max(y), max(c)) + 5)
    fig.tight_layout()
    p = CHART_DIR / "fig8_abundance_hump.png"
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def chart_agent_profiles(rows):
    sel = [r for r in rows if r["run"] in ("aws+build-confrontation", "aws+build-mixed")]
    labels = [r["agent"] for r in sel]
    s = [float(r["skill_pct"]) * 100 for r in sel]
    cl = [float(r["collab_pct"]) * 100 for r in sel]
    res = [float(r["residual_pct"]) * 100 for r in sel]

    y = np.arange(len(sel))
    fig, ax = plt.subplots(figsize=(7.2, 5))
    ax.barh(y, s, label="Skill", color="#2F6B9A")
    ax.barh(y, cl, left=s, label="Collab", color="#D98943")
    ax.barh(y, res, left=np.array(s) + np.array(cl), label="Residual", color="#C7CDD4")
    ax.set_yticks(y, labels, fontsize=7)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Agent behaviour (%)")
    ax.set_title("Fig.9 Agent-level behaviour profiles")
    ax.legend(frameon=False, ncol=3, loc="lower right")
    ax.grid(axis="x", alpha=0.2)
    fig.tight_layout()
    p = CHART_DIR / "fig9_agent_profiles.png"
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _style_run(run):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def _new_para(anchor, text="", bold=False, indent=True):
    """Insert a new paragraph after *anchor* and return it."""
    el = OxmlElement("w:p")
    anchor._p.addnext(el)
    p = Paragraph(el, anchor._parent)
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    if indent:
        pf.first_line_indent = Cm(0.74)
    r = p.add_run(text or "")
    if text:
        r.bold = bold
        _style_run(r)
    return p


def _new_heading(anchor, text, level=2):
    """Insert a heading after *anchor*."""
    p = _new_para(anchor, "", indent=False)
    pf = p.paragraph_format
    pf.space_before = Pt(7)
    r = p.runs[0] if p.runs else p.add_run("")
    r.text = text
    r.bold = True
    r.font.size = Pt(12 if level == 2 else 11)
    return p


def _new_table(anchor, headers, data_rows):
    """Build a w:tbl element, insert after *anchor*, return trailing spacer paragraph."""
    tbl = OxmlElement("w:tbl")

    # tblPr
    tp = OxmlElement("w:tblPr")
    tp2 = OxmlElement("w:tblStyle")
    tp2.set(qn("w:val"), "TableGrid")
    tp.append(tp2)
    tp3 = OxmlElement("w:tblW")
    tp3.set(qn("w:w"), "5000")
    tp3.set(qn("w:type"), "pct")
    tp.append(tp3)
    tbl.append(tp)

    # tblGrid
    tg = OxmlElement("w:tblGrid")
    col_w = 9000 // len(headers)
    for _ in headers:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(col_w))
        tg.append(gc)
    tbl.append(tg)

    # ── helper: make one row ──
    def _row(vals, is_header=False):
        tr = OxmlElement("w:tr")
        for v in vals:
            tc = OxmlElement("w:tc")
            tcp = OxmlElement("w:tcPr")
            tc.append(tcp)
            wp = OxmlElement("w:p")
            wr = OxmlElement("w:r")
            if is_header:
                rpr = OxmlElement("w:rPr")
                b_el = OxmlElement("w:b")
                rpr.append(b_el)
                wr.append(rpr)
            wt = OxmlElement("w:t")
            wt.set(qn("xml:space"), "preserve")
            wt.text = str(v)
            wr.append(wt)
            wp.append(wr)
            tc.append(wp)
            tr.append(tc)
        return tr

    tbl.append(_row(headers, is_header=True))
    for row in data_rows:
        tbl.append(_row(row))

    anchor._p.addnext(tbl)

    # spacer paragraph
    sp_el = OxmlElement("w:p")
    tbl.addnext(sp_el)
    return Paragraph(sp_el, anchor._parent)


def _new_image(anchor, path, width_cm=14.7):
    p = _new_para(anchor, "", indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    return p


def insert_block(doc, keyword, build_fn):
    hits = [p for p in doc.paragraphs if keyword in p.text]
    if hits:
        build_fn(doc, hits[0])
    else:
        print(f"  WARN: anchor not found: {keyword}")


# ---------------------------------------------------------------------------
# 10 technical modules
# ---------------------------------------------------------------------------

def module_dart(doc, a):
    p = _new_heading(a, "5.1.1 TSE管线工程参数与可复现架构", 3)
    p = _new_para(p,
        "生产代码 src/backend/agents/tse/ 实现了完整的DART-Net管线。"
        "Stage 1话语编码使用BLAKE2b确定性哈希（hash_seed=20260716）"
        "对每条Plaza消息的content字段执行character 1/2/3‑gram特征哈希，"
        "映射为256维浮点嵌入并L2归一化。同时将角色（12类）、仪式信号（10类）、"
        "niche角色（6类）和轮次（16档）以辅助嵌入叠加"
        "（权重分别为0.1/0.1/0.1/0.05）。配置：embed_dim=256，"
        "max_utterances=64，max_chars_per_utterance=800。"
    )
    p = _new_para(p,
        "Stage 2采用三层深度可分离膨胀卷积（DilatedConvBlock），每层执行"
        "depthwise conv(C×K)→pointwise 1×1→LayerNorm→ReLU→残差连接。"
        "dilations=[1,2,4]，kernel_size=3；理论感受野RF=1+2×(1+2+4)=15"
        "（单侧）。权重初始化为He式缩放（depthwise: √(2/k)×0.5，"
        "pointwise: √(2/C)×0.5）。"
    )
    p = _new_para(p,
        "Stage 3以5组可学查询探针对TCN输出执行交叉注意力，"
        "输出可回溯的focus_indices。冷启动融合FIELD_KEYWORD_SEEDS先验"
        "（权重0.3）。Stage 4的ConstrainedSkillDecoder支持"
        "ChatHarness在线、chat_fn异步和TSE+local离线三后端，"
        "grammar_retry=1，输出经JSON schema验证。完整TSEConfig参数："
        "tcn_hidden_dim=256，num_heads=4，top_k_utterances=8，"
        "decoder_temperature=0.2。"
    )


def module_extract(doc, a):
    p = _new_heading(a, "5.2.1 三算法技能萃取与知识簇预处理", 3)
    p = _new_para(p,
        "SkillExtractorEngine以三种互补算法独立处理同一Plaza讨论。"
        "算法1（去语境化）剥离AWS服务名、错误码等具体细节，"
        "抽取抽象动词—关系对。算法2（反面模式）将事故日志和异议发言"
        "反转为防御规则。算法3（关键路径）从完整SOP提取3—5个成败判断点。"
    )
    p = _new_para(p,
        "长文档经_chunk_by_structure()分块后以TF-IDF余弦相似度"
        "凝聚聚类（max_clusters=5），每簇独立萃取、经slug去重和"
        "审核队列合并。审核队列以pending→llm_prefilling→ready_for_review"
        "→approved/rejected状态机持久化至storage/skill_extract_queue/。"
    )


def module_memory(doc, a):
    p = _new_heading(a, "6.1.1 四层记忆的精确算法", 3)
    p = _new_para(p,
        "AgentMemoryCore（agent_memory_core.py）以JSON存储。"
        "EpisodicLog的recall()使用复合评分"
        "score=recency+importance+relevance，"
        "recency=0.995^hours，relevance由character bigram Jaccard"
        "重叠度计算。PerceptionStream为500条FIFO；compress()输出"
        "时间窗和各modality计数后清空。IntentionQueue按dueAt排序，"
        "支持drop/escalate/keep超时策略。AffectResidue以τ=72h"
        "指数衰减，重复感受按max(旧×1.2,新)更新。"
    )


def module_transfer(doc, a):
    p = _new_heading(a, "6.2.1 Seal—Will—Export—Import的可审计传递", 3)
    p = _new_para(p,
        "AgentMemoryTransfer遵循\"复制非移动、保留非销毁\"原则。"
        "seal写Schema=ag.legacy/v1的遗体快照；will指定受益方和"
        "handover_intentions（ask_new_owner/auto/drop）；"
        "export生成Schema=ag.memory/v1的层化JSON。"
        "import逐层merge：日志追加\"传递继承\"tag和[from:source]标记；"
        "感知追加并截断500条；意图按策略交接；情绪标签max合并，"
        "valence/arousal各50%加权。系统记录transfer_id审计追踪和"
        "importance=9的\"记忆继承\"事件。"
    )


def module_evolution(doc, a):
    p = _new_heading(a, "7.5 LLM-as-Judge演化评估管线", 3)
    p = _new_para(p,
        "演化引擎对每个变体执行simulate→judge双阶段。"
        "Judge评分三维（instruction_following, output_quality, "
        "conciseness，各0—1），复合F=0.4×following+0.4×quality"
        "+0.2×conciseness。SkillFitnessReport聚合均值和"
        "composite<0.5的失败集。变体长度膨胀惩罚：ratio≤1无惩罚；"
        "1<ratio≤1.5线性扣减至多0.2；ratio>1.5归零。"
    )


def module_classifier(doc, a):
    p = _new_heading(a, "7.6 三池分类器与防抖毕业机制", 3)
    p = _new_para(p,
        "Skill Classifier判定技能归入exclusive/general/reserve。"
        "储备条件：lifecycle=degraded、effectiveness<0.4已有使用、"
        ">90天未用或total_uses=0。通用需≥2团队或≥2类场景且gate_ok。"
        "特有需单团队占比≥0.8、effectiveness≥0.6且meets_rubric。"
    )
    p = _new_table(p,
        ["参数", "值", "含义"],
        [
            ["EXCLUSIVE_TEAM_SHARE", "0.8", "特有单团队使用占比"],
            ["EXCLUSIVE_MIN_EFFECTIVENESS", "0.6", "特有最低效果"],
            ["GENERAL_MIN_TEAMS", "2", "通用跨团队采用"],
            ["GENERAL_MIN_CATEGORIES", "2", "通用跨场景类目"],
            ["RESERVE_MAX_EFFECTIVENESS", "0.4", "低效强制储备"],
            ["STALE_DAYS", "90", "未使用回收"],
            ["GRADUATE_STREAK", "2", "连续达标毕业"],
            ["DEMOTE_GRACE", "1", "降级宽限"],
        ],
    )
    p = _new_para(p,
        "classify_with_history()以防抖机制避免边界振荡：毕业需连续2周期；"
        "降级需1周期宽限。新技能通过seed_reserve_from_extraction()"
        "先写入储备池，由使用、验证和采纳证据驱动毕业。"
    )


def module_experiments(doc, a, charts, summary):
    p = _new_heading(a, "9.7 多智能体生态仿真：群落组装、资源压力与行为涌现", 2)
    p = _new_para(p,
        "§9.6的短周期闭环仅3次迭代。本节在群体尺度上补充定量观测，"
        "分析experiments_raw.tsv的9个运行组、45条Agent级观测"
        "（每组5个Agent）。每条记录报告Agent技能使用、协作交互"
        "与残余行为比例，并附带竞合编排、代际数及生境参数。"
        "所有值取自原文件可复核字段，结论应理解为探索性证据。"
    )

    p = _new_heading(p, "9.7.1 实验设计与主要指标", 3)
    p = _new_para(p,
        "群落组装对照solo/division/confrontation/mixed四种编排；"
        "生境对照固定division在harsh(0.35)/scarce(0.45)/"
        "pressure(0.55)/abundant(1.40)四种资源丰度下运行。"
    )
    p = _new_table(p,
        ["运行组", "编排", "压力", "N", "技能%", "协作%", "代", "最优T"],
        [
            ["aws-solo-division", "solo", "base", "5", "0.0", "0.9", "1", "45"],
            ["build-solo-division", "solo", "base", "5", "10.7", "15.2", "2", "75"],
            ["aws+build-division", "division", "base", "5", "6.8", "15.0", "1", "69"],
            ["aws+build-confrontation", "confront.", "base", "5", "8.8", "10.6", "1", "77"],
            ["aws+build-mixed", "mixed", "base", "5", "5.0", "8.2", "9", "57"],
            ["habitat-harsh", "division", "0.35", "5", "4.8", "2.5", "1", "60"],
            ["habitat-scarce", "division", "0.45", "5", "3.7", "13.6", "1", "61"],
            ["habitat-pressure", "division", "0.55", "5", "7.9", "13.2", "1", "67"],
            ["habitat-abundant", "division", "1.40", "5", "2.8", "4.9", "2", "86"],
        ],
    )

    p = _new_heading(p, "9.7.2 群落组装模式比较", 3)
    p = _new_image(p, charts[0], 14.7)
    p = _new_para(p,
        "图7 不同群落组装模式下的技能与协作行为"
        "（数据源：experiments_raw.tsv）。"
    )
    p = _new_para(p,
        "confrontation的技能行为均值（8.8%）高于division（6.8%）和"
        "mixed（5.0%），而division协作均值（15.0%）高于confrontation"
        "（10.6%）。竞争性筛选更有利于放大可区分技能行为，职责分工"
        "更利于维持协作密度。mixed经历9代后出现5项dominant技能。"
        "因仅一次运行，不解释为显著结论。"
    )

    p = _new_heading(p, "9.7.3 资源丰度与技能涌现", 3)
    p = _new_image(p, charts[1], 14.7)
    p = _new_para(p,
        "图8 资源丰度、生境压力与群体行为的关系。"
    )
    p = _new_para(p,
        "技能行为从4.8%(harsh)升至7.9%(pressure)再降至2.8%(abundant)，"
        "呈现倒U型，与\"中等压力最有利于技能探索与选择\"的假说一致。"
        "协作在稀缺(13.6%)与适中压力(13.2%)下较高，极端环境下降。"
    )

    p = _new_heading(p, "9.7.4 Agent级行为剖面与解释边界", 3)
    p = _new_image(p, charts[2], 14.7)
    p = _new_para(p,
        "图9 对抗与混合模式下的Agent级行为构成。"
    )
    p = _new_para(p,
        "confrontation中aws_lead和aws_mon呈现较高技能/协作份额，"
        "部分Build Agent以residual为主；mixed组技能优势更分散。"
        "该模式支持将技能种群视为角色依赖的生态结构。"
        "每组仅5个Agent且无独立种子重复，不报告显著性检验。"
        "后续应实施多种子重复和跨域复验。"
    )

    p = _new_heading(p, "9.7.5 与统一闭环的关联", 3)
    p = _new_para(p,
        "竞合编排影响Plaza中角色交互和选择压力，进而改变技能种群构成；"
        "生境参数相当于对外生成本、风险和预算的扰动；dominant技能"
        "可由分类器、适应度报告和记忆传递记录回溯。"
        "生态数据是闭环在群体尺度的补充观测而非独立结果。"
    )


def module_appendix(doc, a):
    p = _new_heading(a, "附录A 工程部署架构、API路由与Token治理", 2)
    p = _new_para(p,
        "AgentsGroup2026以Kubernetes部署"
        "（k8s/目录：Namespace/ConfigMap/Secret/Deployment/Service），"
        "Dockerfile多阶段构建，暴露端口8000(HTTP)和8001(SSE)。"
        "核心FastAPI路由：/api/v1/plaza/*、/api/v1/agent-memory/*、"
        "/api/v1/skills/extract/*、/api/v1/skills/evolution/*、"
        "/api/v1/skills/classify/*、/api/v1/cost/*。"
    )
    p = _new_para(p,
        "Token治理子系统在每次LLM调用前组合行为注入、代码图谱桥接、"
        "成本分级、渐进历史、提示词精简和工具输出压缩六类杠杆，"
        "SavingsStore持久化每次优化节省的token与金额。"
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    rows, summary = parse_data()
    charts = [
        chart_competition(summary),
        chart_hump(summary),
        chart_agent_profiles(rows),
    ]

    doc = DocxDocument(ORIGINAL)
    print(f"Original paragraphs: {len(doc.paragraphs)}")

    insert_block(doc, "5.1 层次化编码与约束解码", module_dart)
    insert_block(doc, "5.2 技能Schema与生命周期扩展", module_extract)
    insert_block(doc, "6.1 四层并行记忆", module_memory)
    insert_block(doc, "6.2 Seal\u2014Will\u2014Export\u2014Import遗传链", module_transfer)
    insert_block(doc, "7.3 变异、选择、组合与退役", module_evolution)
    insert_block(doc, "7.4 双轨知识遗传", module_classifier)
    insert_block(
        doc, "9.6 初步闭环收敛",
        lambda d, a: module_experiments(d, a, charts, summary),
    )
    insert_block(doc, "11 结束语", module_appendix)

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Final paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
