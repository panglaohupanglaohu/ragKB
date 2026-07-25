# -*- coding: utf-8 -*-
"""Rebuild the paper from original format, add ten technical modules and new §9.7 charts."""

import csv, os
from pathlib import Path
from statistics import mean

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path('/Users/panglaohu/OpenWorker/5232097c-f7c')
ORIGINAL = Path('/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文.docx')
DATA = Path('/Users/panglaohu/Downloads/AgentsGroup2026/experiments_raw.tsv')
OUT = ROOT / '协商审议_DARTNet_记忆遗传_统一闭环论文_重写含实验绘图版.docx'
CHART_DIR = ROOT / 'paper_charts'
CHART_DIR.mkdir(exist_ok=True)

plt.rcParams['font.sans-serif'] = ['Arial Unicode MS','PingFang SC','Heiti SC','SimHei','DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['figure.dpi'] = 180


def parse_data():
    with DATA.open(encoding='utf-8') as f:
        rows = list(csv.DictReader(f, delimiter='\t'))
    by_run = {}
    for r in rows:
        by_run.setdefault(r['run'], []).append(r)
    summary = {}
    for run, rr in by_run.items():
        r0 = rr[0]
        summary[run] = {
            'run': run, 'n': len(rr),
            'skill': float(r0['run_skill_mean']),
            'collab': float(r0['run_collab_mean']),
            'bestT': int(r0['bestT']), 'gens': int(r0['gens']),
            'regime': r0['regime'] or 'base', 'tournament': r0['tournament'],
            'abundance': float(r0['abundance']) if r0['abundance'] else None,
            'predator_pressure': float(r0['predator_pressure']) if r0['predator_pressure'] else None,
            'drift_prob': float(r0['drift_prob']) if r0['drift_prob'] else None,
            'niche_capacity': float(r0['niche_capacity']) if r0['niche_capacity'] else None,
            'dominant': r0['dominant'], 'n_dominant': int(r0['n_dominant']) if r0['n_dominant'] else 0,
        }
    return rows, summary


def chart_competition(summary):
    names = ['division','confrontation','mixed']
    runs = ['aws+build-division','aws+build-confrontation','aws+build-mixed']
    skills = [summary[r]['skill'] * 100 for r in runs]
    collabs = [summary[r]['collab'] * 100 for r in runs]
    x = np.arange(len(names)); w = 0.36
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    b1 = ax.bar(x - w / 2, skills, w, label='Skill %', color='#2F6B9A')
    b2 = ax.bar(x + w / 2, collabs, w, label='Collab %', color='#D98943')
    ax.set_xticks(x, ['Division','Confrontation','Mixed'])
    ax.set_ylabel('Run mean (%)')
    ax.set_title('Fig.7 Community assembly modes vs. skill & collaboration')
    ax.set_ylim(0, max(collabs + skills) + 5)
    ax.grid(axis='y', alpha=.25); ax.legend(frameon=False, ncol=2, loc='upper right')
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x() + b.get_width() / 2, b.get_height() + 0.4, f'{b.get_height():.1f}', ha='center', va='bottom', fontsize=8)
    fig.tight_layout(); p = CHART_DIR / 'fig7_competition_modes.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig); return p


def chart_hump(summary):
    runs = ['habitat-harsh-aws+build','habitat-scarce-aws+build','habitat-pressure-aws+build','habitat-abundant-aws+build']
    d = sorted([summary[r] for r in runs], key=lambda x: x['abundance'])
    x = np.array([z['abundance'] for z in d]); y = np.array([z['skill'] * 100 for z in d])
    c = np.array([z['collab'] * 100 for z in d])
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    ax.plot(x, y, marker='o', lw=2.1, color='#2F6B9A', label='Skill %')
    ax.plot(x, c, marker='s', lw=1.8, ls='--', color='#D98943', label='Collab %')
    for xi, yi, lab in zip(x, y, ['Harsh (0.35)','Scarce (0.45)','Pressure (0.55)','Abundant (1.40)']):
        ax.annotate(lab, (xi, yi), xytext=(0, 9), textcoords='offset points', ha='center', fontsize=8)
    ax.set_xlabel('Resource abundance'); ax.set_ylabel('Run mean (%)')
    ax.set_title('Fig.8 Resource abundance vs. population behavior (hump-shaped)')
    ax.grid(alpha=.25); ax.legend(frameon=False)
    ax.set_ylim(0, max(max(y), max(c)) + 5)
    fig.tight_layout(); p = CHART_DIR / 'fig8_abundance_hump.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig); return p


def chart_agent_profiles(rows):
    sel = [r for r in rows if r['run'] in ('aws+build-confrontation', 'aws+build-mixed')]
    labels = [f"{r['agent']} ({r['run'].split('-')[-1]})" for r in sel]
    s = [float(r['skill_pct']) * 100 for r in sel]
    cl = [float(r['collab_pct']) * 100 for r in sel]
    res = [float(r['residual_pct']) * 100 for r in sel]
    y = np.arange(len(sel))
    fig, ax = plt.subplots(figsize=(7.2, 5.0))
    ax.barh(y, s, label='Skill', color='#2F6B9A')
    ax.barh(y, cl, left=s, label='Collaboration', color='#D98943')
    ax.barh(y, res, left=np.array(s) + np.array(cl), label='Residual', color='#C7CDD4')
    ax.set_yticks(y, labels, fontsize=7); ax.invert_yaxis(); ax.set_xlim(0, 100)
    ax.set_xlabel('Agent behaviour (%)'); ax.set_title('Fig.9 Agent-level behaviour profiles (confrontation & mixed)')
    ax.legend(frameon=False, ncol=3, loc='lower right'); ax.grid(axis='x', alpha=.2)
    fig.tight_layout(); p = CHART_DIR / 'fig9_agent_profiles.png'; fig.savefig(p, bbox_inches='tight'); plt.close(fig); return p


def _style_run(run):
    run.font.name = 'Times New Roman'
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)


def _add_p(anchor_para, text='', bold=False, style='Normal', first_line=True):
    new_p_el = OxmlElement('w:p')
    anchor_para._p.addnext(new_p_el)
    p = Paragraph(new_p_el, anchor_para._parent)
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    if first_line and style == 'Normal':
        pf.first_line_indent = Cm(0.74)
    r = p.add_run(text or '')
    if text:
        r.bold = bold
        _style_run(r)
    return p


def _add_heading(anchor_para, text, level=2):
    p = _add_p(anchor_para, '', first_line=False)
    pf = p.paragraph_format; pf.space_before = Pt(7)
    r = p.runs[0] if p.runs else p.add_run('')
    r.text = text
    r.bold = True
    r.font.size = Pt(12) if level == 2 else Pt(11)
    return p


def _add_table(doc, anchor_para, headers, data_rows):
    """Add table: we create via doc's table, then move XML after anchor."""
    ncols = len(headers)
    # Create table via doc
    tbl = doc.add_table(rows=0, cols=ncols, style='Table Grid')
    # Header row
    hdr = tbl.add_row()
    for j, h in enumerate(headers):
        cell = hdr.cells[j]; cell.text = str(h)
        for run in cell.paragraphs[0].runs: run.bold = True
    # Data rows
    for row in data_rows:
        cells = tbl.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
    # Move tbl XML after anchor
    anchor_para._p.addnext(tbl._tbl)
    # Remove tbl from doc's body so it's not duplicated
    doc.element.body.remove(tbl._tbl)
    # Spacer
    sp = _add_p(Paragraph(tbl._tbl, anchor_para._parent), '', first_line=False)
    return sp


def _add_img(anchor_para, path, width_cm=14.7):
    p = _add_p(anchor_para, '', first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(str(path), width=Cm(width_cm))
    return p


def insert_block(doc, anchor_text, build_fn):
    hits = [p for p in doc.paragraphs if anchor_text in p.text]
    if not hits:
        print(f'  WARN: anchor not found: {anchor_text}')
        return
    # build_fn receives (doc, anchor_paragraph)
    build_fn(doc, hits[0])


# ── 10 Technical Modules ──────────────────────────────────────────

def tech_dart(doc, a):
    p = _add_heading(a, '5.1.1 TSE管线工程参数与可复现实现', 3)
    p = _add_p(p, 'DART-Net在生产代码中对应TSE（TCN-Skill-Extractor，src/backend/agents/tse/）。Stage 1以BLAKE2b确定性哈希（hash_seed=20260716）对每条Plaza消息的content字段执行character 1/2/3-gram特征哈希，生成256维嵌入并L2归一化。同时编码角色（12类）、仪式信号（10类）、niche角色（6类）和轮次（16档）辅助嵌入，各以0.1/0.1/0.1/0.05权重叠加。配置embed_dim=256，max_utterances=64，max_chars_per_utterance=800。')
    p = _add_p(p, 'Stage 2采用三层深度可分离膨胀卷积（DilatedConvBlock），每层执行depthwise conv(C×K)→pointwise 1×1(C_out×C_in)→LayerNorm→ReLU→residual连接。dilations=[1,2,4]，kernel_size=3；理论感受野RF=1+2×(1+2+4)=15（单侧）≈29（全窗）。输入投影W_in与输出投影W_out均为ℝ²⁵⁶×²⁵⁶，权重初始化为He式缩放（depthwise scale=√(2/k)×0.5，pointwise scale=√(2/C)×0.5）。')
    p = _add_p(p, 'Stage 3以5组可学查询探针{name, description, category, tools, instructions}对TCN输出执行交叉注意力，产生field级skill_repr_k∈ℝ²⁵⁶和可追溯的focus_indices。冷启动以FIELD_KEYWORD_SEEDS先验(权重0.3)初始化注意力偏置。Stage 4的ConstrainedSkillDecoder支持ChatHarness在线、chat_fn异步和TSE+local离线三种后端，grammar_retry=1，输出经JSON schema验证。完整TSEConfig参数：tcn_hidden_dim=256，num_heads=4，top_k_utterances=8，decoder_temperature=0.2。')

def tech_extraction(doc, a):
    p = _add_heading(a, '5.2.1 三算法技能萃取与知识簇预处理', 3)
    p = _add_p(p, '生产SkillExtractorEngine（skill_extractor.py）以三种互补算法独立处理同一段Plaza讨论文本。算法1（去语境化）剥离AWS服务名、错误码等具体细节，抽取抽象动词—关系对，使技能跨平台可迁移。算法2（反面模式萃取）将事故日志、异议发言和失败回溯反转为防御规则和禁止约束。算法3（关键路径与最小动作集）从完整SOP中提取3—5个成败判断点，每个点对应独立技能。')
    p = _add_p(p, '长文档经_chunk_by_structure()按Markdown标题或空行分块，再以TF-IDF余弦相似度的凝聚聚类聚合成≤5个知识簇（max_clusters=5），每簇独立进入萃取管线。审核队列以pending→llm_prefilling→ready_for_review→approved/rejected状态机管理，通过source_fingerprint去重（SHA256+plaza_id+discussion_id+output_id）并持久化至storage/skill_extract_queue/。')

def tech_memory(doc, a):
    p = _add_heading(a, '6.1.1 四层记忆的精确算法', 3)
    p = _add_p(p, 'AgentMemoryCore（agent_memory_core.py）以JSON存储至storage/agent_memory/{team_id}/{agent_id}/。EpisodicLog记录{id,t,subject,action,detail,place,importance(1-10),tags,lastAccessAt}，召回评分score=recency+importance+relevance，其中recency=0.995ʰᵒᵘʳˢ，relevance为query与事件文本的character bigram Jaccard重叠度。')
    p = _add_p(p, 'PerceptionStream为容量500的FIFO缓冲区；compress()将缓冲汇总为单条EpisodicLog事件（action="感知压缩"），输出时间窗、各modality计数和fear均值后清空。IntentionQueue维护pending/confirmed/dropped三态，支持drop/escalate/keep超时策略，按dueAt排序输出daysLeft。AffectResidue以指数衰减τ=72h维护valence∈[-1,1]、arousal∈[0,1]和标签强度；重复感受按max(旧×1.2,新)更新以防止快速遗忘掩盖强情绪信号。')

def tech_transfer(doc, a):
    p = _add_heading(a, '6.2.1 Seal—Will—Export—Import的可审计传递协议', 3)
    p = _add_p(p, 'AgentMemoryTransfer.execute()（agent_memory_transfer.py）执行四阶段DNA式记忆遗传。传递原则为"复制非移动"且原稿默认可凭吊（keep_memorial=True）。seal写Schema=ag.legacy/v1的legacy.json遗体快照；will指定受益方、layers过滤、handover_intentions（ask_new_owner|auto|drop）与keep_memorial标志。')
    p = _add_p(p, 'export产生Schema=ag.memory/v1的完整层化JSON。import采用逐层merge：日志追加"传递继承"tag与[from:source]标记；感知流追加并截断至500条；意图按策略交接或丢弃；情绪标签取max强度合并，valence/arousal新旧各50%加权。系统记录importance=9的"记忆继承"事件及transfer_id审计追踪。')
    p = _add_p(p, '该协议将"身份延续"限制为带source_agent标签的显式来源分区，用户始终知晓哪些内容是继承而来。')

def tech_evolution(doc, a):
    p = _add_heading(a, '7.5 LLM-as-Judge演化评估管线', 3)
    p = _add_p(p, '技能演化引擎（evolution/fitness.py）对每个变体候选执行simulate→judge双阶段评估。Qwen3 Judge独立给予instruction_following、output_quality、conciseness三维0—1评分。复合适应度F=0.4×following+0.4×quality+0.2×conciseness。SkillFitnessReport聚合均值、逐例得分和composite<0.5的失败集。')
    p = _add_p(p, '同时对变体长度膨胀引入惩罚：ratio=evolved_len/original_len≤1时无惩罚，1<ratio≤1.5线性扣减至多0.2，ratio>1.5则归零。该设计抑制"以冗长购买表面完整"的作弊式变异，保留技能简洁度。')

def tech_classifier(doc, a):
    p = _add_heading(a, '7.6 三池分类器与防抖毕业', 3)
    p = _add_p(p, 'Skill Classifier（skill_classifier.py）判定技能归入exclusive/general/reserve三池。强制储备条件：lifecycle=degraded、已有使用但effectiveness<0.4、超90天未用或total_uses=0。通用需≥2团队采用或≥2类场景验证且gate_ok。特有需单团队使用占比≥0.8、effectiveness≥0.6且meets_rubric。')
    p = _add_table(doc, p, ['参数', '值', '含义'],
        [['EXCLUSIVE_TEAM_SHARE','0.8','特有技能单团队占比'],
         ['EXCLUSIVE_MIN_EFFECTIVENESS','0.6','特有最低效果'],
         ['GENERAL_MIN_TEAMS','2','通用跨团队采用数'],
         ['GENERAL_MIN_CATEGORIES','2','通用跨场景类目数'],
         ['RESERVE_MAX_EFFECTIVENESS','0.4','低效强制储备'],
         ['STALE_DAYS','90','未用天数触发储备'],
         ['GRADUATE_STREAK','2','连续达标毕业周期'],
         ['DEMOTE_GRACE','1','降级宽限周期']])
    p = _add_p(p, 'classify_with_history() 防止边界振荡：毕业需连续2周期达标，降级1周期宽限。萃取后的技能通过seed_reserve_from_extraction()首先写入储备池，由使用、验证和采纳证据驱动渐进毕业。')

def tech_appendix(doc, a):
    p = _add_heading(a, '附录A 工程部署架构、API路由与Token治理', 2)
    p = _add_p(p, 'AgentsGroup2026以Kubernetes部署，k8s/目录定义Namespace、ConfigMap、Secret、Deployment（主服务+Teams独立扩展）及Service。Dockerfile采用多阶段构建（Python 3.11），暴露端口8000(HTTP)与8001(SSE/WebSocket)。')
    p = _add_p(p, '核心FastAPI路由包括：/api/v1/plaza/*（Plaza讨论创建/参与/共识）；/api/v1/agent-memory/{team_id}/{agent_id}/*（记忆CRUD/导出/导入/封存）；/api/v1/skills/extract/*（萃取队列/审核/批准）；/api/v1/skills/evolution/*（演化AB测试）；/api/v1/skills/classify/*（三池分类）；/api/v1/cost/*与/api/v1/token/*（成本核算与治理）。')
    p = _add_p(p, 'Token治理子系统（token_governance/）在每次LLM调用前组合六类杠杆：行为注入（成本意识system prompt）、代码图谱桥接（复用上下文）、成本分级（三档模型分配）、渐进历史（滑动窗口截断）、提示词精简（LLM自压缩）和工具输出压缩。SavingsStore持久化每次优化节省的token数与金额。')


# ── New §9.7 Experiment Section ──────────────────────────────────

def tech_experiments(doc, a, charts, summary):
    p = _add_heading(a, '9.7 多智能体生态仿真：群落组装、资源压力与行为涌现', 2)
    p = _add_p(p, '§9.6的短周期闭环收敛仅覆盖3次迭代。为在群体尺度上补充定量观测，本节分析AgentsGroup2026实验记录experiments_raw.tsv中的9个运行组、45条Agent级观测（每组5个Agent）。每条记录报告Agent的技能使用、协作交互与残余行为比例，并附带竞合编排、代际数及生境参数。所有值取自原始文件的可复核字段，结论应理解为探索性证据。')

    p = _add_heading(p, '9.7.1 实验设计与主要指标', 3)
    p = _add_p(p, '群落组装实验对照四种编排：solo（同域团队独立运行）、division（跨团队职责分工）、confrontation（竞争性选择）和mixed（竞合并存）。生境实验固定division编排，在harsh(0.35)、scarce(0.45)、pressure(0.55)和abundant(1.40)四种资源丰度下运行。主要指标为运行均值skill_pct（技能使用的行为比例）与collab_pct（协作交互的行为比例）。')

    p = _add_table(doc, p, ['运行组','编排','压力','Agent数','技能%','协作%','代际','最优T'],
        [['aws-solo-division','solo','base','5','0.0','0.9','1','45'],
         ['build-solo-division','solo','base','5','10.7','15.2','2','75'],
         ['aws+build-division','division','base','5','6.8','15.0','1','69'],
         ['aws+build-confrontation','confrontation','base','5','8.8','10.6','1','77'],
         ['aws+build-mixed','mixed','base','5','5.0','8.2','9','57'],
         ['habitat-harsh','division','harsh(0.35)','5','4.8','2.5','1','60'],
         ['habitat-scarce','division','scarce(0.45)','5','3.7','13.6','1','61'],
         ['habitat-pressure','division','pressure(0.55)','5','7.9','13.2','1','67'],
         ['habitat-abundant','division','abundant(1.40)','5','2.8','4.9','2','86']])

    p = _add_heading(p, '9.7.2 群落组装模式比较', 3)
    p = _add_img(p, charts[0], 14.7)
    p = _add_p(p, '图7 不同群落组装模式下的技能与协作行为（数据源：experiments_raw.tsv）。')
    p = _add_p(p, '图7显示confrontation的技能行为均值（8.8%）高于division（6.8%）和mixed（5.0%），而division的协作均值（15.0%）高于confrontation（10.6%）。这说明竞争性筛选在当前参数下有助于放大可区分的技能行为，职责分工更有利于维持协作密度。mixed组经历9代后出现5项dominant技能，提示更长迭代可能促进跨团队技能占优。由于仅有一次运行，本结果不应解释为显著性结论。')

    p = _add_heading(p, '9.7.3 资源丰度与技能涌现', 3)
    p = _add_img(p, charts[1], 14.7)
    p = _add_p(p, '图8 资源丰度、生境压力与群体行为的关系（数据源：experiments_raw.tsv）。')
    p = _add_p(p, '图8呈现非单调关系：资源丰度从0.35升至0.55时，技能行为由4.8%升至7.9%；丰度进一步升至1.40后降为2.8%。该结果与"中等环境压力最有利于技能探索与选择"的倒U型假说一致。协作行为在稀缺(13.6%)与适中压力(13.2%)下较高，极端环境下均下降，提示过强约束与过度宽松都会弱化群体的有效互动。')

    p = _add_heading(p, '9.7.4 Agent级行为剖面与解释边界', 3)
    p = _add_img(p, charts[2], 14.7)
    p = _add_p(p, '图9 对抗与混合模式下的Agent级行为构成（skill+collab+residual=100%）。')
    p = _add_p(p, '图9揭示运行均值背后的显著异质性：在confrontation中aws_lead和aws_mon呈现较高技能/协作份额，而部分Build Agent的residual占比偏高；mixed组中技能优势分布更分散。这一模式支持将技能种群视为具有角色依赖的生态结构。因每组仅5个Agent且无独立种子重复，本文不报告显著性检验。后续工作应实施多种子重复、预注册阈值和跨域复验。')

    p = _add_heading(p, '9.7.5 与统一闭环的关联', 3)
    p = _add_p(p, '生态实验将§8中的耦合动力学具体化：竞合编排影响Plaza中角色交互和选择压力，进而改变技能种群的使用与协作构成；生境参数相当于对执行成本、失败风险和资源预算的外生扰动；dominant技能及其谱系可由技能分类器、适应度报告和记忆传递记录追溯。因此生态数据是闭环在群体尺度上的补充观测，而非独立结果。')


# ── Main ─────────────────────────────────────────────────────────

def main():
    rows, summary = parse_data()
    charts = [chart_competition(summary), chart_hump(summary), chart_agent_profiles(rows)]

    doc = DocxDocument(ORIGINAL)
    print(f'Original: {len(doc.paragraphs)} paragraphs')

    # 10 modules
    insert_block(doc, '5.1 层次化编码与约束解码', tech_dart)
    insert_block(doc, '5.2 技能Schema与生命周期扩展', tech_extraction)
    insert_block(doc, '6.1 四层并行记忆', tech_memory)
    insert_block(doc, '6.2 Seal—Will—Export—Import遗传链', tech_transfer)
    insert_block(doc, '7.3 变异、选择、组合与退役', tech_evolution)
    insert_block(doc, '7.4 双轨知识遗传', tech_classifier)
    insert_block(doc, '9.6 初步闭环收敛', lambda d,a: tech_experiments(d,a,charts,summary))
    insert_block(doc, '11 结束语', tech_appendix)

    doc.save(OUT)
    print(f'Saved: {OUT}')
    print(f'Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.inline_shapes)} inline shapes')
    for c in charts:
        print(f'  Chart: {c}')


if __name__ == '__main__':
    main()
