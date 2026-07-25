# -*- coding: utf-8 -*-
"""Deepen the rewritten AgentsGroup2026 paper while preserving the source DOCX format."""
from __future__ import annotations

import copy
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path("/Users/panglaohu/OpenWorker/5232097c-f7c")
SRC = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文_重写含实验绘图版.docx")
DST = Path("/Users/panglaohu/Downloads/协商审议_DARTNet_记忆遗传_统一闭环论文_架构深化版.docx")
CHARTS = ROOT / "deep_charts"
CHARTS.mkdir(exist_ok=True)

plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "PingFang SC", "Heiti SC", "SimHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def draw_box(ax, xy, w, h, title, subtitle, color):
    x, y = xy
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.015,rounding_size=0.02",
                           linewidth=1.3, edgecolor=color, facecolor=color + "18")
    ax.add_patch(patch)
    ax.text(x + w/2, y + h*0.64, title, ha="center", va="center", fontsize=9.2, weight="bold", color=color)
    ax.text(x + w/2, y + h*0.30, subtitle, ha="center", va="center", fontsize=7.2, color="#333333", wrap=True)


def arrow(ax, a, b, color="#4B5563", rad=0):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=12,
                                 linewidth=1.25, color=color,
                                 connectionstyle=f"arc3,rad={rad}"))


def generate_closed_loop_chart():
    fig, ax = plt.subplots(figsize=(10, 5.8))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    boxes = [
        ((.03,.68),.17,.17,"Plaza协商审议","ORID / 角色 / 共识","#275D8C"),
        ((.24,.68),.17,.17,"ExecutionPlan","步骤、依赖、验收","#275D8C"),
        ((.45,.68),.17,.17,"TaskHabitatContract","niche / demand / budget","#287A5A"),
        ((.66,.68),.14,.17,"孪生竞标","C0 + R1—R5","#B66B24"),
        ((.84,.68),.13,.17,"Ratchet","质量门+成本棘轮","#A33A3A"),
        ((.82,.27),.15,.17,"生产派发","胜出团队/技能/模型","#6A4C93"),
        ((.61,.27),.17,.17,"EvidenceRun","质量、成本、轨迹、哈希","#4B6478"),
        ((.39,.27),.17,.17,"技能演化/路由","TSE→验证→Router","#287A5A"),
        ((.17,.27),.17,.17,"记忆生命周期","共享、传递、回放","#6A4C93"),
        ((.03,.27),.10,.17,"反馈","讨论/计划","#275D8C"),
    ]
    for args in boxes: draw_box(ax, *args)
    top_centers=[(.20,.765),(.41,.765),(.62,.765),(.80,.765),(.97,.765)]
    for a,b in zip(top_centers[:-1],top_centers[1:]): arrow(ax,a,b)
    arrow(ax,(.905,.68),(.895,.44))
    arrow(ax,(.82,.355),(.78,.355))
    arrow(ax,(.61,.355),(.56,.355))
    arrow(ax,(.39,.355),(.34,.355))
    arrow(ax,(.17,.355),(.13,.355))
    arrow(ax,(.08,.44),(.08,.68))
    arrow(ax,(.69,.27),(.54,.17),color="#8B5CF6",rad=.15)
    ax.text(.58,.12,"执行结果同时回注技能适应度、记忆与下一轮审议",ha="center",fontsize=8.2,color="#5B3F86")
    ax.set_title("计划驱动的协商—孪生竞标—技能与记忆统一闭环", fontsize=13, weight="bold", pad=12)
    p=CHARTS/"fig_plan_driven_closed_loop.png"
    fig.tight_layout(); fig.savefig(p,dpi=220,bbox_inches="tight"); plt.close(fig)
    return p


def generate_bidding_chart():
    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off")
    draw_box(ax,(.04,.67),.19,.18,"基线 C0","团队/技能/顺序/模型","#275D8C")
    labels=[("R1","换角色"),("R2","换技能"),("R3","并行化"),("R4","加Review"),("R5","模型降档")]
    xs=[.28,.42,.56,.70,.84]
    for x,(r,t) in zip(xs,labels): draw_box(ax,(x,.68),.11,.15,r,t,"#B66B24")
    for x in xs: arrow(ax,(.23,.76),(x,.76),rad=(x-.5)*.15)
    draw_box(ax,(.23,.37),.22,.16,"孪生试炼","success / quality / token","#287A5A")
    draw_box(ax,(.53,.37),.18,.16,"双门槛","success≥0.9 且 quality≥0.9","#A33A3A")
    draw_box(ax,(.78,.37),.18,.16,"达标排序","token升序；同token质量降序","#B66B24")
    arrow(ax,(.56,.68),(.38,.53)); arrow(ax,(.45,.45),(.53,.45)); arrow(ax,(.71,.45),(.78,.45))
    draw_box(ax,(.31,.08),.24,.16,"RatchetLedger","效率=quality/max(token,1)\n只进不退","#6A4C93")
    draw_box(ax,(.66,.08),.23,.16,"回流与入账","胜者回流Plaza；simulation成本分账","#4B6478")
    arrow(ax,(.87,.37),(.55,.24)); arrow(ax,(.55,.16),(.66,.16))
    ax.set_title("候选组合竞标、质量门禁与棘轮锁定",fontsize=13,weight="bold",pad=10)
    p=CHARTS/"fig_bidding_ratchet.png"
    fig.tight_layout(); fig.savefig(p,dpi=220,bbox_inches="tight"); plt.close(fig)
    return p


def style_run(run, size=10.5, bold=False):
    run.font.name="Times New Roman"
    rpr=run._element.get_or_add_rPr()
    rf=rpr.find(qn("w:rFonts"))
    if rf is None:
        rf=OxmlElement("w:rFonts"); rpr.insert(0,rf)
    rf.set(qn("w:eastAsia"),"宋体")
    run.font.size=Pt(size); run.bold=bold


def p_after(anchor,text="",heading=False,size=10.5,indent=True,align=None):
    el=OxmlElement("w:p"); anchor._p.addnext(el); p=Paragraph(el,anchor._parent)
    p.paragraph_format.space_after=Pt(4)
    if indent: p.paragraph_format.first_line_indent=Cm(.74)
    if heading: p.paragraph_format.space_before=Pt(7); p.paragraph_format.first_line_indent=Cm(0)
    if align is not None: p.alignment=align
    r=p.add_run(text); style_run(r,size,bold=heading)
    return p


def image_after(anchor,path,width=14.8):
    p=p_after(anchor,"",indent=False,align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path),width=Cm(width)); return p


def table_after(anchor,headers,rows):
    tbl=OxmlElement("w:tbl")
    pr=OxmlElement("w:tblPr")
    st=OxmlElement("w:tblStyle"); st.set(qn("w:val"),"TableGrid"); pr.append(st)
    w=OxmlElement("w:tblW"); w.set(qn("w:w"),"5000"); w.set(qn("w:type"),"pct"); pr.append(w); tbl.append(pr)
    grid=OxmlElement("w:tblGrid")
    for _ in headers:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(9000//len(headers))); grid.append(gc)
    tbl.append(grid)
    def row(vals,head=False):
        tr=OxmlElement("w:tr")
        for val in vals:
            tc=OxmlElement("w:tc"); tcp=OxmlElement("w:tcPr");tc.append(tcp)
            pp=OxmlElement("w:p"); rr=OxmlElement("w:r")
            rpr=OxmlElement("w:rPr")
            rfs=OxmlElement("w:rFonts");rfs.set(qn("w:ascii"),"Times New Roman");rfs.set(qn("w:eastAsia"),"宋体");rpr.append(rfs)
            sz=OxmlElement("w:sz");sz.set(qn("w:val"),"18");rpr.append(sz)
            if head:rpr.append(OxmlElement("w:b"))
            rr.append(rpr);tt=OxmlElement("w:t");tt.text=str(val);rr.append(tt);pp.append(rr);tc.append(pp);tr.append(tc)
        return tr
    tbl.append(row(headers,True))
    for r in rows: tbl.append(row(r))
    anchor._p.addnext(tbl); sp=OxmlElement("w:p");tbl.addnext(sp);return Paragraph(sp,anchor._parent)


def find(doc,text):
    for p in doc.paragraphs:
        if text in p.text:return p
    raise RuntimeError(f"anchor not found: {text}")


def insert_method_extensions(doc,closed_chart,bid_chart):
    a=find(doc,"3.2 跨模块身份与来源关系")
    p=p_after(a,"3.3 从执行计划到任务生境契约",heading=True,size=12,indent=False)
    p=p_after(p,"统一闭环并不直接把Plaza结论交给生态仿真，而是先编译为TaskHabitatContract。输入ExecutionPlan中的步骤经依赖拓扑排序，形成有序NicheWindow；每个窗口包含step_id、demanded_skills、responsible_role、acceptance、base_ticks、depends_on与inferred_skills。显式required_skills优先；缺失时由步骤标题、描述和角色启发式推断并保留inferred_skills标记。")
    p=p_after(p,"第j个生态位的基础时钟定义为b_j=max(8,12+4·max(|K_j|,1))。总步数预算B=clamp(Σ_j b_j,40,500)，世代预算G=clamp(2+⌈|N|/3⌉,1,10)。契约以SHA-256截断指纹绑定plan_id、revision与niche三元组，从而保证演练结果能够回溯到具体计划版本。该契约把“讨论如何做”转化为“环境选择什么”，但不改变survival_ticks作为生态演练唯一原生适应度。")
    p=image_after(p,closed_chart)
    p=p_after(p,"图 1A 计划驱动的协商—数字孪生竞标—技能与记忆统一闭环",indent=False,align=WD_ALIGN_PARAGRAPH.CENTER)

    a=find(doc,"7.4 双轨知识遗传")
    p=p_after(a,"7.7 SkillRouter两阶段检索、生命周期重排与反馈学习",heading=True,size=12,indent=False)
    p=p_after(p,"SkillRouter把已验证技能注入后续Agent上下文。Stage 1在全技能池上执行BM25、TF-IDF余弦、中文bigram/trigram、描述短语、指令深匹配与同义词扩展，取max(20,3K)个候选；Stage 2按name、description、instructions、category/tools字段联合重排。最终基础分为S=0.45S_retrieval+0.55S_rerank。")
    p=p_after(p,"生命周期乘子进一步修正排序：solidified=1.14、verified=1.12、published=1.08、team_local=1.00、draft=0.90、degraded=0.72。用户rating/revoke反馈更新(agent,category)亲和度，限制在[-0.5,0.5]；技能赋予同时把数字孪生熟练度先验提升到至少0.8，但不等同于已经通过验证。该区分避免把“被路由”误当作“有效”。")

    a=find(doc,"8.3 三维控制界面")
    p=p_after(a,"8.4 计划驱动竞标与正向棘轮",heading=True,size=12,indent=False)
    p=p_after(p,"对同一任务契约，竞标编排器由基线C0生成最多4个单算子候选：R1替换角色、R2替换技能绑定、R3并行化无依赖步骤、R4增加Review回边、R5降低模型档。每个候选在数字孪生中产生(success_rate,quality_score,token_consumed,collab_heat)。只有success_rate≥0.9且quality_score≥0.9的候选进入达标组；达标组按token升序、同token按质量降序排名，不达标组仅作诊断。")
    p=p_after(p,"胜者效率定义为η=quality/max(token,1)。RatchetLedger以scenario_best:{task_type}为共享键，仅在η不低于当前值且满足最小增量时推进generation；退步被拒绝，容忍区间内保持held。竞标token记为simulation成本，不混入生产效能；竞标结论可回写Plaza讨论时间线与ExecutionPlan。")
    p=image_after(p,bid_chart)
    p=p_after(p,"图 5A 候选组合竞标、双门槛质量筛选与棘轮锁定",indent=False,align=WD_ALIGN_PARAGRAPH.CENTER)

    p=p_after(p,"8.5 EvidenceRun证据链与唯一适应度约束",heading=True,size=12,indent=False)
    p=p_after(p,"EvidenceRun是技能验证、任务执行、演化比较和成本门禁共享的追加式证据对象。它关联team_id、agent_id、skill_id、task_id、evolution_item_id、cost_target_id、plaza_topic_id、session_id与request_id，并保存runtime、command、exit_code、artifact_dir、stdout/stderr、metrics_before/after及detail。对象采用SHA-256截断evidence_hash进行完整性校验，按月落盘。")
    p=p_after(p,"生态演练中T_i=survival_ticks仍是唯一原生适应度。skill_ticks、collab_ticks与residual_ticks只是对每个存活tick主因的解释性归因，满足三者之和等于T_i；采样缺失计入residual。归因优先级为收到分享→自身技能成功觅食→跟随协作成功→跟随减损→残余行为。该机制不引入第二适应度，避免用人工质量分替代生态选择。")

    p=table_after(p,["工程对象","核心字段/规则","论文角色"],[
        ["TaskHabitatContract","niches, demanded_skills, step_budget, fingerprint","把计划编译为选择环境"],
        ["SkillRouter","BM25/TF-IDF/ngram→field rerank→lifecycle mult","把验证技能路由到Agent"],
        ["BiddingOrchestrator","C0+R1—R5; success/quality≥0.9","同计划下选择成本有效候选"],
        ["RatchetLedger","generation,value,evidence; reject regression","锁定单调改进"],
        ["EvidenceRun","object ids, metrics, artifacts, evidence_hash","统一可验证证据链"],
        ["SurvivalDecompose","skill+collab+residual=T_i","唯一适应度的可解释归因"],
    ])
    p=p_after(p,"表 2A 计划驱动闭环新增的核心工程对象与约束",indent=False,align=WD_ALIGN_PARAGRAPH.CENTER)

    a=find(doc,"6.3 选择性遗传与受控遗忘")
    p=p_after(a,"6.4 生命周期状态机、共享ACL与运行时反馈",heading=True,size=12,indent=False)
    p=p_after(p,"记忆层不仅支持seal与transfer，还实现unbound、active、shared、sealed、archived和destroyed状态约束，并以audit.jsonl与tombstone保留生命周期证据。共享接口使用share_grants ACL与layer_mask控制可见层，默认不共享affect；人格配置支持xiaoman、shenmian与hybrid三类Persona及autonomy默认策略。")
    p=p_after(p,"运行时由chat_harness注入tone_hint与recall，但Plaza phase跳过该注入以避免个体情绪污染协商议程。任务完成/失败事件自动写入EpisodicLog与AffectResidue；tool_loop将工具观察写入PerceptionStream并在阈值达到时自动压缩。该设计使Memory→Plaza反馈是受阶段、权限和来源约束的，而非无条件拼接历史。")


def fix_texts(doc):
    replacements={
        "图7 不同群落组装模式下":"图8 不同群落组装模式下",
        "图8 资源丰度":"图9 资源丰度",
        "图9 对抗与混合":"图10 对抗与混合",
        "图7显示confrontation":"图8显示confrontation",
        "图8呈现非单调关系":"图9呈现非单调关系",
        "图9揭示运行均值":"图10揭示运行均值",
        "暴露端口8000(HTTP)和8001(SSE)":"后端运行端口8080，前端开发端口5173",
    }
    for p in doc.paragraphs:
        old=p.text
        new=old
        for a,b in replacements.items():new=new.replace(a,b)
        if new!=old:
            for r in p.runs:r.text=""
            if p.runs:p.runs[0].text=new
            else:p.add_run(new)


def copy_block_after(ref, elements):
    cur=ref._element
    for el in elements:
        new=copy.deepcopy(el)
        cur.addnext(new);cur=new


def move_experiment_block(doc):
    body=doc.element.body
    children=list(body)
    # locate by paragraph text, then move the block from 9.7 through before Table 5 caption
    start=None; end=None; anchor=None
    for i,el in enumerate(children):
        if el.tag==qn("w:p"):
            txt="".join(el.itertext())
            if "9.7 多智能体生态仿真" in txt:start=i
            if start is not None and "表 5 闭环技能数量与质量变化" in txt:
                end=i;anchor=el;break
    if start is None or end is None:return False
    block=children[start:end]
    # remove current block
    for el in block: body.remove(el)
    # re-find anchor after removals. Insert after Table 5 English caption, not before it.
    children=list(body);anchor_idx=None
    for i,el in enumerate(children):
        if el.tag==qn("w:p") and "Table 5 Closed-loop changes" in "".join(el.itertext()):
            anchor_idx=i;break
    if anchor_idx is None:
        for i,el in enumerate(children):
            if el.tag==qn("w:p") and "表 5 闭环技能数量与质量变化" in "".join(el.itertext()):anchor_idx=i
    cur=children[anchor_idx]
    for el in block:
        cur.addnext(el);cur=el
    return True


def move_appendix(doc):
    body=doc.element.body;children=list(body)
    start=None;end=None
    for i,el in enumerate(children):
        if el.tag==qn("w:p"):
            txt="".join(el.itertext())
            if txt.startswith("附录A 工程部署架构"):start=i
            if start is not None and "参 考 文 献" in txt:end=i;break
    if start is None or end is None:return False
    block=children[start:end]
    for el in block:body.remove(el)
    children=list(body);ref=None
    for el in children:
        if el.tag==qn("w:p") and "参 考 文 献" in "".join(el.itertext()):ref=el;break
    # insert immediately before references
    idx=body.index(ref)
    for el in block:
        body.insert(idx,el);idx+=1
    return True


def main():
    closed=generate_closed_loop_chart();bid=generate_bidding_chart()
    doc=Document(SRC)
    insert_method_extensions(doc,closed,bid)
    fix_texts(doc)
    moved_exp=move_experiment_block(doc)
    moved_app=move_appendix(doc)
    doc.save(DST)
    print("saved",DST)
    print("moved_exp",moved_exp,"moved_appendix",moved_app)
    print("paragraphs",len(doc.paragraphs),"tables",len(doc.tables),"shapes",len(doc.inline_shapes),"sections",len(doc.sections))

if __name__=="__main__":main()
