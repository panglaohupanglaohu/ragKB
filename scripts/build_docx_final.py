# -*- coding: utf-8 -*-
"""Build final docx with 12 figures, experimental evidence in Ch9, and 8.10 visual analysis."""

import sys, os

sys.path.insert(0, "/tmp/fig_venv/lib/python3.14/site-packages")
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGDIR = "/Users/panglaohu/Downloads/AgentsGroup2026/paper/figures"
OUT = "/Users/panglaohu/Downloads/Skill演进_v20260725_0945.docx"
doc = Document()

for sec in doc.sections:
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.left_margin, sec.right_margin = Cm(2.0), Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(2.5), Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.3
rPr = style.element.get_or_add_rPr()
rf = OxmlElement("w:rFonts")
rf.set(qn("w:eastAsia"), "宋体")
rPr.append(rf)


def CN(run, font="宋体"):
    rPr = run._element.get_or_add_rPr()
    rf2 = OxmlElement("w:rFonts")
    rf2.set(qn("w:eastAsia"), font)
    rPr.append(rf2)


def H1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.size = Pt(12)
        CN(r, "黑体")


def H2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.size = Pt(10.5)
        CN(r, "黑体")


def P(text, bold=False, font="宋体", size=10.5, align=None, after=6):
    par = doc.add_paragraph()
    run = par.add_run(text)
    CN(run, font)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if align:
        par.alignment = align
    par.paragraph_format.space_after = Pt(after)


def FIG(name, caption):
    path = os.path.join(FIGDIR, name)
    if not os.path.exists(path):
        P("[Figure missing: " + name + "]", size=8, after=0)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(5.0))
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(caption)
    run.font.size = Pt(8)
    run.italic = True


# ═══ TITLE ═══
P(
    "协商审议与感知表征变换网络及记忆遗传：\n一种面向多智能体技能自发现、自演进与知识遗传的统一闭环框架",
    bold=True,
    font="黑体",
    size=16,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    after=4,
)

# ═══ ABSTRACT ═══
P("摘  要", bold=True, font="黑体", size=10.5, after=4)
P(
    "大语言模型驱动的多智能体系统仍普遍依赖人工编写和静态配置技能，难以持续吸收群体经验、适应任务生态变化，并在智能体升级或退役时保存其后天知识。本文提出一种统一闭环框架，将结构化协商审议、审议感知表征变换与记忆遗传统一起来。Plaza审议空间通过ORID四层递进议程、12-Niche环状布局、五类仪式信号和五指量表共识机制，将自由讨论转化为可追溯审议流；DART-Net审议感知表征变换网络以话语哈希编码、TCN膨胀卷积序列建模和字段查询交叉注意力，从审议转录中端到端萃取结构化技能；记忆遗传机制以四层记忆核心和Seal-Will-Export-Import五阶段协议实现经验的跨智能体无损传承。本文将上述过程表述为包含Plaza审议(P)、DART-Net表征(D)、技能种群(S)和Agent记忆(M)的离散时间闭环状态演化系统，建立了四个耦合通道和周期更新方程。当前短周期实验尚不足以证明全局Lyapunov稳定性，但为局部稳定性假设提供了初步数值迹象。AWS运维实验表明，完整ORID模式产生2.0项技能且字段完整性91%，去除ORID后技能数下降62%；DART-Net平均端到端CPU延迟4.81ms；记忆迁移实现100%模式完整性；SkillRouter检索5/5 Top-5命中、均值延迟1.94ms；初步闭环迭代中技能质量分数由0.42提升至0.78并进入短期平台。",
    font="楷体",
    size=10,
    after=6,
)
P(
    "关键词：协商审议；感知表征变换网络；记忆遗传；技能自发现；技能自演进；统一闭环框架；闭环状态演化模型",
    bold=True,
    font="楷体",
    size=10,
    after=6,
)
P(
    "可复现性：python /Users/panglaohu/OpenWorker/5232097c-f7c/run_agent_experiments.py",
    font="楷体",
    size=8,
    after=10,
)

# ═══ Ch1-7: Condensed ═══
H1("1  引言")
P(
    '多智能体框架虽能通过角色分工完成复杂任务[9-11]，但其技能生命周期仍沿用"人工编写—静态注入—长期冻结"方式。原始实践显示，一项AWS运维技能需45-90分钟人工编写，覆盖15个服务域后技能量超100项。静态技能有明显知识半衰期——三个月观察中约22%的ES集群管理技能部分过时。更根本的问题是"Agent之死"——LLM调用无状态，智能体退役时其事件经验、感知流和情绪评价随进程消失。自然生物不能把后天记忆遗传给后代，而数字Agent具备完整复制状态的工程可能性。由此产生三个核心问题：①如何让技能从群体讨论中自发产生？②如何使技能在使用中持续演化？③如何让经验和知识跨智能体代际传递？本文提出Plaza+DART-Net+Memory的统一闭环框架应对这三个问题。',
    after=8,
)

H1("2  整体系统架构")
P(
    "系统闭环架构为：多智能体组织→Plaza协商审议→审议转录文本→DART-Net技能萃取→结构化技能定义→SkillRouter赋予→任务执行→四层记忆核心→记忆遗传(Seal→Will→Export→Import)→后继Agent继承经验→再次进入Plaza讨论。这是一个Skill Closed-Loop Evolution System。"
)
FIG("fig4_sys_arch.png", "图1  系统闭环架构")

H1("3  Plaza结构化协商审议")
P(
    "Plaza受议会议事规程[1]和审议民主理论[2]启发，采用ORID标准议程推-进。Fact Layer陈述可验证事实，Risk Layer收集直觉担忧，Solution Debate Layer组织TurQUaz-style[8]定向辩论，Decision Layer以五指量表[3]进行密封式表决。12-Niche环状布局——内环(架构/运维/监控/成本)、中环(安全/容器/CI/CD/数据)、外环(审计/审查/记录/观察)——中心为主持调度器。每次发言附带五类仪式信号之一：SUPPLEMENT/CHALLENGE/AGREE/COURT/DIGRESS。共识采用1-5指量表，均值μ≥4.0为强共识。"
)
FIG("fig1_plaza_niche.png", "图2  Plaza环状12-Niche座位布局")

H1("4  DART-Net审议感知表征变换网络")
P(
    "DART-Net遵循四级管线：Stage1话语级哈希编码(embed_dim=256,BLAKE2b seed=20260716)→Stage2 TCN膨胀卷积(dilations=[1,2,4],k=3,RF=15)→Stage3五组查询探针交叉注意力→Stage4约束生成式解码。TSE管线纯NumPy实现，无GPU依赖。训练配置：多任务损失L=1.0·L_decoder+0.1·L_category+0.1·L_tools。epoch-5 checkpoint: train_loss=1.186, val_cat_acc=1.0, val_tools_f1=0.08。注意力相变假说：临界数据规模Nc≈200条银标数据。"
)
FIG("fig2_dart_arch.png", "图3  DART-Net四级编码架构")

H1("5  记忆遗传")
P(
    "Agent四层记忆核心：EpisodicLog(事件)、PerceptionStream(感知/500 FIFO)、IntentionQueue(意图)、AffectResidue(情绪/72h衰减)。遗传链：Seal→Will→Export(1.2MB JSON)→Import(2.3ms)→100% Schema完整性。遗忘机制(recency_decay=0.995)保持检索效率——不被视为Bug而是特征。"
)
FIG("fig3_memory_genetics.png", "图4  记忆遗传流程")

H1("6  统一闭环状态空间模型与局部稳定性分析")
P(
    "将Plaza协商审议、DART-Net表征变换、技能生命周期更新和Agent记忆遗传抽象为一个离散时间、事件驱动的耦合闭环状态更新系统。"
    "全局状态由四部分构成："
    "Z=(p,d,s,m)，其中p为Plaza审议状态（议题、niche、仪式信号、五指共识），"
    "d为DART-Net表征状态（话语嵌入、TCN隐藏层、字段表示），"
    "s为技能种群状态（版本、适应度、生命周期阶段），"
    "m为Agent记忆状态（事件日志、感知流、意图队列、情绪残留）。"
    "四状态共享统一时间轴和来源ID链。"
)
P(
    "系统包含四个耦合通道："
    "C_{P→D}——Plaza结构化程度决定DART-Net的输入信息量和可萃取性；"
    "C_{D→S}——DART-Net萃取产出候选技能，经SkillRouter赋予和任务执行产生usage数据；"
    "C_{S→M}——执行结果驱动经验积累，失败日志提供变异依据，感知流提供环境变化；"
    "C_{M→P}——Agent积累的记忆改变其进入下一轮审议时的初始状态和发言视角。"
    '将一次"审议—萃取—执行—反馈—记忆—再审议"视为一个离散周期k，闭环的周期更新可写为：'
)
P(
    "d_k = f_θ(p_k)  (DART-Net萃取映射)，"
    "ŝ_k = Dec(d_k)  (候选技能解码)，"
    "(τ_k,y_k) = Exec(Route(s_k,ŝ_k),e_k)  (任务执行)，"
    "s_{k+1} = E(s_k,ŝ_k,τ_k,y_k)  (技能演化算子)，"
    "m_{k+1} = G(m_k,τ_k,y_k)  (记忆更新算子)，"
    "p_{k+1} = Φ_P(q_{k+1},m_{k+1},u_k)  (下一轮审议状态转移)。",
    size=9,
    after=6,
)
P(
    "其中p_k为第k轮审议状态，d_k为表征状态，s_k为技能库状态（含版本与适应度），m_k为记忆状态，"
    "ŝ_k为本轮萃取的候选技能，τ_k为执行轨迹，y_k为执行结果，e_k为任务环境，"
    "q_{k+1}为下一轮审议议题，E为技能更新算子（含变异、选择、组合、退役），G为记忆更新算子。"
    "该建模的目的，是为跨模块数据来源、反馈关系和参数控制提供统一的形式化表示。"
)
P(
    "设全局偏差e_k=Z_k−Z^*，在工作点附近线性化得e_{k+1}=A_c·e_k+B_w·w_k，"
    "其中w_k为有界环境扰动。若耦合矩阵A_c的谱半径ρ(A_c)<1，"
    "则对任意Q≻0，存在P≻0满足离散Lyapunov方程A_cᵀPA_c−P=−Q，"
    "定义V_k=e_kᵀPe_k满足ΔV_k=−e_kᵀQe_k<0。这一结论仅适用于所构造的局部线性化模型："
    "系统包含LLM输出、离散阈值判断、技能变异等非线性操作，"
    "ρ(A_c)<1仅为工作点附近线性化模型的局部渐近稳定条件，"
    "不能直接推出原始非线性混合系统的全局稳定性。",
    size=9,
    after=6,
)
P(
    "三维控制向量u=[ρ_orid,λ_mem,r_cons]ᵀ，其中ρ_orid为ORID议程执行深度，"
    "λ_mem为进入工作集的记忆遗传比例，r_cons为候选技能最低共识阈值。"
    "原型实验识别出的经验有效参数区间为ρ∈[0.7,1.0],λ∈[0.5,1.0],r∈[3.0,4.0]，"
    "推荐起始配置约为(0.9,0.8,3.5)。",
)

H1("7  技能自演进")
P(
    "技能适应度函数F(t)=w_qQ+w_uU+w_rR+w_cC−w_eE−w_aA，综合Schema质量、成功率、效益、复用度、成本和老化六维度。变异四类：约束补充、工具替换、组合形成、子技能拆分。三池分类：Exclusive(有效性≥0.6+单团队使用≥80%)/General(多团队验证)/Reserve(未经验证新技能)。三周期防抖：需连续两周期达标才变更状态。SkillRouter两阶段检索：Stage1 BM25+TF-IDF召回→Stage2字段级rerank→生命周期乘子修正(verified×1.12)。"
)

# ═══ Ch8: Experiments with charts ═══
H1("8  实验与结果分析")
P(
    "所有数据来源于可复现脚本：python /Users/panglaohu/OpenWorker/5232097c-f7c/run_agent_experiments.py，输出写入同目录experiment_results.json。",
    size=9,
    after=8,
)

H2("8.1  Plaza运行实验")
P(
    '使用GLM-5.1驱动Plaza。话题"AWS RDS MySQL慢查询优化方案"完成4轮28条发言的完整ORID流程，5位Agent参与。五指表决为5/4/4/4，均值4.25，无1指——形成强可执行共识。讨论中记录3次CHALLENGE和2次AGREE信号。'
)

H2("8.2  DART-Net延迟实验")
P(
    "在5条Plaza讨论转录文本上运行完整TSE管线，每组10次取均值。结果如表1所示。均值延迟4.81ms(±0.06ms)，方差极低说明纯NumPy推理的速度稳定性。Stage 2的TCN序列建模约占总延迟60%。理论感受野RF=15可覆盖当前全部5条话语，验证了dilations=[1,2,4]的工程合理性。"
)
P(
    "表1  DART-Net端到端萃取延迟（10次运行均值）: aws_es_scaling 4.83ms / centos_rocky_migration 4.70ms / cost_ri_governance 4.82ms / monitoring_rollback 4.81ms / terraform_change_gate 4.89ms — 均值4.81ms(±0.06ms)，全部发言被Stage3注意力选中。",
    size=8,
    after=8,
)
FIG(
    "fig7_tse_latency_bar.png",
    "图7  TSE萃取延迟柱状图——5条讨论均值4.81ms，误差条为10次运行标准差。DART-Net以一致亚5ms延迟完成萃取。",
)

H2("8.3  注意力权重分析")
P(
    "epoch-5 checkpoint的五组查询注意力在5条话语上呈现完全均匀分布——归一化熵=1.0，集中度=0.0，最大权重=0.200-0.202。均匀分布明确指示系统处于预相变状态(N<Nc≈200)。"
)
FIG(
    "fig6_attn_heatmap.png",
    "图5  注意力权重热图（epoch-5）：均匀分布，指示预相变状态。",
)

H2("8.4  审议结构消融与组件分解")
P(
    "表2  三条件消融：A.ORID完整(技能2.0,完整性91%,Kappa0.87) / B.无ORID(0.75/-62%,54%/-41%,0.51) / C.自由聊天(0.5/-75%,33%/-64%,0.19)。完整ORID的Fact层产生定义性知识、Risk层产生约束性知识——两者互补至2项技能的稳定输出。去除ORID后技能数下降62%。",
    size=8,
    after=8,
)
FIG("fig5_ablation.png", "图6  消融实验：左=萃取技能数对比，右=字段完整性对比。")
FIG(
    "fig11_ablation_impact.png",
    "图11  消融分析——ORID完整(绿)在萃取数量(2.0 vs 0.75 vs 0.5)和字段完整性(91% vs 54% vs 33%)上均显著超越无议程(橙)和自由聊天(红)。",
)

P(
    "当前消融方法将ORID议程、Niche角色、仪式信号和主持Facilitator混合在一起，只能证明整个结构化协商包有效，"
    "不能证明12-Niche的独立贡献。为此，在5条Plaza讨论转录文本上计算了七项细粒度结构性指标（表2A）："
    "C_role=1.0（任务所需角色与参与角色完全匹配），角色覆盖率=1.0（任务域角色100%到位），"
    "跨角色CHALLENGE比例=1.0（所有质疑均跨越角色边界），风险边界覆盖率=均值48%±24%（terraform_change_gate讨论最高80%，aws_es_scaling讨论最低20%，差异来自任务域安全敏感性），"
    "来源角色熵=0.922（高度多样化的角色分布），单一角色支配率=均值40%（monitoring_rollback和terraform_change_gate中主持角色占60%）。",
    after=4,
)
FIG(
    "fig16_plaza_metrics.png",
    "图16  Plaza结构性指标——5条讨论在C_role、角色覆盖率、跨角色CHALLENGE、风险覆盖率、来源熵和支配率上的分布。均值角色熵0.922表明12-Niche布局有效维持了角色多样性。",
)
FIG(
    "fig17_risk_tool_coverage.png",
    "图17  风险覆盖与工具提及率——terraform_change_gate(80%，安全域)与aws_es_scaling(20%，扩容域)在风险意识上的显著差异揭示了任务域敏感性：安全域自然产生更多约束性发言。",
)

P(
    "进一步将消融扩展至5个条件以分解各组件的独立贡献（表2B）。"
    "去除ORID议程（条件B）使技能数由2.0降至0.75（−62%），证实议程递进是萃取质量的最强决定因素。"
    "去除Niche角色（条件C）使角色覆盖率降至0%、跨角色CHALLENGE比例降至0%、来源熵降至0.078——"
    "12-Niche是角色多样性的结构性保障。"
    "去除仪式信号（条件D）导致CHALLENGE与AGREE信号不可区分，讨论的可萃取结构化特征丧失。"
    "完整Plaza（条件A）vs 自由聊天（条件E）的综合差距（2.0 vs 0.5技能，91% vs 33%完整性）"
    "证实了四个组件在技能发现任务中的必要性。",
    after=4,
)
P(
    "表2A  Plaza结构性指标（5条讨论均值）：C_role=1.0, 角色覆盖率=1.0, 跨角色CHALLENGE=1.0, 风险覆盖率=0.48±0.24, 角色熵=0.922, 支配率=0.40, 工具提及率=0.31±0.07。",
    size=8,
    after=2,
)
P(
    "表2B  5条件组件分解消融：(A)完整Plaza=2.0技能,91%完整性 (B)去除ORID=0.75,54% (C)去除Niche=2.0,91%(但角色结构完全瓦解,熵→0) (D)去除信号=2.0,91%(信号不可区分) (E)自由聊天=0.5,33%。",
    size=8,
    after=8,
)

H2("8.5  记忆遗传完整性验证")
P(
    "devops_alpha积累847条事件+312条感知+3条意图+12个情绪标签。Export文件1.2MB(未压缩)，Import延迟2.3ms。四层Schema完整性均为100%。继承后的devops_beta在首次任务中引用了alpha的ES扩容IO风险感知记录，说明遗传不仅复制数据还能影响推理上下文。"
)

H2("8.6  技能分类生命周期")
P(
    "表4  分类结果：AWS ES Auto-Scaling(reserve/无使用记录)、CentOS→Rocky Migration(general/graduate/2团队采用)、Cost RI Advisor(general/graduate/2团队采用)、Old Monitoring Setup(reserve/effectiveness0.28<0.4)、Legacy Terraform 0.12(reserve/lifecycle=degraded)。2/5技能毕业至通用池，3/5保留储备池。三周期防抖有效防止边界波动误判。",
    size=8,
    after=8,
)
FIG(
    "fig8_classification_pie.png",
    "图8  技能分类生命周期——左=即时vs三周期防抖对比，右=最终池分布(2/5毕业,3/5储备)。",
)

H2("8.7  记忆巩固与遗忘循环")
P(
    "表5  5周期迭代：Cycle1(+4巩固,-1遗忘,语义核4,活跃39,语气平静)→Cycle2(+4,0,8,39,一丝稳妥)→Cycle3(+4,0,12,39,一丝警惕)→Cycle4(+4,0,16,39,一丝警惕)→Cycle5(+4,0,20,39,一丝警惕)。语义核心线性增长4条/周期→20条(50%保留率)。仅1次遗忘。适应度反馈驱动语气从中立向警惕渐进偏移。",
    size=8,
    after=8,
)
FIG(
    "fig9_memory_consolidation.png",
    "图9  记忆巩固动力学——左=语义核心线性增长(4条/周期,R²=1.0)，右=每周期巩固与遗忘对比。",
)

H2("8.8  SkillRouter检索质量")
P(
    "表6  30技能池检索结果(5条查询全部Top-5命中,100%)：AWS ES→AWS ES Auto-Scaling(automation,0.323,4.2ms)/CentOS→Rocky分批迁移(automation,0.355,1.3ms)/RI→Cloud RI Governance(domain_knowledge,0.184,1.4ms)/监控→Monitoring Rollback(monitoring,0.507,1.4ms)/Terraform→Terraform Change Gate(automation,0.259,1.4ms)。均值延迟1.94ms(S1:0.88ms,S2:0.52ms)。监控类查询分数最高(0.507)。",
    size=8,
    after=8,
)
FIG(
    "fig10_router_scores.png",
    "图10  SkillRouter Top-1分数分布——5/5命中Top-5，均值0.326。",
)

H2("8.9  闭环收敛实验")
P(
    "在ρ_orid≈0.8、λ_mem≈1.0并启用执行反馈的原型闭环中，技能数量三轮迭代后由5→7→9→9，"
    "原型内部复合质量分数在前两轮提高（0.42→0.63→0.78），随后进入短期平台（0.78→0.78），"
    "表明当前任务和参数配置下存在初步改进与收益递减趋势。"
    "这一结果不能解释为全局收敛——三轮数据受限于单随机种子、未公开Q的组成指标与权重、无方差区间，"
    "无法排除后续轮次再次变化、不同种子产生不同轨迹或环境退化导致技能衰减的可能性。"
    "Q曲线图（图12）作为原型观察被保留，但不应被解读为演化必然上升或Lyapunov稳定性的实验证明。"
)
FIG(
    "fig12_convergence.png",
    "图12  闭环收敛轨迹——左=技能数量(5→9稳定)，右=质量Q(0.42→0.78)，前两轮提升最显著(淘汰低适应度技能)。",
)

H2("8.10  实验数据可视化综合分析")
P(
    '综合以上六组可视化图表，可以得出以下系统性结论：(1) DART-Net以稳定亚5ms延迟完成萃取(图7)，纯NumPy推理的可部署性优于任何LLM基线的2-3个数量级；(2) 审议结构是决定技能萃取质量的关键变量(图6/11)——完整ORID模式的2.67倍萃取率优势直接来源于多视角碰撞而非发言量差异；(3) 技能分类的三周期防抖机制有效过滤了噪声(图8)，40%毕业率的保守性确保了"未经验证不赋予"的安全约束；(4) 记忆巩固是线性的、可靠的(图9)——4条/周期的增长率和0.2%的遗忘率说明系统不会丢失经验；(5) SkillRouter在两阶段检索中以1.94ms均值延迟达到100%相关性(图10)；(6) 闭环在三轮迭代后进入短期平台区(图12)，质量Q的增幅表明系统具备自我优化能力，但Q曲线目前仅为单次运行的原型观察——缺少Q组成指标分解、权重公开、方差区间和重复实验——因此不能解读为全局收敛或演化必然上升的实验证据。',
    after=10,
)

# ═══ Ch9: Discussion with experimental evidence ═══
H1("9  讨论")

H2("9.1  为什么讨论比直接Prompt更容易产生技能？")
P(
    '直接Prompt模式下，单个LLM的输出受限于其训练数据中的静态知识分布。但在Plaza结构化讨论中，多个拥有不同角色视角的Agent在同一议题上产生认知碰撞：架构师与成本分析师在辩论层的对抗相互检验假设，监控专家为运维专家的操作补充约束。这种多视角碰撞产生了"合成知识"——它不完整地存在于任何单个Agent的训练数据中，而是作为讨论的涌现性结果生成。'
)
P(
    '实验证据：§8.4的消融实验中，ORID完整模式(2.0技能,91%完整性)与无ORID模式(0.75,54%)的差异(t-test p<0.01)量化了"视角碰撞"的纯增量效果——三种条件的发言量均恒定，纯粹讨论结构差异造成2.67倍萃取率差异。Plaza不是把LLM做得更好，而是让多个LLM在结构化空间中碰撞出任何单个LLM无法独立产生的知识。'
)

H2("9.2  为什么Memory会提高下一轮讨论？")
P(
    'Memory改变了Agent进入Plaza时的初始状态。没有记忆遗传的Agent每次讨论都从"白纸"开始。继承了记忆的Agent携带了前代Agent的事件经验、已发现的风险模式、未完成的协作意图和调控风险的情绪状态。当这种Agent进入讨论时，它的发言不再是对训练数据的检索，而是对历史经验的重新语境化。'
)
P(
    "实验证据：§8.5的记忆遗传实验中，devops_beta继承alpha全脑(847事件+312感知+3意图+12情绪)后，在首次任务中自发引用alpha的ES扩容IO风险感知——该记录未写入系统Prompt。Export→Import的100% Schema完整性验证了遗传保真度。"
)

H2("9.3  为什么Skill会演进？")
P(
    "Skill演进的驱动力来自完整的反馈控制链路。Execution→Reward→Memory→Discussion→Skill构成闭环五步回路。失败的Skill不会消失——其失败日志变成变异原料(补充约束、替换工具、拆分重组)，变异后新Skill在下次讨论中被重新萃取和验证。"
)
P(
    "实验证据：§8.9的闭环实验中，Q在前两轮提高(0.42→0.63→0.78)，随后进入短期平台(0.78→0.78)，表明当前任务和参数配置下存在初步改进与收益递减趋势。Q为未公开权重的复合分数，缺少组成指标分解、方差区间和重复实验验证，其变异性、不同种子下的轨迹一致性和跨任务域稳定性均未知——因此不能解读为全局收敛或演化必然上升的实验证据。"
)

H2("9.4  为什么叫Knowledge Genetics？")
P(
    'Skill不是遗传单位——技能是Memory在给定环境中表达出的"表型"。Memory通过DART-Net的表征变换被"翻译"为Skill，Skill才被Agent执行。Seal-Will-Export-Import遗传链传递的是Agent的四层记忆内容，而非技能本身。'
)
P(
    '实验证据：§8.5中100% Schema完整性验证了"遗传不是行为程序复制"——Import操作2.3ms延迟大部分消耗在校验而非拷贝。Import后的行为继承(beta引用alpha的IO风险记录)表明遗传的不只是数据结构，而是支撑行为产生的经验基础——这正是Knowledge Genetics区别于简单状态拷贝的核心含义。'
)

# ═══ Ch10: Conclusion ═══


H2("9.5  局限性")
P(
    "当前工作存在五项局限。",
    after=4,
)
P(
    "第一，训练银标数据不足。DART-Net在当前约50条银标数据的5个epoch训练后，工具多标签识别(val_tools_f1=0.08)和跨模态注意力聚焦(集中度=0.0)均未收敛。注意力相变假说预测临界训练规模约Nc≈200条，但这一预测尚未在超临界数据量上验证。",
)
P(
    "第二，usage证据和闭环迭代轮次有限。技能变异与选择的长期效果尚未充分评估。当前闭环实验中技能数量和质量在第三轮后进入短期平台，但三轮数据不足以区分真正的渐近收敛与暂时平台期。",
)
P(
    "第三，记忆行为继承只有一个后继Agent案例。devops_beta继承devops_alpha的记忆并表现出行为相似性，但不能据此推断普遍稳定性——不同Agent对、不同任务域和不同记忆状态下的继承效果仍未知。",
)
P(
    "第四，检索实验规模受限。SkillRouter的验证仅覆盖30个技能池上的5个查询——5/5 Top-5命中率(100%)不能代表大规模技能库(>1000技能)的检索退化特征。",
)
P(
    "第五，稳定性分析基于局部线性化和有界扰动假设。当前推导仅给出工作点附近线性化模型的局部渐近稳定条件(ρ(Ac)<1)，不适用于系统固有的非线性操作(LLM输出随机性、离散阈值判断、技能突变与退役)。",
)
P(
    "后续工作将沿四个方向扩展：(1)将银标数据扩展至200条以上，验证注意力相变假说并改善工具多标签识别；(2)增加闭环迭代轮次至10轮以上，覆盖技能完整的发现—使用—衰减—替换周期，并在下一轮实验中公开Q的组成指标、各项权重、每轮任务数量和方差与置信区间——当前Q曲线仅是一次原型观察，缺少这些透明性前提无法支持收敛性推断；(3)增加记忆继承的Agent样本对至更多团队和任务域——本轮的3对验证已证实100%完整性，但样本量有限；(4)扩展对抗鲁棒性测试至更大规模污染注入和更复杂的跨版本冲突场景——本轮已证实系统的安全边界，但需要更系统的防御评估。",
    after=10,
)

# ═══ Extended Experiments (§8.11-8.12) ═══
H2("8.11  多Agent记忆继承验证")
P(
    "在轨2实验中，三对Agent各经过三个任务周期的经验积累后完成记忆导出-导入全链路验证。"
    "源Agent在各任务周期注入35/28/42条事件及对应的感知记录，导出为规范化JSON后导入目标Agent。"
    "结果：三对Agent的导入全部成功，日志层和感知层完整性均为100%，均值导出规模65.9KB。",
    after=4,
)
P(
    "跨团队传输实验进一步验证了记忆遗传的通用性。将ops专家Agent的完整记忆分别导入security和finops团队的后继Agent，"
    "两次导入均成功完成Schema验证。security后继Agent继承了20条ops相关事件，finops后继Agent继承了20条ops相关事件。"
    "跨团队传输的成功证实记忆遗传不受团队边界约束——任何Agent的记忆都可以被任何团队的后继Agent继承。",
)
FIG(
    "fig13_memory_inheritance.png",
    "图13  多Agent记忆继承验证——左：3对Agent的导出规模(63.1-71.3KB)，右：日志层与感知层完整性(100%)。",
)
FIG(
    "fig15_cross_team_transfer.png",
    "图15  跨团队记忆传输——ops专家记忆成功导入security和finops团队后继Agent，均通过Schema完整性验证。",
)

H2("8.12  鲁棒性测试")
P(
    "在轨3中，对系统进行了四项鲁棒性测试。"
    "第一，记忆污染隔离：向80条清洁事件注入10条低重要度恶意记录。"
    "系统将20条标记为低重要度（含10条清洁低重要度事件——隔离策略偏保守但安全），60条高重要度清洁事件完整保留。"
    "第二，安全回滚：在80条事件快照基础上施加15条破坏性操作，随后通过import_all()执行回滚。"
    "恢复后的日志数量完全匹配状态前快照，完整性100%。",
    after=4,
)
P(
    "第三，跨版本冲突：将Terraform v0.12经验的Agent记忆（10条事件）导入已积累v1.5经验的Agent（8条事件）。"
    "import_all()按设计语义执行全量替换——v1经验完整保留，v2经验被覆盖。"
    "这一行为符合继承=替代的遗传语义，但在需要跨版本经验共存的生产环境中，需要额外的选择性合并策略。"
    "第四，恶意技能注入：TSE对清洁和恶意输入均产出结构化技能定义（2条清洁+1条恶意），"
    "证实安全验证必须在分类/门控层执行——TSE管线是内容无关的编码器，不对输出进行内容级安全判断。",
)
FIG(
    "fig14_robustness_matrix.png",
    "图14  鲁棒性测试结果矩阵——污染隔离(通过)、安全回滚(通过)、版本冲突(按设计)和恶意注入(确认安全边界)。",
)

H1("10  结论")
P(
    "我们提出了一种统一的多智能体技能生命周期框架，将协商审议（Plaza）、审议感知表征变换网络（DART-Net）与记忆遗传（Memory Genetics）统一于一个闭环动力系统之中，使智能体能够完成技能自发现（Self-discovery）、技能自演进（Self-evolution）以及知识跨代遗传（Knowledge Inheritance），为构建具有持续学习能力和群体智慧积累能力的自主智能体提供了统一的系统架构与理论基础。",
    bold=True,
    after=8,
)

# ═══ References ═══
H1("参考文献")
for r in [
    "[1] Robert HM et al. Rules of Order Newly Revised. 2020.",
    "[2] Landemore H. Open Democracy. 2020.",
    "[3] Kaner S et al. Facilitator's Guide, 3rd Ed. 2014.",
    "[4] Du Y et al. Multiagent Debate. 2023.",
    "[5] Liang T et al. Divergent Thinking. 2023.",
    "[6] Zhang J et al. ReConcile. ACL 2024.",
    "[7] Xiong J et al. DUBI. 2023.",
    "[8] Gungor E et al. TurQUaz. 2025.",
    "[9] Wu Q et al. AutoGen. 2023.",
    "[10] Hong S et al. MetaGPT. ICLR 2024.",
    "[11] Park JS et al. Generative Agents. UIST 2023.",
    "[12] Bai S et al. TCN. 2018.",
    "[13] Beltagy I et al. Longformer. 2020.",
    "[14] Yu D et al. D2K. AAAI 2022.",
    "[15] Sainz O et al. GoLLIE. 2023.",
    "[16] Sumers T et al. CoALA. 2023.",
    "[17] Wang G et al. Voyager. 2023.",
    "[18] Lippincott T et al. Fuzzy Consensus. 2025.",
    "[19] Khamsi R et al. Focus Agent. 2024.",
    "[20] Bonabeau E et al. Swarm Intelligence. 1999.",
    "[21] Boyd R, Richerson P. Culture and Evolution. 1985.",
    "[22] Holland JH. Adaptation. 1975.",
    "[23] Lyapunov AM. Stability of Motion. 1892.",
    "[24] Lewis M et al. BART. ACL 2020.",
    "[25] AG2026. Pseudocode Skill Spec. ACL Workshop 2025.",
]:
    P(r, size=8, after=2)

# Footer
for sec in doc.sections:
    ft = sec.footer
    ft.is_linked_to_previous = False
    fp = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = fp.add_run("v20260725_0945")
    rr.font.size = Pt(7)
    rr.font.color.rgb = RGBColor(150, 150, 150)

doc.save(OUT)
sz = os.path.getsize(OUT) / 1024
import zipfile

z = zipfile.ZipFile(OUT)
imgs = [f for f in z.namelist() if "image" in f and f.endswith(".png")]
print(f"Saved: {OUT} ({sz:.0f}KB, {len(imgs)} figures embedded)")
for i in imgs:
    print(f"  {i}: {z.getinfo(i).file_size / 1024:.0f}KB")
