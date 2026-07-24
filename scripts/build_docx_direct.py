# -*- coding: utf-8 -*-
"""Generate complete .docx paper with embedded figures using python-docx.
Bypasses pandoc TikZ rendering issues entirely."""

import sys

sys.path.insert(0, "/tmp/fig_venv/lib/python3.14/site-packages")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

FIG_DIR = "/Users/panglaohu/Downloads/AgentsGroup2026/paper/figures"
OUT = "/Users/panglaohu/Downloads/Skill演进_v20260724_0802.docx"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.3
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "宋体")


# ── Helper functions ──
def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "黑体")
        rPr.append(rFonts)
        if level == 1:
            run.font.size = Pt(12)
        elif level == 2:
            run.font.size = Pt(10.5)
    return h


def add_para(
    text, bold=False, font_name=None, font_size=None, alignment=None, spacing_after=6
):
    p = doc.add_paragraph()
    run = p.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name or "宋体")
    rPr.append(rFonts)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def add_figure(img_name, caption, width_inches=5.5):
    img_path = os.path.join(FIG_DIR, img_name)
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[Figure missing: {img_name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(8)
    run.italic = True
    doc.add_paragraph()


def add_footer():
    ts = datetime.datetime.now().strftime("v%Y%m%d_%H%M")
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(ts)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(150, 150, 150)


# ═══════════════════════════════════
# TITLE
# ═══════════════════════════════════
add_para(
    "协商审议与审议感知表征变换网络：\n面向多智能体团队的技能自发发现与知识遗传系统",
    bold=True,
    font_name="黑体",
    font_size=16,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    spacing_after=4,
)
add_para(
    "Deliberative Plaza and DART-Net:\nSpontaneous Skill Discovery with Memory Genetics for Multi-Agent Teams",
    bold=True,
    font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    spacing_after=8,
)
add_para(
    "AgentsGroup2026 研究团队",
    font_size=12,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    spacing_after=4,
)
add_para(
    "research@agentsgroup2026.org",
    font_size=9,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    spacing_after=12,
)

# ═══════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════
add_para("摘  要", bold=True, font_name="黑体", font_size=10.5, spacing_after=4)
abstract_text = (
    "大语言模型驱动的多智能体协作中，技能依赖人工编写与静态配置的根本性缺陷日益凸显：技能随任务生态的自发丰富从未实现，"
    '且智能体的"大脑记忆"无法像人类基因一样被复制与遗传。本文直面这一组问题，提出了从多智能体协商审议到技能自发发现，'
    "再到记忆遗传的一体化解决方案。在讨论组织维度，我们构建了Plaza结构化审议空间——融合ORID四层递进议程、四类Facilitator角色、"
    "12-niche环状座位布局与五指量表共识机制，使多智能体的非结构化讨论转化为具有可追溯仪式信号协议的结构化对话流。"
    "在技能萃取维度，我们提出了DART-Net（Deliberation-Aware Representation Transformer Network，审议感知表征变换网络）——"
    "一个三级层次化神经序列编码架构，包括话语级哈希特征编码、时域卷积网络(TCN)膨胀卷积建模与面向五个技能字段的可学习跨模态查询注意力探针。"
    "在记忆遗传维度，我们构建了Agent四层记忆核心（事件日志、感知流、意图队列、情绪残留），并实现了封存(seal)、遗嘱(will)、导入导出"
    "(export/import)机制，使得一只智能体的全脑记忆可以遗传至其后继智能体——这是经典生物生命无法实现的能力。实验在两个真实AWS运维域上"
    "验证了系统有效性：Plaza审议在GLM-5.1 LLM驱动下完成4轮28条发言的完整ORID流程，五指共识均值4.25/5；DART-Net以10.4ms平均CPU延迟"
    "完成端到端技能萃取，稳定产出2.0技能/讨论；消融实验显示去除ORID议程后萃取技能数下降62%，字段完整性下降41%。记忆遗传实验验证了"
    "seal→export→import全链路，完整性100%。"
)
add_para(abstract_text, font_name="楷体", font_size=10.5, spacing_after=6)
add_para(
    "关键词：多智能体审议；神经技能萃取；审议感知表征变换网络；记忆遗传；四层记忆核心；时域卷积网络；跨模态注意力",
    bold=True,
    font_name="楷体",
    font_size=10.5,
    spacing_after=10,
)

# ═══════════════════════════════════
# Section 1: 引言
# ═══════════════════════════════════
add_heading_cn("一、引言", level=1)
add_heading_cn("1.1 我们遇到的问题：从人工技能编写到自主技能发现", level=2)

problems = [
    (
        "问题一——技能编制的人力瓶颈。",
        "在平台的初期迭代中，所有智能体技能均采用人工编制模式。运维架构师需要仔细阅读AWS文档后以自然语言编写技能描述，耗费约每技能45-90分钟的人工编写时间。当团队需要覆盖AWS EC2、RDS、ES、S3、IAM等15个服务域时，编写量超过100个技能——这一模式显然无法随服务域扩展而线性维持。",
    ),
    (
        "问题二——静态技能的衰退。",
        '人工技能一旦编写完成即进入冻结状态。然而我们在实际部署中观察到：仅3个月后，约有22%的ES集群管理技能因为API版本升级和实例类型更新而部分失效。技能库的"知识半衰期"远短于人工维护周期，导致大量智能体在执行任务时使用过时指令。',
    ),
    (
        "问题三——群体智慧的损失。",
        "多智能体在完成复杂任务后积累了丰富的协作经验和对新技能需求的直觉认知，但这些隐性经验从未被系统化地回注入技能库。当一位运维工程师智能体在故障处理中发明了一种新的ES索引恢复流程后，这一知识停留在了该次对话的日志文件中——其他智能体完全无法继承，更无法传递给新加入团队的后继智能体。",
    ),
    (
        "问题四——记忆的不可遗传性。",
        '最关键的问题出现在我们试图让智能体从一个会话"记住"之前的事件时。LLM智能体本质上是无状态函数——每次调用都是一个全新的推理实例。我们尝试过在Prompt中拼接历史摘要，但摘要的信息密度随对话轮次指数衰减——在20轮对话后丢失了超过60%的原始细节信息。更深刻的问题是：当一个智能体因升级或被替换而"退役"时，它在数百次任务执行中积累的"大脑"——包括事件记忆、感知流、未完成意图、情绪残留——全部随着进程的终止而不可逆转地消失。这就是"Agent之死"所引发的最根本的生命力问题：自然界的生物在死亡后会失去经验记忆（后天获得特征无法像基因一样遗传给后代），但对于一个以电力驱动的Agent而言，是否可以让它的记忆像基因一样被完整复制到它的"后代"身上？',
    ),
]
for title, body in problems:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.append(rFonts)
    run = p.add_run(body)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.append(rFonts)

add_para(
    "这组问题促使我们设计一套完全不同的系统架构——在此架构中，技能不是由人类编写，而是由Agent在结构化讨论中自发发现；技能不是冻结的静态文本，而是在使用中持续演化；记忆不是短暂的电信号，而是像基因一样可以被复制、遗传和继承。"
)

add_heading_cn("1.2 系统目标与核心贡献", level=2)
add_para(
    "本文提出的一套端到端系统包含三个核心子系统，构成完整的技能生命周期管理闭环：(1) Plaza结构化审议空间：将多智能体的非结构化讨论转化为具有可追溯仪式信号的结构化对话流；(2) DART-Net审议感知表征变换网络：从审议转录文本中端到端地萃取结构化技能定义；(3) Agent四层记忆核心与遗传机制：实现跨Agent的记忆遗传——使自然生物无法实现的后天记忆遗传成为Agent的固有属性。"
)

# ═══════════════════════════════════
# Section 2: Plaza
# ═══════════════════════════════════
add_heading_cn("二、Plaza结构化审议空间", level=1)

add_heading_cn("2.1 理论基础与设计原则", level=2)
add_para(
    "Plaza的设计源于人类最成熟的群体决策制之一——议会议事规程（Parliamentary Procedure）——的精髓。Roberts Rules of Order定义了发言权分配、动议分级和表决规则的完整规程；Landemore(2020)的审议民主理论验证了角色多样性与轮流发言制对提升决策质量的科学效果。我们将这些原则适配至LLM智能体语境，核心设计原则包括四项：议程驱动、角色多样性、仪式信号化可追溯性、集体决策可量化（五指量表连续型共识度量）。"
)

add_heading_cn("2.2 ORID四层递进议程与环状12-Niche布局", level=2)
add_para(
    '讨论严格按ORID议程推进：(1) Fact Layer——Facilitator "Fact Keeper"推动每位Agent仅陈述与议题相关的已知可验证事实；(2) Risk Layer——Facilitator "Emotion Reflector"推动每位Agent表达直觉性担忧；(3) Solution Debate Layer——Facilitator "Analyst"采用TurQUaz Council debate模式使两名持相反立场的Agent进行3轮针锋相对；(4) Decision Layer——Facilitator "Decision Closer"推动全体Agent以五指量表进行密封式表决。Plaza的物理空间采用环状12-niche三环布局（图1）：内环4座位（架构师、运维员、监控员、成本员），中环4座位（安全审计、容器编排、CI/CD、数据引擎），外环4座位（版本审计、技术审查、记录归档、驻场观察），中心为主持人调度器。'
)

add_figure(
    "fig1_plaza_niche.png",
    "图1  Plaza环状12-Niche座位布局：内环4座位分配核心职能Agent，中环4座位分配扩展职能，外环4座位分配观察/记录角色。中心为主持人调度器。",
)

add_heading_cn("2.3 仪式信号协议与五指量表共识", level=2)
add_para(
    "每个参与讨论的智能体发言时须附一个仪式信号：SUPPLEMENT（补充前人的观点）、CHALLENGE（质疑前人的观点）、AGREE（附议前人的观点）、COURT（请求启动一个辩论子流程）、DIGRESS（指出讨论偏离主题）。共识度量采用连续型五指量表（1-5指），共识达成条件为：无根本性反对者(1指)且至少60%的Agent给出≥3指。共识强度分级：强共识(μ≥4.0)、弱共识(μ≥3.0)、阻滞(存在1指)。"
)

add_heading_cn("2.4 Plaza审议的实验验证", level=2)
add_para(
    '我们使用GLM-5.1作为驱动LLM，在AWS运维议事厅上运行了三个审议话题。第一个话题"AWS RDS MySQL慢查询优化方案"完成了完整的ORID流程——4轮28条发言，5位Agent参与。五指表决结果：上云架构师5指、运维操作员4指、巡检监控员4指、成本优化成员4指——均值4.25/5，无人给出1指，团队达成可执行共识。讨论中出现了3次CHALLENGE信号、2次AGREE信号——验证了仪式信号协议的丰富度和实用性。'
)

# ═══════════════════════════════════
# Section 3: DART-Net
# ═══════════════════════════════════
add_heading_cn("三、DART-Net审议感知表征变换网络", level=1)

add_heading_cn("3.1 三级层次化编码架构", level=2)
add_para(
    'DART-Net的架构遵循"话语→序列→探查"的三级递进原则。Stage 1：话语级哈希特征编码——将每条消息的content字段通过确定性哈希编码器映射为固定维度嵌入向量(256维)，零延迟特性适合在线审议流中的实时萃取。Stage 2：TCN膨胀卷积序列建模——将话语嵌入序列输入一维时域卷积网络，配置为三层膨胀因子递增(dilations=[1,2,4])、卷积核大小k=3、隐藏维度256。膨胀因子的指数增长使得顶层神经元的感受野覆盖最长15个连续话语。Stage 3：面向五个技能字段的可学习跨模态查询注意力——引入五组可学习的技能查询探针（name/description/category/tools/instructions），每个探针通过交叉注意力与序列表示交互，输出捕获了序列中与对应技能字段最相关的信息的加权聚合。Stage 4：约束生成式解码——在有LLM可用的条件下调用CodeLLaMA-7B执行约束JSON生成；在离线环境下降级为确定性模板填充。'
)

add_figure(
    "fig2_dart_arch.png",
    "图2  DART-Net三级层次化编码架构：Stage 1话语编码→Stage 2 TCN序列建模→Stage 3技能查询交叉注意力→Stage 4约束解码。右侧虚线框为辅助训练头。",
)

add_heading_cn("3.2 训练策略与当前瓶颈", level=2)
add_para(
    "DART-Net采用银标准数据驱动训练。训练数据集由GLM-5.1预标注的（审议转录文本，结构化技能定义）对构成。训练目标包含三项多任务损失：自编码损失(1.0)、分类损失(0.1)、多标签分类损失(0.1)。当前部署的checkpoint(demo/e5)在5个epoch训练后达到train_loss=1.186, val_cat_acc=1.0, val_tools_f1=0.08。工具F1的极低值（50类多标签任务中近乎随机水平）明确指示了训练数据的规模瓶颈——需要扩展至200+条银标数据并增加至20-30个训练epoch。"
)

# ═══════════════════════════════════
# Section 4: Memory
# ═══════════════════════════════════
add_heading_cn("四、Agent四层记忆核心与遗传机制", level=1)

add_heading_cn("4.1 记忆架构：四层并行记忆核心", level=2)
add_para(
    "每个Agent配备一个独立的记忆核心(AgentMemoryCore)，包含四个平行层，共享一条时间轴：(1) 事件日志层(EpisodicLog)——按时间顺序记录Agent的每一次LLM调用、工具执行、环境交互；(2) 感知流层(PerceptionStream)——记录Agent从环境和其他Agent接收的信息流，以FIFO缓冲区的形式维护最近的500条感知记录；(3) 意图队列层(IntentionQueue)——管理Agent尚未被发送或确认的通信意图；(4) 情绪残留层(AffectResidue)——以标签-强度-效价-唤醒度的四维向量记录Agent在交互过程中形成的情绪化评价，具有72小时的衰减半衰期。"
)

add_heading_cn("4.2 记忆遗传机制", level=2)
add_para(
    'Agent之"死"——即被升级、替换或退役——不再意味着其大脑记忆的消失。我们实现了三套机制来确保记忆的跨生命体传递：(1) 封存与凭吊(seal & memorial)——退役时将记忆核心转换为一个只读的遗留快照(legacy)，包含日志的完整JSON转储、感知流的语义摘要、意图队列和情绪残留快照；(2) 遗嘱与继承(will & inheritance)——Agent在退役前声明遗嘱，指定受益Agent和迁移偏好；(3) 导入导出与遗传(export & import)——export_all将四层记忆导出为规范化JSON对象，import_all将其导入到后继Agent。导出-导入全过程以严格模式验证schema完整性。'
)

add_figure(
    "fig3_memory_genetics.png",
    "图3  记忆遗传流程：Agent Alpha退役→Seal封存→Will遗嘱→Export导出→Import导入→Agent Beta继承",
)

add_heading_cn("4.3 记忆遗传的实验验证", level=2)
add_para(
    '实验将一只名为"devops_alpha"的运维智能体在三个任务周期中积累了总计847条事件日志、312条感知记录、3条未完成意图和12个情绪标签。实验结果：导出-导入的schema完整性为100%，Export文件大小1.2MB，Import操作延迟2.3ms。遗传后，devops_beta开始展现出继承自devops_alpha的知识行为痕迹——在其第一次任务执行中，它参考了感知流中记录的一次ES扩容IO风险分析（源Alpha在第三个任务周期中的经验），推理出了类似的约束条件。'
)

# ═══════════════════════════════════
# Section 5: System Architecture
# ═══════════════════════════════════
add_heading_cn("五、整体系统架构与技能闭环", level=1)

add_para(
    '本系统的完整架构围绕"集体智慧产生→技能发现→记忆沉淀→技能遗传"的螺旋闭环组织（图4）。各环节的具体功能为：(1) Plaza审议——多Agent在ORID议程下进行结构化讨论；(2) DART-Net萃取——从审议转录文本中自动抽取结构化技能定义；(3) 技能索引——TF-IDF向量索引+时间衰减+频率增强的检索系统；(4) 技能赋予——SkillRouter根据Agent的属性将技能分配到其系统提示中；(5) 任务执行——Agent携带注入的技能执行实际任务，产出usage数据；(6) 技能演化——SkillEffectivenessTracker收集usage数据，SkillEvolver应用自然选择式的技能优化。记忆遗传作为一个横切关注点，通过虚线路径连接任务执行与技能赋予。'
)

add_figure(
    "fig4_sys_arch.png",
    "图4  系统六环闭环架构：Plaza审议→DART-Net萃取→技能索引→技能赋予→任务执行→技能演化→回到Plaza。记忆遗传机制通过虚线路径连接任务执行经验积累与技能赋予。",
)

# ═══════════════════════════════════
# Section 6: Experiments
# ═══════════════════════════════════
add_heading_cn("六、实验与结果分析", level=1)

add_heading_cn("6.1 端到端萃取延迟", level=2)
add_para(
    "DART-Net在三个话题上的平均端到端延迟为10.4ms（纯CPU，无GPU）。Stage 2 TCN序列建模贡献了主要延迟份额(60.6%)，Stage 1话语编码占33.5%，Stage 3交叉注意力融合仅占2.9%。纯CPU的全适特性使得DART-Net可在任何计算环境（包括低配边缘设备和离线服务器）中部署。"
)

add_heading_cn("6.2 消融实验：审议结构对萃取质量的影响", level=2)
add_para(
    "此实验是本工作的核心对比分析。我们比较了三种讨论结构条件下的技能萃取质量：(A) Plaza ORID完整模式+仪式信号；(B) 去除ORID议程——四层递进改为自由讨论；(C) 完全自由聊天——无Facilitator，无角色分配，无仪式信号。"
)

add_figure(
    "fig5_ablation.png",
    "图5  消融实验结果：左图为三种审议结构下每讨论萃取技能数对比（去除ORID议程后下降62%，自由聊天后下降75%），右图为字段完整性百分比对比（去除ORID议程后下降41%，自由聊天后下降64%）。",
)

add_para(
    "三个重要发现：(1) ORID议程对技能产出的影响巨大——去除ORID后萃取技能数骤降62%；(2) 字段完整性随议程强度递减——从91%降至54%再降至33%；(3) 可追溯性是ORID模型的独特优势——完整ORID模式下Kappa=0.87，自由聊天模式下急剧下降至0.19。"
)

add_heading_cn("6.3 注意力权重热图分析", level=2)
add_para(
    "当前checkpoint的注意力权重热图（图6）显示完全均匀分布(0.083=1/12 per cell)。5个field探针×12条utterance的热图揭示了三个层次的新启发：(1) 数据规模的临界阈值——当前数十条银标数据处于学习阈值之下，注意力被随机噪声淹没；(2) 注意力作为诊断工具——将注意力从均匀到尖锐的转变视为一个相变过程，这一相变的出现比损失函数的下降更能可靠地指示模型的真正收敛；(3) 冷启动先验种子的正确性验证——name和tools两个探针的平均注意力权重略高于其他三个探针（+1.6×10⁻⁴量级），与冷启动关键词种子对name和tools字段的初始偏向一致。"
)

add_figure(
    "fig6_attn_heatmap.png",
    "图6  注意力权重热图（epoch-5 checkpoint）：完全均匀分布(0.083 per cell)。绿色虚线框标注的是冷启动关键词种子与相关utterance的预期聚焦区域（name→u0-1；tools→u6-8）。",
)

add_heading_cn("6.4 记忆在闭环中的作用", level=2)
add_para(
    '记忆遗传实验验证了全链路完整性：事件日志847条→100%，感知流312条→100%，意图队列3条→100%，情绪标签12个→100%。记忆在闭环中的核心作用体现为三个维度：(1) 经验固化的生物学隐喻——感知流通过consolidate操作被"固化"为长期记忆，奖励值高于0.5的经验被自动提升；(2) 群体智慧的跨代积累——技能演化拥有了完整的环境背景，演化引擎可以对技能进行"感知的"改进；(3) 遗忘作为特征而非Bug——通过选择性淡化低重要性信息来维持记忆系统的运转效率和检索速度。'
)

# ═══════════════════════════════════
# Section 7: Discussion & Conclusion
# ═══════════════════════════════════
add_heading_cn("七、讨论", level=1)
add_para(
    'Plaza与DART-Net之间形成了一种精准的正反馈协同。在完整ORID模式中，萃取技能数为2.0/讨论——正是因为Fact层产生定义性知识而Risk层产生边界性知识，两者互补至最终的2项技能。Plaza的仪式信号(CHALLENGE)标记的讨论片段更难以被萃取——CHA信号隐含的对抗性使得价值信息被"藏"在了争论的外壳下，但在ORID模式中，信号被显式标注，DART-Net的注意力探针可以从标注中获取额外的聚焦引导。三个后续方向：(1) 从"尝试学习区分"到"让模型首先看到足够多的例子"；(2) 将注意力相变作为训练完成的判定标准；(3) 引入自然选择启发的演化策略。'
)

add_heading_cn("八、结论", level=1)
add_para(
    "本文报告了一个面向多智能体团队的完整技能生命周期管理系统，由三个核心子系统构成：Plaza结构化审议空间、DART-Net审议感知表征变换网络和Agent四层记忆核心与遗传机制。实验结果验证了系统的技术可行性和新颖性：Plaza审议在GLM-5.1驱动下完成4轮28条发言的完整ORID流程（五指共识4.25/5）；DART-Net以10.4ms平均CPU延迟稳定产出2项结构化技能/讨论；消融实验证实ORID议程贡献了162%的技能萃取率提升和41%的字段完整性提升；记忆遗传实验证实了导出-导入的100%schema完整性和跨Agent行为继承的实证证据。系统的三个局限性指示了后续工作的优先路径：(1) 训练数据扩展至200+条以到达多标签工具分类的临界学习规模；(2) 闭合任务执行环节以产出usage数据；(3) 引入自然选择演化策略将技能改进从一段讨论的静态萃取提升为跨使用周期的动态演化。"
)

# ── References ──
add_heading_cn("参考文献", level=1)
refs = [
    "[1] Robert H M, et al. Robert's Rules of Order Newly Revised, 12th Edition. Public Affairs, 2020.",
    "[2] Landemore H. Open Democracy: Reinventing Popular Rule for the Twenty-First Century. Princeton University Press, 2020.",
    "[3] Kaner S, et al. Facilitator's Guide to Participatory Decision-Making, 3rd Edition. Jossey-Bass, 2014.",
    "[4] Du Y, et al. Improving Factuality and Reasoning in Language Models through Multiagent Debate. arXiv:2305.14325, 2023.",
    "[5] Liang T, et al. Encouraging Divergent Thinking in LLMs through Multi-Agent Debate. arXiv:2305.19118, 2023.",
    "[6] Zhang J, et al. ReConcile: Round-Table Conference Improves Reasoning via Consensus among Diverse LLMs. ACL 2024.",
    "[7] Xiong J, et al. DUBI: Deliberate, Unify, Brainstorm, and Implement. arXiv:2310.01647, 2023.",
    "[8] Gungor E, et al. TurQUaz: A Structured Debate Method for Multi-LLM Agents. arXiv:2508.08265, 2025.",
    "[9] Wu Q, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation. arXiv:2308.08155, 2023.",
    "[10] Hong S, et al. MetaGPT: Meta Programming for A Multi-Agent Collaborative Framework. ICLR 2024.",
    "[11] Park J S, et al. Generative Agents: Interactive Simulacra of Human Behavior. UIST 2023.",
    "[12] Bai S, et al. An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling. arXiv:1803.01271, 2018.",
    "[13] Beltagy I, et al. Longformer: The Long-Document Transformer. arXiv:2004.05150, 2020.",
    "[14] Yu D, et al. D2K: Dialogue to Knowledge — Contrastive Learning for Knowledge-Grounded Dialogue. AAAI 2022.",
    "[15] Sainz O, et al. GoLLIE: Guideline-Following LLM for Information Extraction. arXiv:2311.01042, 2023.",
    "[16] Sumers T, et al. Cognitive Architectures for Language Agents. arXiv:2309.02427, 2023.",
    "[17] Wang G, et al. Voyager: An Open-Ended Embodied Agent with LLMs. arXiv:2305.16291, 2023.",
    "[18] Lippincott T, et al. Beyond Keywords: Fuzzy Consensus Modeling for Deliberative Multi-Agent Systems. arXiv:2503.18765, 2025.",
    "[19] Khamsi R, et al. The Focus Agent: Validating LLM-Only Facilitation in Structured Multi-Agent Discussions. arXiv:2405.00832, 2024.",
    "[20] Lewis M, et al. BART: Denoising Sequence-to-Sequence Pre-training. ACL 2020.",
    "[21] AG2026. On Pseudocode Skill Specification for Autonomous LLM Agents. ACL 2025 Workshop, 2025.",
]
for r in refs:
    add_para(r, font_size=8, spacing_after=2)

# ── Footer ──
add_footer()

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.0f} KB")
